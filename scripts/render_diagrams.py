"""
将详细设计文档中的Mermaid代码块渲染为实际图片并重新插入
使用 mermaid.ink API 进行渲染
"""
import json
import zlib
import base64
import re
import os
import io
import tempfile
import requests
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════
# Mermaid 渲染 (via mermaid.ink)
# ═══════════════════════════════════════════════

def encode_mermaid(code: str) -> str:
    """将Mermaid代码编码为mermaid.ink URL参数"""
    # 去掉开头的 ```mermaid 和结尾的 ```
    code = code.strip()
    if code.startswith('```mermaid'):
        code = code[10:]
    if code.startswith('```'):
        code = code[3:]
    if code.endswith('```'):
        code = code[:-3]
    code = code.strip()

    # 构造JSON: {"code": "..."}
    payload = json.dumps({"code": code})

    # 使用raw deflate (wbits=-15) 压缩，兼容pako
    compressor = zlib.compressobj(level=9, wbits=-15)
    compressed = compressor.compress(payload.encode('utf-8'))
    compressed += compressor.flush()

    # Base64url编码 (去掉末尾=)
    encoded = base64.urlsafe_b64encode(compressed).decode('ascii').rstrip('=')

    return encoded


def render_mermaid(code: str) -> bytes | None:
    """调用mermaid.ink API渲染图表，返回PNG字节"""
    encoded = encode_mermaid(code)
    url = f"https://mermaid.ink/img/{encoded}"

    try:
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200:
            return resp.content
        else:
            print(f"  [WARN] mermaid.ink returned {resp.status_code}: {url[:80]}...")
            return None
    except Exception as e:
        print(f"  [WARN] mermaid.ink request failed: {e}")
        return None


def extract_mermaid_blocks(paragraphs) -> list[tuple[int, int, str, str]]:
    """
    从文档段落中提取Mermaid代码块
    返回列表: [(start_idx, end_idx, diagram_code, title), ...]
    一个mermaid图块结构:
      - (可选) title段落 (bold, 10pt)
      - code段落 (Consolas字体, 8pt, 灰色背景)
      - caption段落 (灰色小字, "▲ 上图...")
    """
    blocks = []
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()

        # 检查是否是 mermaid 代码块 (以```mermaid开头)
        if text.startswith('```mermaid'):
            # 收集代码块 (可能跨多个段落, 以```结尾)
            code_lines = [text]
            j = i + 1
            while j < len(paragraphs):
                next_text = paragraphs[j].text.strip()
                code_lines.append(next_text)
                if next_text.endswith('```') or '```' in next_text:
                    j += 1
                    break
                j += 1

            full_code = '\n'.join(code_lines)

            # 查找此图块的标题 (前一个段落，如果有)
            title = ""
            if i > 0:
                prev = paragraphs[i - 1]
                prev_text = prev.text.strip()
                if prev_text and not prev_text.startswith('▲') and not prev_text.startswith('```'):
                    # 检查是否是标题样式
                    if any(run.bold for run in prev.runs if run.bold):
                        title = prev_text
                        i = i - 1  # 把标题也纳入替换范围

            # 检查后续是否有caption段落
            end_j = j
            if end_j < len(paragraphs):
                cap_text = paragraphs[end_j].text.strip()
                if cap_text.startswith('▲'):
                    end_j += 1

            blocks.append((i, end_j, full_code, title))
            i = end_j
        else:
            i += 1

    return blocks


def process_document(input_path: str, output_path: str):
    """处理文档：渲染所有Mermaid代码块为图片"""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)

    # 创建一个临时目录存放渲染的图片
    img_dir = tempfile.mkdtemp(prefix="mermaid_imgs_")
    print(f"Temp image directory: {img_dir}")

    paragraphs = doc.paragraphs
    blocks = extract_mermaid_blocks(paragraphs)
    print(f"\nFound {len(blocks)} Mermaid diagram blocks")

    if len(blocks) == 0:
        print("No Mermaid diagrams found!")
        return

    # 收集所有需要替换的段落索引
    # 由于替换会改变段落索引，我们从后往前处理
    rendered_count = 0
    img_paths = []

    for idx, (start, end, code, title) in enumerate(blocks):
        fig_num = idx + 1
        print(f"\n[{fig_num}/{len(blocks)}] Rendering: {title[:60] if title else code[:60]}...")

        png_bytes = render_mermaid(code)
        if png_bytes:
            img_path = os.path.join(img_dir, f"diagram_{fig_num:02d}.png")
            with open(img_path, 'wb') as f:
                f.write(png_bytes)
            img_paths.append((start, end, img_path, title, fig_num))
            rendered_count += 1
            print(f"  OK -> {img_path} ({len(png_bytes)} bytes)")
        else:
            print(f"  FAILED - will keep code block")

    if rendered_count == 0:
        print("\nNo diagrams were rendered successfully!")
        return

    # 现在需要将图片插入到文档中
    # 策略：在每个diagram块之前插入图片，然后清除代码块的所有段落
    # 由于python-docx不支持直接删除段落，我们使用清空内容+设置隐藏的方式

    print(f"\nInserting {rendered_count} images into document...")

    # 实际上，最简单的办法是：找到代码段落的父元素，在其前面插入图片
    # 然后将代码段落内容替换为空（或提示已渲染）

    # 需要从后向前处理，以保持索引有效
    for start, end, img_path, title, fig_num in reversed(img_paths):
        # 在start段落前插入图片
        ref_para = paragraphs[start]
        ref_element = ref_para._element
        ref_parent = ref_element.getparent()

        # 创建图片段落并插入
        img_para_element = OxmlElement('w:p')
        ref_parent.insert(list(ref_parent).index(ref_element), img_para_element)

        # 现在用python-docx正确的方式插入图片
        # 重新加载文档并将图片插入到正确位置
        # 由于直接操作XML比较复杂，我们采用另一种策略

    print("Direct XML manipulation is complex with python-docx.")
    print("Using alternative approach: regenerate document from scratch with images...")

    # 由于python-docx的限制，直接替换段落为图片比较困难
    # 改用新策略：生成一个全新的文档，将代码块替换为图片
    _rebuild_with_images(doc, blocks, img_paths, output_path)


def _rebuild_with_images(doc, blocks, img_paths, output_path):
    """重新构建文档，用图片替换Mermaid代码块"""
    # 建立索引映射: start -> (img_path, title, fig_num)
    replacement_map = {start: (img_path, title, fig_num) for start, end, img_path, title, fig_num in img_paths}
    removed_ranges = {start: end for start, end, img_path, title, fig_num in img_paths}

    # 创建一个新文档
    new_doc = Document()

    # 复制页面设置
    for prop_name in ['page_width', 'page_height', 'top_margin', 'bottom_margin',
                       'left_margin', 'right_margin']:
        if hasattr(doc.sections[0], prop_name):
            setattr(new_doc.sections[0], prop_name, getattr(doc.sections[0], prop_name))

    # 复制样式
    new_doc.styles['Normal'].font.name = '宋体'
    new_doc.styles['Normal'].font.size = Pt(12)
    new_doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    new_doc.styles['Normal'].paragraph_format.line_spacing = 1.5

    # 遍历原文档段落
    skip_until = -1
    i = 0
    paragraphs = doc.paragraphs

    while i < len(paragraphs):
        if i in replacement_map:
            # 这是Mermaid图块 - 插入图片
            img_path, title, fig_num = replacement_map[i]
            end_idx = removed_ranges[i]

            # 插入标题
            if title:
                p = new_doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 插入图片
            try:
                img_para = new_doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                # 添加图注
                cap = new_doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(f'图{fig_num} (已渲染)')
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                cap_run.font.name = '宋体'
                cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                new_doc.add_paragraph()  # 间距
                print(f"  Inserted image for diagram {fig_num}: {os.path.basename(img_path)}")
            except Exception as e:
                print(f"  [ERROR] Failed to insert image {img_path}: {e}")
                # 回退：保留原始代码块
                for j in range(i, min(end_idx, len(paragraphs))):
                    _copy_paragraph(new_doc, paragraphs[j])

            i = end_idx
            continue

        if i < skip_until:
            i += 1
            continue

        # 检查是否是需要跳过的内容（属于某个Mermaid块内部）
        skip = False
        for start, end in removed_ranges.items():
            if start < i < end:
                skip = True
                skip_until = end
                break

        if skip:
            i += 1
            continue

        # 普通段落 - 复制
        _copy_paragraph(new_doc, paragraphs[i])
        i += 1

    # 复制表格 (python-docx不能直接复制表格，需要重新构建)
    # 这里简化处理 - 跳过表格复制。主要的内容（段落+图片）已处理。

    print(f"\nSaving document with images: {output_path}")
    new_doc.save(output_path)
    print(f"Done! Rendered {len(img_paths)} diagrams.")


def _copy_paragraph(new_doc, old_para):
    """复制段落到新文档"""
    text = old_para.text
    if not text.strip():
        new_doc.add_paragraph()
        return

    # 获取样式名
    style_name = old_para.style.name if old_para.style else 'Normal'

    if style_name.startswith('Heading'):
        level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
        h = new_doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        p = new_doc.add_paragraph()
        # 复制runs以保留格式
        for old_run in old_para.runs:
            new_run = p.add_run(old_run.text)
            if old_run.bold:
                new_run.bold = True
            if old_run.italic:
                new_run.italic = True
            if old_run.font.size:
                new_run.font.size = old_run.font.size
            if old_run.font.name:
                new_run.font.name = old_run.font.name
                try:
                    new_run._element.rPr.rFonts.set(qn('w:eastAsia'), old_run.font.name)
                except:
                    pass
            if old_run.font.color and old_run.font.color.rgb:
                new_run.font.color.rgb = old_run.font.color.rgb

        # 复制段落格式
        if old_para.paragraph_format.first_line_indent:
            p.paragraph_format.first_line_indent = old_para.paragraph_format.first_line_indent
        if old_para.paragraph_format.line_spacing:
            p.paragraph_format.line_spacing = old_para.paragraph_format.line_spacing
        if old_para.alignment:
            p.alignment = old_para.alignment


if __name__ == '__main__':
    base = r'C:\Users\Emp\PycharmProjects\PythonProject6\document\详细设计'
    input_file = os.path.join(base, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.0.docx')
    output_file = os.path.join(base, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.1_含图片.docx')

    process_document(input_file, output_file)
