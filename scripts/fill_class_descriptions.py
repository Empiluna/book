#!/usr/bin/env python3
"""
Fill in the missing class detailed descriptions in the design document.
Each "详细描述" placeholder gets replaced with a comprehensive class-level
description followed by "详细描述如下：" and the existing table.
"""
import copy
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ── Content mapping: key = section identifier, value = list of paragraph texts ──
# Each list represents the multi-paragraph descriptive text for that class

CONTENT = {
    # ─── API Controllers (控制层) ───
    "UserAPI": [
        "UserAPI是用户管理模块的RESTful API路由组，路由前缀为/api/v1/user，负责提供用户注册、登录认证、画像获取、阅读行为记录、收藏管理和评分功能的HTTP端点。该路由组是系统与前端交互的入口层，所有端点通过FastAPI的Depends依赖注入机制获取数据库会话和当前用户身份。",
        "设计要点：注册接口通过Pydantic的UserRegister模型进行请求验证，在写入MySQL前校验用户名和邮箱的唯一性；登录接口采用bcrypt密码哈希验证，成功后签发24小时有效的JWT Token；/profile接口是模块间接口契约的关键实现——模块三（推荐引擎）通过此接口获取用户画像数据（标签偏好权重向量、偏好作者和标签的ID列表、高分图书ID列表），该接口返回的UserProfileForRecommend结构是模块一与模块三之间的数据契约；评分接口采用Upsert模式（存在则更新、不存在则新建），评分后自动触发图书平均分的重算（通过_update_book_avg_rating内部方法聚合user_ratings表实现）；阅读进度接口同样采用Upsert模式，支持多端同步。",
        "涉及的数据表：users（用户认证）、reading_history（阅读历史）、search_logs（搜索日志）、bookmarks（收藏）、user_ratings（评分）、reading_progress（阅读进度）。所有接口除/register和/login外均需携带有效JWT Token。"
    ],
    "GraphAPI": [
        "GraphAPI是知识图谱模块的RESTful API路由组，路由前缀为/api/v1/graph，负责提供基于Neo4j图数据库的图谱路径查询、实体与关系管理、可视化子图获取和统计功能。该路由组是系统知识图谱能力对外的统一接口层，底层通过GraphService（静态方法类）封装Cypher查询逻辑。",
        "设计要点：POST /paths接口是模块二（知识图谱）向模块三（推荐引擎）提供的核心契约接口——接收源图书ID和可配置的路径权重，通过GraphService.find_paths()执行五类Cypher路径查询（同作者/同标签/同系列/同出版社/多跳），返回按final_score降序排列的候选图书列表及其路径详情；实体管理和关系管理接口需要管理员权限（通过get_admin_user依赖注入进行角色校验），使用Cypher的MERGE语义实现幂等的创建或更新操作；可视化接口支持可变深度（1~4层）的子图遍历，返回去重的节点集合（含标签类型和属性）和边集合（含关系类型），直接适配前端ECharts/D3.js的图渲染数据格式。",
        "涉及的数据源：Neo4j图数据库（5种节点类型：Book/Author/Tag/Publisher/Series，5种关系类型：AUTHORED/TAGGED/PUBLISHED/SERIES_OF/SIMILAR）。图谱统计数据通过Cypher的count()聚合函数实时计算。"
    ],
    "RecommendAPI": [
        "RecommendAPI是个性化推荐模块的RESTful API路由组，路由前缀为/api/v1/recommend，负责提供首页个性化推荐流、相似图书发现、热门推荐和推荐权重动态调整功能。该路由组是推荐引擎（RecommendService）的对外HTTP接口层，是系统核心推荐能力的统一入口。",
        "设计要点：GET /home是系统的核心推荐接口——对已登录用户，编排完整的推荐流水线：(1)调用模块一的GET /profile获取用户画像（标签权重、偏好作者/标签、高分图书列表），(2)针对用户的高分图书列表中的每本图书，并行调用模块二的POST /paths获取图谱路径候选集，(3)构造UserProfileForRecommend和GraphPathsForRecommend传入RecommendService.recommend_hybrid()，执行知识图谱(KG 40%)、协同过滤(CF 40%)、热度(Hot 10%)、新书(New 10%)四种策略的加权融合，(4)为每条推荐结果调用generate_reason()生成自然语言推荐理由（如"因为你也喜欢刘慈欣的作品"）；对未登录用户，直接调用recommend_hot()返回热门推荐作为冷启动策略。",
        "GET /similar接口基于图谱路径发现与指定图书相似的候选图书，服务于图书详情页的"你可能也喜欢"模块。PUT /weights接口支持运行时调整四维推荐权重，影响后续所有推荐计算的策略倾向。",
        "推荐策略权重支持通过系统设置页动态调整，四维权重之和归一化为1.0。推荐结果带可解释的推荐理由，增强用户信任和推荐透明度。"
    ],
    "EcosystemAPI": [
        "EcosystemAPI是阅读生态模块的RESTful API路由组，路由前缀为/api/v1/ecosystem，负责提供试读权限控制、评论互动（含点赞Toggle和置顶）、多平台购书链接配置、书架管理（含默认书架和自定义书架）以及多维阅读统计功能。该路由组是系统"阅读闭环"（发现→阅读→评价→管理）的核心交互层。",
        "设计要点：试读权限采用基于登录状态的分级策略——未登录用户允许试读3页，已登录用户允许试读10页（页数限制通过settings.TRIAL_PAGES_ANONYMOUS和TRIAL_PAGES_LOGGED_IN配置，可动态调整），试读内容当前以图书简介字段作为文本源；评论系统采用置顶优先→点赞数降序→时间降序的三级排序策略，点赞采用Toggle机制（同一用户重复点赞即取消），通过comment_likes中间表记录点赞关系并实时更新book_comments.likes_count计数器；购书链接支持京东/当当/淘宝三平台配置，仅管理员可更新。",
        "书架系统默认提供"想读/在读/已读"三个不可删除的书架，用户可创建自定义书架并在书架间移动图书（move_book_to_shelf）或完全移除（remove_book_from_shelf）。删除自定义书架时同时删除其下的所有Bookmark记录（CASCADE行为）。阅读统计接口聚合总阅读量、完成/在读/想读分布、评分统计、书架分布、评论数、7日阅读趋势（每日按15分钟估算）和Top8偏好标签。",
        "涉及的数据表：bookmarks（收藏/书架）、book_comments（评论）、comment_likes（点赞）、reading_progress（阅读进度）、books（图书元数据）。评论删除操作支持作者本人和管理员两种权限路径。"
    ],
    "ChatAPI": [
        "ChatAPI是AI智能问答模块的RESTful API路由组，路由前缀为/api/v1/chat，负责提供基于大语言模型（LLM）的自然语言对话交互接口。该路由组将AIChatService的完整对话流水线（意图识别→边界检查→上下文检索→LLM生成→历史持久化）封装为HTTP服务。",
        "设计要点：POST /send是对话核心接口——接收用户自然语言消息（可附带图片URL），通过AIChatService.process_message()执行完整的处理流水线：(1)角色判定（匿名用户/普通用户/管理员），(2)LLM意图识别（7类意图：book_rec书籍推荐/book_qa书籍查询/personal_qa个人数据查询/function_qa功能问答/admin_help管理员帮助/kg_assist图谱辅助/out_of_scope超范围），LLM不可用时自动降级为关键词规则匹配，(3)边界检查——拦截天气/新闻/股票等超范围问题并返回礼貌拒绝，(4)根据意图从MySQL检索增强上下文（热门图书、标签目录、用户阅读历史、书架数据、评分记录等），(5)结合最近30条对话历史和系统提示词调用LLM生成自然语言回答，(6)保存用户消息和AI回复到chat_history表，(7)返回AI回答、意图类型和1~2条建议追问问题。",
        "对话历史管理接口支持用户查询最近对话和清空全部历史（DELETE语义，返回删除条数）。系统提示词设计了角色设定（"你是知书AI助手"），限制回答范围于图书推荐、书籍信息和阅读管理相关话题，确保对话不偏离系统业务范围。"
    ],

    # ─── Services (业务逻辑层) ───
    "UserService": [
        "UserService是用户管理模块的业务逻辑层，位于app/services/user_service.py，采用模块级函数组织方式（无类封装）。该模块封装了用户认证、画像构建、阅读行为记录和进度管理等全部核心业务逻辑，是模块一的业务处理中枢。",
        "设计要点：认证流程采用bcrypt+JWT双层安全机制——密码通过passlib的CryptContext（bcrypt方案）进行哈希存储，明文密码在内存中仅短暂存在；登录成功后通过create_access_token()签发HS256签名的JWT令牌，载荷包含sub（用户ID）和username，默认有效期24小时（通过settings.ACCESS_TOKEN_EXPIRE_MINUTES配置）。用户画像构建（build_user_profile）是模块一与模块三之间的核心契约实现——通过聚合阅读历史中的标签频率（_compute_tag_preferences，归一化为0~1权重向量）、高频作者ID（_compute_favorite_authors，统计阅读频次Top-N）、高频标签ID和高分图书ID（_get_high_rated_books，阈值默认为4.0），形成完整的UserProfileForRecommend数据结构。",
        "评分系统设计了自动触发机制：rate_book()在创建/更新评分后自动调用_update_book_avg_rating()重算图书的平均分和评分人数，确保推荐引擎获取的图书评分数据始终是最新的。进度管理采用Upsert模式（update_reading_progress），通过检查user_id+book_id的组合唯一性决定INSERT还是UPDATE，支持多端阅读进度的无缝同步。阅读统计（get_reading_stats）提供简洁的已完成/在读数摘要。",
        "涉及的数据表：users、reading_history、search_logs、bookmarks、user_ratings、reading_progress、books。所有数据库操作通过SQLAlchemy ORM Session完成，支持事务回滚。"
    ],
    "GraphService": [
        "GraphService是知识图谱模块的业务逻辑层，位于app/services/graph_service.py，采用静态方法类（@staticmethod）的组织方式。该类封装了Neo4j图数据库的全部操作，通过Cypher查询语言实现图谱路径发现、实体管理、关系管理、子图可视化和统计功能。",
        "设计要点：GraphService设计为无状态的静态方法集合，所有方法直接接收Neo4j Session对象作为参数（通过FastAPI的Depends依赖注入传入），不持有内部状态。find_paths()是系统的核心图谱算法——从源图书节点出发，沿五类预定义路径执行Cypher模式匹配查询：P1（同作者，MATCH (b1)-[:AUTHORED]-(a)-[:AUTHORED]-(b2)，权重1.0）、P2（同标签，MATCH (b1)-[:TAGGED]-(t)-[:TAGGED]-(b2)，权重0.8）、P3（同系列，MATCH (b1)-[:SERIES_OF]-(s)-[:SERIES_OF]-(b2)，权重0.6）、P4（同出版社，MATCH (b1)-[:PUBLISHED]-(p)-[:PUBLISHED]-(b2)，权重0.5）、P5（多跳：同作者→同标签，MATCH (b1)-[:AUTHORED]-(a)-[:AUTHORED]-(b_mid)-[:TAGGED]-(t)-[:TAGGED]-(b2)，权重0.7）。路径权重可通过API动态调整（path_weights参数），每条候选图书的final_score为所有命中路径的权重之和，按降序排列截取top_k。P5多跳路径仅在max_hops>=2时执行。",
        "create_relation()采用动态Cypher构造模式——根据传入的source_type/target_type动态生成MATCH子句的类型标签和MERGE关系语句，避免对每种关系类型编写硬编码查询。get_subgraph()使用Neo4j的可变长度路径遍历（variable-length pattern matching），返回去重后的节点集合（含labels和properties）和边集合（含source/target/type），直接适配前端D3.js/ECharts的力导向图渲染格式。",
        "数据源：仅操作Neo4j图数据库，不直接访问MySQL。5种节点类型分别通过init_graph_constraints()创建唯一性约束（book_id/author_id/tag_id/publisher_id/series_id），确保图谱数据的完整性。"
    ],
    "RecommendService": [
        "RecommendService是个性化推荐模块的核心业务逻辑类，位于app/services/recommend_service.py，采用实例类设计。该类封装了混合推荐引擎的完整算法逻辑，组合知识图谱推理（KG）、物品协同过滤（CF）、热度排名（Hot）和新书推荐（New）四种独立策略，按可配置权重进行加权融合，生成最终的个性化推荐列表。",
        "设计要点：RecommendService采用构造函数依赖注入模式——初始化时注入user_profile（UserProfileForRecommend，模块一的用户画像输出）和可选的graph_paths（GraphPathsForRecommend，模块二的图谱路径发现结果），四种策略方法（recommend_kg/recommend_cf/recommend_hot/recommend_new）各自独立计算候选推荐列表。recommend_hybrid()是混合推荐的核心方法——并行调用四种策略→使用内部_merge函数按策略权重将各候选图书的得分加权累加到统一的字典accumulator中→按累计得分（accumulated_score）降序排列→截取top_n→为每条结果附加推荐理由来源标记（reason字段）。",
        "协同过滤策略（recommend_cf）基于物品的协同过滤（ItemCF）算法：构建用户-图书评分稀疏矩阵（scipy.sparse.csr_matrix）→使用余弦相似度计算图书间相似度（sklearn.metrics.pairwise.cosine_similarity）→从用户高评分图书（≥4.0）出发，每本高分图书取Top-N相似图书→多本高分图书命中同一候选时按相似度×评分累加得分→排除用户已读/已评图书。相似度矩阵存储在Redis缓存（TTL=6小时），支持惰性刷新策略（新增评分>100条时触发重算）。",
        "知识图谱策略（recommend_kg）消费graph_paths中的路径发现结果，按final_score降序排列返回候选图书。热度策略（recommend_hot）按hot_score降序排列，新书策略（recommend_new）按入库时间降序排列。",
        "推荐理由生成（generate_reason）是提升推荐可解释性的关键——维护8种自然语言模板（同作者："因为你也喜欢{作者}的作品"、同标签："因为你对'{标签}'类别感兴趣"、同系列："这是'{系列}'系列的又一力作"等），根据path_info中的path_type自动选择模板并填充via信息。"
    ],
    "EcosystemService": [
        "EcosystemService是阅读生态模块的业务逻辑层，位于app/services/ecosystem_service.py，采用模块级函数组织方式。该模块封装了试读权限控制、评论互动、购书链接管理、书架系统和多维阅读统计的全部业务逻辑。",
        "设计要点：试读权限采用了分级授权模型——通过user_id是否为None判定登录状态，从settings读取对应的试读页数上限（TRIAL_PAGES_LOGGED_IN=10页/TRIAL_PAGES_ANONYMOUS=3页），试读内容使用图书的description字段截取（登录用户2000字/未登录600字），这种设计在生产环境可平滑升级为对接PDF.js等专业文档渲染服务。评论系统实现了完整的社交互动功能链：create_comment发表评论→like_comment点赞（Toggle机制，同一用户重复点赞即取消，通过comment_likes表记录关系并实时更新likes_count计数器）→get_book_comments获取列表（置顶优先→点赞降序→时间降序的三级排序）→pin_comment管理置顶→delete_comment删除（作者本人或管理员权限校验）。",
        "书架系统采用虚拟+物理混合设计：三个默认书架（"想读""在读""已读"）始终出现在get_user_bookshelves()的返回结果中（即使count为0），自定义书架通过Bookmark记录的shelf_name字段区分。书架操作支持图书移动（move_book_to_shelf，修改Bookmark的shelf_name）、图书移除（remove_book_from_shelf，删除Bookmark记录）、书架删除（delete_shelf，删除该书架下的所有Bookmark，默认书架返回-1表示拒绝删除）。每项书架操作均以user_id作为隔离边界，确保数据安全。",
        "阅读统计（get_reading_stats）是多维聚合的典型实现：统计总阅读记录数、完成/在读/想读各状态的分布、评分均值和总次数、书架分布计数、评论总数、7日阅读趋势（基于reading_history的read_at字段按日分组，每日按15分钟阅读时长估算）和Top8偏好标签（聚合阅读历史和收藏中的图书标签频次）。",
        "涉及的数据表：bookmarks（书架/收藏）、book_comments（评论）、comment_likes（点赞关系）、reading_progress（阅读进度）、reading_history（阅读历史）、user_ratings（评分）、books（图书元数据）。所有操作通过SQLAlchemy ORM的Session管理事务一致性。"
    ],
    "AIChatService": [
        "AIChatService是AI智能问答模块的业务逻辑层，位于app/services/ai_chat_service.py，采用模块级函数组织方式。该模块封装了基于大语言模型（LLM）的智能对话服务，实现"意图识别→边界检查→上下文检索→LLM生成→历史持久化"的完整对话流水线，是系统智能化交互的核心引擎。",
        "设计要点：AIChatService采用了多层降级容错架构——第一层：LLM意图识别（_classify_intent），通过OpenAI兼容API调用LLM进行7类意图分类（book_rec/book_qa/personal_qa/function_qa/admin_help/kg_assist/out_of_scope），返回结构化JSON（intent/entities/confidence）；当LLM不可用（API Key未配置、网络异常、openai包未安装）时自动降级为第二层：关键词规则匹配（_fallback_intent），基于"推荐/介绍、作者/出版社、我的/收藏、如何/怎么、添加/删除"等模式进行规则分类，置信度固定为0.5。",
        "上下文检索（_build_context）是RAG（检索增强生成）模式的核心——根据识别的意图类型，从MySQL数据库检索不同的增强数据：书籍推荐意图→返回Top10热门图书（含作者/标签/评分）+全部可用标签目录+用户画像信息；书籍查询意图→按标题/描述LIKE模糊搜索（最多5条）；个人查询意图→返回用户的阅读历史（最近10条）+书架分布+高评分图书；功能问答→返回系统功能列表说明；管理员帮助→返回后台操作指南；图谱辅助→返回标签提取模板。",
        "答案生成（_generate_answer）构建包含系统角色提示词（"你是知书AI助手，专注于图书推荐和阅读建议"）+最近20条对话历史+意图特定上下文的完整消息列表，调用LLM生成自然语言回答；LLM不可用时同样降级为模板回复（_fallback_answer）。边界检查（_check_bounds）维护一个超范围关键词黑名单（天气、新闻、股票、游戏、电影等），对命中词条返回统一的礼貌拒绝回复。",
        "对话管理支持完整的CRUD：process_message()自动保存用户消息和AI回复到chat_history表（role='user'/'assistant'），get_chat_history()按时间正序返回最近N条记录，delete_chat_history()物理删除用户全部历史（返回删除条数）。LLM客户端采用懒加载单例模式（_get_llm_client），通过settings.LLM_API_KEY和LLM_BASE_URL配置。"
    ],
    "SecurityUtils": [
        "SecurityUtils是系统级安全工具模块，位于app/core/security.py，采用模块级函数组织方式。该模块提供密码哈希、密码验证、JWT令牌生成与解码四项基础安全能力，是系统中所有认证流程的底层依赖。",
        "设计要点：密码安全采用bcrypt算法（通过passlib.context.CryptContext实现），bcrypt内建盐值（salt）机制，每次哈希结果不同但均可通过verify_password()验证；默认配置为bcrypt方案的自动选择（schemes=["bcrypt"]），标记为"deprecated"="auto"以支持未来算法升级。JWT令牌采用HS256对称签名算法，载荷（payload）包含业务数据字段（sub=用户ID, username=用户名）和exp过期时间戳（UTC时区），默认有效期通过settings.ACCESS_TOKEN_EXPIRE_MINUTES配置（24小时）；密钥（settings.SECRET_KEY）和算法（settings.ALGORITHM）从统一配置中心读取，确保全系统一致。",
        "create_access_token()支持自定义过期时间（expires_delta参数），当不传入时使用默认配置值，灵活支持不同场景（登录令牌/重置密码令牌/短期验证码）的时效需求。decode_access_token()设计为安全失败模式——任何解码异常（PyJWT的JWTError，包括签名不匹配、令牌过期、格式错误等）均返回None而非抛出异常，调用方通过判空即可判断令牌有效性，简化了API层的认证逻辑。",
        "该模块无数据库依赖，为纯函数式设计，所有函数均为同步操作（不涉及I/O），适合在API中间件和依赖注入中高频调用。"
    ],
}


def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph element."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para


def find_and_fill(doc):
    """Find each '详细描述' placeholder and insert detailed text after it."""
    body = doc.element.body
    counts = {}

    # We need to iterate over paragraphs AND tables in body order
    # Build ordered list of (index_in_body, type, element)
    children = []
    para_index_map = {}  # maps paragraph index to its body position
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para_index_map[len([c for c in children if c[1] == 'p'])] = i
            children.append((i, 'p', child))
        elif tag == 'tbl':
            children.append((i, 'tbl', child))

    filled_sections = set()

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text != "详细描述":
            continue

        # Find what this "详细描述" belongs to — look at the previous paragraphs
        # to determine section identifier
        section_name = None
        for j in range(i - 1, max(i - 6, 0), -1):
            prev_text = doc.paragraphs[j].text.strip()
            if not prev_text:
                continue
            # Match API sections
            if "UserAPI" in prev_text or "用户管理路由" in prev_text:
                section_name = "UserAPI"
                break
            elif "GraphAPI" in prev_text or "知识图谱路由" in prev_text:
                section_name = "GraphAPI"
                break
            elif "RecommendAPI" in prev_text or "推荐路由" in prev_text:
                section_name = "RecommendAPI"
                break
            elif "EcosystemAPI" in prev_text or "阅读生态路由" in prev_text:
                section_name = "EcosystemAPI"
                break
            elif "ChatAPI" in prev_text or "智能问答路由" in prev_text:
                section_name = "ChatAPI"
                break
            elif "UserService" in prev_text or "用户管理服务" in prev_text:
                section_name = "UserService"
                break
            elif "GraphService" in prev_text or "知识图谱服务" in prev_text:
                section_name = "GraphService"
                break
            elif "RecommendService" in prev_text or "推荐引擎服务" in prev_text:
                section_name = "RecommendService"
                break
            elif "EcosystemService" in prev_text or "阅读生态服务" in prev_text:
                section_name = "EcosystemService"
                break
            elif "AIChatService" in prev_text or "智能问答服务" in prev_text:
                section_name = "AIChatService"
                break
            elif "SecurityUtils" in prev_text or "安全工具" in prev_text:
                section_name = "SecurityUtils"
                break

        if section_name and section_name not in filled_sections and section_name in CONTENT:
            # Insert paragraphs after the "详细描述" paragraph in body order
            current_element = p._element
            insert_point = current_element

            # Reverse order insertion so they end up in correct order
            paragraphs_text = CONTENT[section_name]
            for para_text in reversed(paragraphs_text):
                new_p = OxmlElement("w:p")
                insert_point.addnext(new_p)
                new_para = docx.text.paragraph.Paragraph(new_p, doc)
                new_para.text = para_text
                # Set font for Chinese compatibility
                for run in new_para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)

            filled_sections.add(section_name)
            counts[section_name] = len(paragraphs_text)
            print(f"  ✅ Filled {section_name}: {len(paragraphs_text)} paragraphs added")

    return counts


def main():
    src = "document/详细设计/《基于知识图谱的个性化荐书系统》详细设计说明书_v2.1_最终版.docx"
    print(f"Loading: {src}")
    doc = Document(src)

    print("\nFilling class detailed descriptions...")
    counts = find_and_fill(doc)

    if not counts:
        print("\n⚠️  No placeholders found! Checking what happened...")
        # Debug: print contexts around "详细描述"
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text == "详细描述":
                # Print surrounding paragraphs
                start = max(0, i - 2)
                end = min(len(doc.paragraphs), i + 2)
                print(f"\n  '详细描述' at paragraph {i}:")
                for j in range(start, end):
                    print(f"    P{j}: [{doc.paragraphs[j].style.name}] {doc.paragraphs[j].text.strip()[:100]}")
        return

    print(f"\n✅ Done! Filled {len(counts)} sections:")
    for name, n in counts.items():
        print(f"  - {name}: {n} paragraphs")

    doc.save(src)
    print(f"\n📄 Saved to: {src}")


if __name__ == "__main__":
    main()