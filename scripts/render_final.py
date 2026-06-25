"""
Final render script - fixes all bugs from v1/v2:
1. Proper extraction of single-paragraph mermaid blocks
2. Fix mermaid class diagram syntax (parentheses issues)
3. Fix function ordering
4. Clean error handling
"""
import os, re, sys, subprocess, tempfile, shutil
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

BASE_DIR = r'C:\Users\Emp\PycharmProjects\PythonProject6'
MMDC = r'C:\Users\Emp\AppData\Roaming\npm\mmdc.cmd'
PUPPETEER = os.path.join(BASE_DIR, 'scripts', 'puppeteer-config.json')
DOC_DIR = os.path.join(BASE_DIR, 'document', '详细设计')
INPUT = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v2.0_struct.docx')
OUTPUT = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v2.1_最终版.docx')
WORK = tempfile.mkdtemp(prefix='mermaid_')

# ── helpers (must be defined before use) ──
def copy_para(nd, old):
    """Copy a paragraph to new document preserving formatting"""
    text = old.text
    sn = old.style.name if old.style else 'Normal'
    if not text.strip():
        nd.add_paragraph()
        return
    if sn.startswith('Heading'):
        lv = int(sn.split()[-1]) if sn.split()[-1].isdigit() else 1
        h = nd.add_heading(text, level=lv)
        for r in h.runs:
            r.font.name = '黑体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return
    p = nd.add_paragraph()
    for or_ in old.runs:
        nr = p.add_run(or_.text)
        try:
            if or_.bold: nr.bold = True
            if or_.italic: nr.italic = True
            if or_.font.size: nr.font.size = or_.font.size
            if or_.font.name:
                nr.font.name = or_.font.name
                nr._element.rPr.rFonts.set(qn('w:eastAsia'), or_.font.name)
            if or_.font.color and or_.font.color.rgb:
                nr.font.color.rgb = or_.font.color.rgb
        except: pass
    try:
        pf = old.paragraph_format
        if pf.first_line_indent: p.paragraph_format.first_line_indent = pf.first_line_indent
        if pf.line_spacing: p.paragraph_format.line_spacing = pf.line_spacing
        if pf.alignment: p.alignment = pf.alignment
    except: pass

def fix_mermaid(code):
    """Fix mermaid syntax issues that prevent rendering"""
    lines = code.split('\n')
    fixed = []
    for line in lines:
        # Fix 1: class diagrams - `{param}` in URL paths conflicts with mermaid struct syntax
        # Simply remove curly braces from URL path params
        # e.g. /subgraph/{book_id} -> /subgraph/book_id
        if re.match(r'\s*[+\-~]', line) and '{' in line:
            line = line.replace('{', '').replace('}', '')
        # Fix 2: class diagrams - remove Chinese parenthetical descriptions from method names
        # e.g. +GET /path (中文说明) -> +GET /path
        if re.match(r'\s*[+\-~]', line):
            line = re.sub(r'\s*\([^)]*[一-鿿　-〿][^)]*\)', '', line)
        # Fix 3: URL-like labels in class diagram relationships break mermaid
        # e.g. Neo4jDriver --> Neo4j : "bolt://localhost:7687" has :// confusion
        # Remove the URL label part
        if '-->' in line and '://' in line:
            line = re.sub(r'\s*:.*$', '', line)
        fixed.append(line)
    return '\n'.join(fixed)

# ═══════════════════════════
# Step 1: Extract ALL mermaid blocks from docx
# ═══════════════════════════
print("=" * 60)
print("Step 1: Extracting Mermaid diagrams from document")
print("=" * 60)

doc = Document(INPUT)
paras = doc.paragraphs

# Find all paragraphs containing mermaid code
# Each mermaid block is ONE paragraph containing ```mermaid...``` (with newlines)
diagrams = []  # [{code_idx, caption_idx, title_idx, code, title}]

for i, para in enumerate(paras):
    text = para.text
    if not text.startswith('```mermaid'):
        continue

    # This paragraph contains the full mermaid code + closing ```
    code = text.strip()
    if code.startswith('```mermaid\n'):
        code = code[11:]  # Remove ```mermaid\n
    elif code.startswith('```mermaid'):
        code = code[10:]  # Remove ```mermaid

    # Remove trailing ```
    if code.endswith('```'):
        code = code[:-3]
    code = code.strip()

    # Find title (bold paragraph before code)
    title = ""
    title_idx = None
    for j in range(i - 1, max(i - 5, -1), -1):
        prev = paras[j]
        if prev.text.strip() and not prev.text.startswith('▲') and not prev.text.startswith('```'):
            is_bold = any(r.bold for r in prev.runs if r.bold)
            if is_bold and len(prev.text) < 120:
                title = prev.text.strip()
                title_idx = j
                break

    # Find caption (next paragraph starting with ▲)
    caption_idx = None
    if i + 1 < len(paras) and paras[i + 1].text.strip().startswith('▲'):
        caption_idx = i + 1

    diagrams.append({
        'code': code,
        'code_idx': i,
        'title': title,
        'title_idx': title_idx,
        'caption_idx': caption_idx,
    })

print(f"Found {len(diagrams)} mermaid diagrams")

# ═══════════════════════════
# Step 2: Render each diagram
# ═══════════════════════════
print(f"\n{'='*60}")
print("Step 2: Rendering diagrams with mmdc")
print("=" * 60)

rendered = {}  # code_idx -> png_path

for idx, d in enumerate(diagrams):
    raw_code = d['code']
    clean_code = fix_mermaid(raw_code)
    title = d['title'][:100] if d['title'] else raw_code[:60].replace('\n', ' ')

    mmd = os.path.join(WORK, f'd{idx+1:02d}.mmd')
    png = os.path.join(WORK, f'd{idx+1:02d}.png')

    with open(mmd, 'w', encoding='utf-8') as f:
        f.write(clean_code)

    print(f"[{idx+1:02d}/{len(diagrams)}] {title}")

    try:
        r = subprocess.run(
            [MMDC, '-i', mmd, '-o', png, '-p', PUPPETEER, '-s', '2', '--backgroundColor', 'white'],
            capture_output=True, timeout=120,
            cwd=BASE_DIR,
            encoding='utf-8', errors='replace'
        )
        if os.path.exists(png) and os.path.getsize(png) > 500:
            rendered[d['code_idx']] = png
            kb = os.path.getsize(png) / 1024
            print(f"  [OK] {kb:.0f} KB")
        else:
            err = r.stderr.strip()[:200] if r.stderr else '(silent)'
            print(f"  [FAIL] {err}")
    except subprocess.TimeoutExpired:
        print(f"  [TIMEOUT]")
    except Exception as e:
        print(f"  [ERR] {e}")

print(f"\nRendered: {len(rendered)}/{len(diagrams)}")

# ═══════════════════════════
# Step 3: Build new docx with images
# ═══════════════════════════
print(f"\n{'='*60}")
print("Step 3: Building new document with embedded images")
print("=" * 60)

# Collect indices to remove (code paragraphs + titles + captions for rendered diagrams)
remove_indices = set()
for d in diagrams:
    if d['code_idx'] in rendered:
        remove_indices.add(d['code_idx'])
        if d['title_idx'] is not None:
            remove_indices.add(d['title_idx'])
        if d['caption_idx'] is not None:
            remove_indices.add(d['caption_idx'])

nd = Document()

# Page setup
for attr in ['page_width', 'page_height', 'top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
    setattr(nd.sections[0], attr, getattr(doc.sections[0], attr))

# Default font
nd.styles['Normal'].font.name = '宋体'
nd.styles['Normal'].font.size = Pt(12)
nd.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
nd.styles['Normal'].paragraph_format.line_spacing = 1.5

# Build skip set: paragraph indices to skip (rendered diagrams + titles + captions)
skip_paras = set()
render_info = {}  # code_idx -> (png_path, title)
for d in diagrams:
    if d['code_idx'] in rendered:
        skip_paras.add(d['code_idx'])
        if d['title_idx'] is not None:
            skip_paras.add(d['title_idx'])
        if d['caption_idx'] is not None:
            skip_paras.add(d['caption_idx'])
        # Store render info at the earliest index of the diagram block
        insert_at = d['code_idx']
        if d['title_idx'] is not None and d['title_idx'] < insert_at:
            insert_at = d['title_idx']
        render_info[insert_at] = (rendered[d['code_idx']], d['title'])

# Build mapping: paragraph index -> body child index
# Body children are a mix of <w:p> (paragraphs) and <w:tbl> (tables)
from lxml import etree
NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
body = doc.element.body
body_children = list(body)

# Map paragraph indices to body child positions
para_idx_to_body = {}
para_count = 0
for bi, child in enumerate(body_children):
    if child.tag == f'{NS}p':
        para_idx_to_body[para_count] = bi
        para_count += 1

print(f"  Body children: {len(body_children)} (paras: {para_count}, tables: {len(body_children) - para_count})")

# Iterate through body children in order
img_counter = 0
processed_para_indices = set()

for bi, child in enumerate(body_children):
    if child.tag == f'{NS}tbl':
        # Copy table via deep copy of XML element
        try:
            import copy
            new_tbl = copy.deepcopy(child)
            nd.element.body.append(new_tbl)
        except Exception as e:
            print(f"  [WARN] Table copy failed: {e}")
        continue

    if child.tag != f'{NS}p':
        continue

    # It's a paragraph - find its index
    para_idx = None
    for pi, pbi in para_idx_to_body.items():
        if pbi == bi:
            para_idx = pi
            break

    if para_idx is None:
        continue

    if para_idx in processed_para_indices:
        continue

    # Check if this paragraph should trigger diagram insertion
    if para_idx in render_info:
        png_path, title = render_info[para_idx]
        img_counter += 1

        # Find all diagram-related paragraph indices to mark as processed
        for d in diagrams:
            if d['code_idx'] in rendered and (para_idx == d['code_idx'] or para_idx == d['title_idx']):
                processed_para_indices.add(d['code_idx'])
                if d['title_idx'] is not None:
                    processed_para_indices.add(d['title_idx'])
                if d['caption_idx'] is not None:
                    processed_para_indices.add(d['caption_idx'])
                break

        # Insert image
        try:
            if title:
                p = nd.add_paragraph()
                r = p.add_run(title)
                r.bold = True; r.font.size = Pt(10); r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            ip = nd.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ir = ip.add_run()
            pil = Image.open(png_path); pw, ph = pil.size
            dpi = 144; max_w = 5.8
            w_in = max_w if pw/dpi > max_w else pw/dpi
            h_in = (ph/dpi)*(max_w/(pw/dpi)) if pw/dpi > max_w else ph/dpi
            ir.add_picture(png_path, width=Inches(w_in), height=Inches(h_in))

            cp = nd.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(f'图{img_counter} (已渲染)')
            cr.font.size = Pt(8); cr.font.color.rgb = RGBColor(0x99,0x99,0x99); cr.italic = True
            nd.add_paragraph()
            print(f"  [{img_counter:02d}] {os.path.basename(png_path)} ({w_in:.1f}x{h_in:.1f} in)")
        except Exception as e:
            print(f"  [{img_counter:02d}] ERROR: {e}")

        continue

    # Skip paragraphs that are part of rendered diagrams
    if para_idx in skip_paras:
        processed_para_indices.add(para_idx)
        continue

    # Normal paragraph - copy
    processed_para_indices.add(para_idx)
    if para_idx < len(paras):
        copy_para(nd, paras[para_idx])

# Save
print(f"\nSaving to: {OUTPUT}")
nd.save(OUTPUT)
print(f"Done! Embedded {img_counter} diagram images, preserved {len(doc.tables)} tables.")

# Cleanup
shutil.rmtree(WORK, ignore_errors=True)
print(f"Output: {OUTPUT}")
