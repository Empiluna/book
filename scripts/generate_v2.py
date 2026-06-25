"""
生成《基于知识图谱的个性化荐书系统》详细设计说明书 v2.0
完全对照模板结构：03-《升学空间站》详细设计文档_软工方向样例.docx
"""
import os, re, json, sys, subprocess, tempfile, shutil
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

BASE = r'C:\Users\Emp\PycharmProjects\PythonProject6'
OUTPUT_DIR = os.path.join(BASE, 'document', '详细设计')
MMDC = r'C:\Users\Emp\AppData\Roaming\npm\mmdc.cmd'
PUPPETEER_CFG = os.path.join(BASE, 'scripts', 'puppeteer-config.json')

# ═══════════════════════════════
# Helpers
# ═══════════════════════════════
def S(cell, color):
    """Set cell shading"""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color); shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

def styled_table(doc, headers, rows, col_widths=None):
    """Create styled table with dark header"""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9); r.bold = True; r.font.name = '宋体'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.color.rgb = RGBColor(255,255,255)
        S(c, '2F5496')
    for rd in rows:
        row = t.add_row()
        for i, txt in enumerate(rd):
            row.cells[i].text = str(txt)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = '宋体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return t

def H(doc, text, level):
    """Add styled heading"""
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def P(doc, text):
    """Add body paragraph with first-line indent"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def code_block(doc, text):
    """Add code block with gray background"""
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def diagram(doc, mermaid_code, title, description=""):
    """Add a diagram as code block (rendering done by separate render_final.py pipeline)"""
    if description: P(doc, description)
    diagram.counter += 1

    # Clean up mermaid code
    code = mermaid_code.strip()
    if code.startswith('```mermaid'): code = code[10:]
    if code.endswith('```'): code = code[:-3]
    code = code.strip()

    # Fix syntax issues
    fixed_lines = []
    for line in code.split('\n'):
        if re.match(r'\s*[+\-~]', line):
            line = line.replace('{','').replace('}','')
            line = re.sub(r'\s*\([^)]*[一-鿿　-〿][^)]*\)', '', line)
        if '-->' in line and '://' in line:
            line = re.sub(r'\s*:.*$', '', line)
        fixed_lines.append(line)
    code = '\n'.join(fixed_lines)

    # Always use code block (rendering done later by render_final.py)
    if title:
        p = doc.add_paragraph(); run = p.add_run(title); run.bold = True
        run.font.size = Pt(10); run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    code_block(doc, f'```mermaid\n{code}\n```')
    cp = doc.add_paragraph(); cr = cp.add_run('▲ 上图可通过Mermaid工具渲染')
    cr.font.size = Pt(8); cr.font.color.rgb = RGBColor(0x99,0x99,0x99); cr.italic = True

# counter for diagrams
diagram.counter = 1

# ═══════════════════════════════
# Generate Document
# ═══════════════════════════════
TEMP_DIR = tempfile.mkdtemp(prefix='dd_v2_')
print(f"Temp: {TEMP_DIR}")

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

style = doc.styles['Normal']; style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ═══ COVER ═══
for _ in range(6): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('基于知识图谱的个性化荐书系统'); r.font.size = Pt(28)
r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.bold = True; r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('详 细 设 计 说 明 书'); r.font.size = Pt(22)
r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph(); doc.add_paragraph()
for line in [f'版本：v2.0', f'日期：{datetime.now().strftime("%Y年%m月%d日")}',
             '项目团队：成员A · 成员B · 成员C · 成员D', '文档状态：已评审']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.size = Pt(14)

doc.add_page_break()

# ═══ PART 1: 引言 ═══
H(doc, '第一部分  引言', 1)

H(doc, '一、编写目的', 2)
P(doc, '编写本文的主要目的是把需求分析得到的用例模型和概要设计确定的技术方案，转换为具体的软件结构和数据结构。设计软件结构的具体任务是：将一个复杂系统按功能进行模块划分、建立模块的层次结构及调用关系、确定模块间的接口及人机界面等。数据结构设计包括数据特征的描述、确定数据的结构特性、以及数据库的设计。')
P(doc, '本设计是指导详细设计和项目实施的重要指导性文件，也是后续项目实施、系统集成测试的核心指导文件以及系统集成测试的重要依据。本文档面向项目开发团队（成员A/B/C/D），为各模块编码实现提供精确的技术规范。')

H(doc, '二、项目背景', 2)
P(doc, '在信息爆炸的时代，读者面临"选书难"的痛点——面对海量图书，如何快速找到适合自己的读物？传统的图书推荐多基于简单的分类浏览或畅销榜单，缺乏个性化和可解释性。')
P(doc, '本系统"基于知识图谱的个性化荐书系统"利用知识图谱技术，将图书、作者、标签、出版社、丛书等实体及其关系构建成语义网络，并结合用户画像和协同过滤算法，实现可解释的个性化图书推荐。系统支持Web端和移动端（uni-app），提供从发现图书、在线阅读到书评互动的完整阅读生态。')
P(doc, '系统采用FastAPI作为后端框架，MySQL+Neo4j双数据库架构，Redis缓存加速，并集成大语言模型提供智能问答助手。系统划分为五大模块：用户画像（模块一）、知识图谱（模块二）、个性化推荐（模块三）、阅读生态（模块四）、智能问答助手（模块五）。')

H(doc, '三、定义', 2)

H(doc, '1、技术类', 3)
tech_terms = [
    ('FastAPI', '基于Python的现代Web框架，支持异步处理、自动生成Swagger文档、Pydantic数据验证。本系统选择FastAPI而非Flask/Django，主要因为其异步支持和高性能特性。'),
    ('Neo4j', '图数据库管理系统，使用Cypher查询语言。本系统使用Neo4j存储图书、作者、标签、出版社等实体及关系图谱，支持多跳路径推理。'),
    ('MySQL 8.0', '关系型数据库管理系统。本系统使用MySQL存储用户、图书属性、评分、评论、阅读记录等结构化数据，通过SQLAlchemy ORM进行操作。'),
    ('Redis', '内存键值存储数据库。本系统使用Redis缓存热门推荐结果、用户Session、ItemCF相似度矩阵等高频访问数据。'),
    ('SQLAlchemy', 'Python ORM框架，提供对象关系映射。本系统使用SQLAlchemy管理MySQL数据库操作，支持连接池、会话管理和声明式模型定义。'),
    ('JWT (JSON Web Token)', '无状态用户认证机制。用户登录后获得Token，后续API请求在Header中携带Bearer Token进行身份验证，过期时间24小时，使用HS256算法签名。'),
    ('ItemCF', '基于物品的协同过滤算法。通过计算图书间的余弦相似度，从用户已评分的高分图书出发，推荐相似图书。相比UserCF更适合图书数量相对稳定的场景。'),
    ('uni-app', '跨平台前端框架，基于Vue 3，一套代码编译到Android/iOS/微信小程序/H5。本系统使用uni-app开发移动端，配合Pinia状态管理和axios网络请求。'),
    ('LLM (大语言模型)', '本系统模块五集成LLM（通过OpenAI兼容API），实现智能问答助手，支持意图识别、知识检索增强生成和自然语言荐书。'),
    ('ElasticSearch', '分布式全文搜索引擎。本系统使用ES为图书提供中文分词全文检索能力，支持按书名、作者、标签、简介等多字段搜索。'),
]
for name, desc in tech_terms:
    H(doc, f'（{tech_terms.index((name,desc))+1}）{name}', 4)
    P(doc, desc)

H(doc, '2、业务类', 3)
biz_terms = [
    ('用户画像', '通过分析用户的阅读历史、搜索记录、收藏偏好和评分数据，构建用户的标签偏好向量、作者偏好和类别偏好。是推荐引擎的数据底座，由模块一负责。'),
    ('知识图谱推理', '从用户感兴趣的高分图书出发，沿图谱中的AUTHORED（作者）、TAGGED（标签）、SERIES_OF（系列）、PUBLISHED（出版社）等关系路径进行多跳探索，发现候选推荐图书并生成可解释的推理路径。由模块二负责。'),
    ('混合推荐', '融合知识图谱推理推荐（40%）、协同过滤推荐（40%）、热门推荐（10%）和新书推荐（10%）四种策略，按可配置权重进行加权融合，生成最终推荐列表。由模块三负责。'),
    ('推荐理由生成', '根据推荐来源路径类型（同作者/同标签/同系列/协同过滤/热门等），自动生成自然语言推荐理由，如"因为你也喜欢刘慈欣的作品"，增强推荐的可解释性。'),
    ('阅读生态', '涵盖电子书试读（权限分级：未登录3页/登录10页）、书评社区（发表/点赞/置顶/删除）、实体书购书链接（京东/当当/淘宝多平台比价）、书架管理（自定义书架/移动/删除）和阅读统计分析。由模块四负责。'),
    ('智能问答助手', '基于大语言模型的对话式助手。支持六种意图：功能问答（如何收藏图书）、自然语言荐书（推荐科幻小说）、图书知识问答（三体的作者是谁）、个人阅读查询（我收藏了哪些书）、管理员帮助、知识图谱辅助。由模块五负责。'),
]
for i, (name, desc) in enumerate(biz_terms, 1):
    H(doc, f'（{i}）{name}', 4)
    P(doc, desc)

H(doc, '四、参考资料', 2)
refs = [
    '《基于知识图谱的个性化荐书系统-需求说明书》v1.0',
    '《基于知识图谱的个性化荐书系统》概要设计说明书 v3.2',
    '《升学空间站》详细设计文档_软工方向样例',
    '实训任务04-详细设计-1（D04）— 详细设计要求与规范',
    '实训任务05-详细设计-2（D05）— 详细设计内容与交付物',
    'FastAPI官方文档 https://fastapi.tiangolo.com/',
    'Neo4j Python Driver文档 https://neo4j.com/docs/python-manual/current/',
    'SQLAlchemy 2.0文档 https://docs.sqlalchemy.org/',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(); r = p.add_run(f'[{i}] {ref}')
    r.font.size = Pt(10.5)

doc.add_page_break()

# ═══ PART 2: 项目概述 ═══
H(doc, '第二部分  项目概述', 1)
P(doc, '《基于知识图谱的个性化荐书系统》项目旨在应对读者"选书难"的痛点，利用知识图谱技术和混合推荐算法，为读者提供可解释的个性化图书推荐。系统面向广大阅读爱好者，支持Web端和移动端访问。')
P(doc, '该平台具备多维度功能：读者可浏览个性化推荐流、搜索图书、在线试读；通过书架管理收藏和阅读进度；在书评社区中发表评论、互动点赞；还能通过AI问答助手以自然语言获取荐书建议和功能帮助。管理员可通过后台管理图书数据、用户权限、评论内容和购书链接。')
P(doc, '平台采用B/S架构与前后端分离技术，后端基于FastAPI框架，前端采用uni-app跨平台方案，数据层采用MySQL+Neo4j双库架构并配合Redis缓存和ElasticSearch搜索引擎。系统遵循分层架构原则（API层→服务层→数据层），通过JWT进行无状态认证，模块间通过明确定义的Pydantic Schema接口契约进行通信。')

doc.add_page_break()

# ═══ PART 3: 总体设计 ═══
H(doc, '第三部分  总体设计', 1)

H(doc, '一、技术架构设计', 2)

H(doc, '1、前端技术框架设计', 3)
H(doc, '1.1 前端技术栈概述', 4)
P(doc, '前端采用uni-app跨平台框架，基于Vue 3 Composition API进行组件化开发。主体框架使用uni-app 3.x（兼容Android/iOS/微信小程序/H5多端编译），视图层采用Vue 3响应式数据绑定和组合式API，状态管理使用Pinia进行全局用户信息和书架数据管理，路由使用uni-app内置路由（pages.json配置Tab栏和页面栈），网络请求使用axios封装（统一BaseURL、JWT Token自动注入、响应拦截错误处理），构建工具使用Vite 5提供快速HMR热更新。')

H(doc, '1.2 前端层次结构', 4)
P(doc, '前端采用MVVM分层架构：View层（pages/页面组件和components/复用组件）负责UI展示和用户交互，通过数据绑定与ViewModel层通信；ViewModel层（store/Pinia状态管理 + utils/工具函数）负责状态管理和业务逻辑处理；Model层（api/数据接口层）负责与后端RESTful API的数据交换。组件采用单文件组件（.vue）组织，每个页面由template（HTML结构）、script（JS逻辑）和style（CSS样式）三部分组成。')

H(doc, '2、后端逻辑架构设计', 3)
H(doc, '2.1 分层架构设计', 4)
P(doc, '后端采用三层架构设计：')
P(doc, '（1）表现层（API层，app/api/）：基于FastAPI框架，负责HTTP请求/响应处理。endpoints/目录按模块组织路由，使用Pydantic模型进行请求体验证和响应体序列化。通过Depends依赖注入管理数据库会话和用户认证。FastAPI自动生成Swagger/ReDoc交互式API文档。')
P(doc, '（2）业务逻辑层（Service层，app/services/）：封装核心业务逻辑。各模块的业务逻辑独立在各自的service文件中实现（user_service/graph_service/recommend_service/ecosystem_service/ai_chat_service），通过函数调用或类方法组织。服务层不依赖HTTP上下文，便于单元测试。')
P(doc, '（3）数据访问层（Data层，app/models/ + app/core/database.py）：通过SQLAlchemy ORM操作MySQL（声明式模型定义、Session依赖注入），通过Neo4j Python Driver操作图数据库（Cypher参数化查询），通过redis-py操作缓存（全局单例模式）。')

H(doc, '2.2 核心模块设计', 4)
P(doc, '系统采用模块化设计，划分为五大核心模块。模块间通过接口契约进行通信，各模块可独立开发、测试和部署。')
styled_table(doc,
    ['模块编号', '模块名称', '负责人', '核心职责', '上游依赖', '下游消费者'],
    [
        ['模块一', '用户画像', '成员A', '用户认证、行为采集、兴趣建模、进度同步', '无（数据底座）', '模块三、模块四、模块五'],
        ['模块二', '知识图谱', '成员B', '图谱构建、路径推理、候选发现、可视化', '无', '模块三'],
        ['模块三', '个性化推荐', '成员C', 'CF推荐、KG推荐、混合融合、理由生成', '模块一、模块二', '前端'],
        ['模块四', '阅读生态', '成员D', '试读、书评、购书、书架、统计', '模块一(JWT认证)', '前端'],
        ['模块五', '智能问答', 'ALL', '意图识别、LLM对话、历史管理', '模块一、模块三', '前端'],
    ]
)
P(doc, '系统整体技术架构如图3-1所示。')

diagram(doc, '''```mermaid
graph TB
    subgraph P[表现层]
        Web[Web前端 uni-app]
        Mobile[移动端 Android/iOS]
        H5[H5端]
    end
    subgraph G[网关层]
        Nginx[Nginx 反向代理]
    end
    subgraph B[业务逻辑层 FastAPI]
        M1[模块一 用户画像]
        M2[模块二 知识图谱]
        M3[模块三 个性化推荐]
        M4[模块四 阅读生态]
        M5[模块五 智能问答]
    end
    subgraph D[数据服务层]
        MySQL[(MySQL 8.0)]
        Neo4j[(Neo4j)]
        Redis[(Redis)]
        ES[(ElasticSearch)]
    end
    Web --> Nginx; Mobile --> Nginx; H5 --> Nginx
    Nginx --> M1; Nginx --> M2; Nginx --> M3; Nginx --> M4; Nginx --> M5
    M1 --> MySQL; M1 --> Redis
    M2 --> Neo4j; M2 --> MySQL; M2 --> ES
    M3 --> MySQL; M3 --> Redis
    M4 --> MySQL
    M5 --> MySQL
    M3 -.-> M1; M3 -.-> M2; M4 -.-> M1; M5 -.-> M1
```''', '图3-1 系统技术架构图')

H(doc, '二、核心控制流程', 2)

H(doc, '1、用户注册登录流程', 3)
P(doc, '用户注册时输入用户名、邮箱和密码，系统校验唯一性后通过bcrypt算法哈希密码并存入MySQL的users表。登录时验证密码哈希，通过后生成JWT Token（有效期24小时）返回前端。前端将Token存入本地存储，后续所有需认证的API请求在Header中携带Authorization: Bearer <token>。')
diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant F as 前端(uni-app)
    participant A as FastAPI
    participant D as MySQL
    U->>F: 输入用户名密码
    F->>A: POST /api/v1/user/login
    A->>D: 查询users表
    D-->>A: 用户记录
    A->>A: bcrypt验证密码
    alt 验证通过
        A->>A: 生成JWT Token(24h)
        A-->>F: Token + 用户信息
        F->>F: 存储Token
        F-->>U: 进入首页
    else 验证失败
        A-->>F: 401 错误
        F-->>U: 提示错误
    end
```''', '图3-2 用户登录认证时序图')

H(doc, '2、用户注册与登录活动图', 3)
P(doc, '用户注册流程包括表单验证、用户名邮箱唯一性校验、bcrypt密码哈希和JWT Token生成。登录流程通过查询users表验证凭证并返回Token。')
diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 访问登录页
    访问登录页 --> 已有账号: 是
    访问登录页 --> 注册新用户: 否
    注册新用户 --> 填写注册信息
    填写注册信息 --> 验证表单
    验证表单 --> 提交注册: 通过
    提交注册 --> bcrypt哈希密码
    bcrypt哈希密码 --> 存入MySQL
    存入MySQL --> 生成JWT
    生成JWT --> 登录成功
    已有账号 --> 输入用户名密码
    输入用户名密码 --> 提交登录
    提交登录 --> 验证密码
    验证密码 --> 登录成功: 匹配
    验证密码 --> 输入用户名密码: 错误
    登录成功 --> [*]
```''', '图3-2b 用户注册登录活动图')

H(doc, '3、个性化推荐流程', 3)
P(doc, '用户访问首页时，推荐模块首先从模块一获取用户画像（标签偏好、偏好作者、高分图书），然后对每本高分图书并行调用模块二的图谱路径查询接口，同时执行ItemCF协同过滤计算、热门推荐和新书推荐。四种策略结果按可配置权重（KG 40% + CF 40% + Hot 10% + New 10%）进行加权融合，去除重复后按得分降序排列，生成Top-N推荐列表并附带自然语言推荐理由。')
diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant R as 推荐模块(三)
    participant M1 as 用户画像(一)
    participant M2 as 知识图谱(二)
    U->>R: GET /api/v1/recommend/home
    R->>M1: 获取用户画像
    M1-->>R: tag_weights, authors, books
    loop 每本高分图书
        R->>M2: POST /graph/paths
        M2-->>R: 候选图书+推理路径
    end
    R->>R: ItemCF + Hot + New
    R->>R: 加权融合(0.4+0.4+0.1+0.1)
    R->>R: 生成推荐理由
    R-->>U: 个性化推荐列表
```''', '图3-3 个性化推荐时序图')

H(doc, '3、智能问答对话流程', 3)
P(doc, '用户输入自然语言问题后，AI问答模块首先通过LLM（或关键词降级方案）识别意图类型，然后进行边界检查（拒绝超范围问题）。对于合法意图，从MySQL检索相关业务上下文（热门图书列表、用户阅读数据、功能说明文档），获取最近30条对话历史，将System Prompt + 历史消息 + 业务上下文 + 用户问题组装后发送给LLM生成回答。最后将用户消息和AI回复保存到chat_history表。')
diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant C as 问答模块(五)
    participant L as LLM API
    participant D as MySQL
    U->>C: 自然语言问题
    C->>C: 意图识别
    alt 超范围
        C-->>U: 礼貌拒绝+引导
    else 合法意图
        C->>D: 检索业务上下文
        C->>D: 获取对话历史
        C->>L: System+上下文+问题
        L-->>C: AI回答
        C->>D: 保存对话记录
        C-->>U: 回答+建议追问
    end
```''', '图3-4 智能问答对话时序图')

H(doc, '4、图谱路径推理流程', 3)
P(doc, '图谱路径推理从源图书节点出发，并行执行五条Cypher查询路径：P1同作者路径（权重1.0）、P2同标签路径（权重0.8）、P3同系列路径（权重0.6）、P4同出版社路径（权重0.5）。当max_hops>=2时额外执行P5多跳推理路径（权重0.7）。查询结果按book_id合并去重，同一图书被多条路径命中时final_score累加，最后按得分降序返回Top-K候选。')
diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收查询请求
    接收查询请求 --> 解析参数
    解析参数 --> 并行五条Cypher查询
    state 并行五条Cypher查询 {
        P1_同作者 --> P1结果
        P2_同标签 --> P2结果
        P3_同系列 --> P3结果
        P4_同出版社 --> P4结果
        P5_多跳推理 --> P5结果
    }
    P1结果 --> 合并去重
    P2结果 --> 合并去重
    P3结果 --> 合并去重
    P4结果 --> 合并去重
    P5结果 --> 合并去重
    合并去重 --> 按final_score排序
    按final_score排序 --> 取TopK返回
    取TopK返回 --> [*]
```''', '图3-5 图谱路径推理活动图')

H(doc, '三、开发环境的集成配置', 2)

H(doc, '3.1 前端开发环境配置', 3)
P(doc, '前端基于uni-app框架开发，需要安装HBuilderX IDE或使用Vue CLI命令行工具。运行环境依赖Node.js v18+和npm。开发调试使用HBuilderX内置浏览器或微信开发者工具（小程序预览）。编译发布时通过HBuilderX云端打包或本地离线打包生成Android APK/iOS IPA/微信小程序/H5静态资源。')
code_block(doc, '''# 前端环境配置
node -v          # >= v18.0.0
npm install      # 安装依赖(vue3, pinia, axios等)
npm run dev:mp-weixin  # 微信小程序开发
npm run dev:h5   # H5开发
npm run build    # 生产构建''')

H(doc, '3.2 后端开发环境配置', 3)
P(doc, '后端基于Python 3.11+和FastAPI框架。使用Docker Compose编排MySQL、Neo4j、Redis、ElasticSearch四个基础服务。Python虚拟环境管理依赖，Uvicorn作为ASGI服务器。')
code_block(doc, '''# 后端环境配置
python -m venv .venv                   # 创建虚拟环境
.venv\\Scripts\\activate               # 激活(Windows)
pip install -r requirements.txt        # 安装依赖
docker-compose up -d                   # 启动基础设施
uvicorn app.main:app --reload --port 8000  # 启动开发服务器''')

H(doc, '3.3 Docker Compose配置', 3)
code_block(doc, '''# docker-compose.yml 核心服务
services:
  mysql:
    image: mysql:8.0
    environment: {MYSQL_ROOT_PASSWORD: root123, MYSQL_DATABASE: book_recommender}
    ports: ["3306:3306"]
  neo4j:
    image: neo4j:5.x
    environment: {NEO4J_AUTH: neo4j/password123}
    ports: ["7474:7474", "7687:7687"]
  redis:
    image: redis:7-alpine
    ports: ["6379:6379"]
  elasticsearch:
    image: elasticsearch:8.x
    ports: ["9200:9200"]''')

H(doc, '3.4 维护和监控', 3)
P(doc, '系统通过FastAPI内置的/health端点进行健康检查，返回各服务连接状态。API文档通过/docs (Swagger)和/redoc (ReDoc)自动生成和更新。日志使用Python标准logging模块，开发环境输出到控制台，生产环境接入日志收集系统。数据库备份通过Docker卷挂载实现持久化，定期执行mysqldump和neo4j-admin dump进行数据备份。')

doc.add_page_break()

# ═══ PART 4: 界面设计 ═══
H(doc, '第四部分  界面设计', 1)

H(doc, '一、界面设计', 2)

H(doc, '1、页面框架设计', 3)
P(doc, '前端采用底部Tab导航结构，包含5个Tab页面（首页、书架、AI问答、个人中心、管理后台）和1个独立页面（图书详情）。Tab栏配置在pages.json中定义，使用uni-app原生TabBar组件。全局状态通过Pinia管理，用户登录状态和书架数据跨页面共享。')
code_block(doc, '''frontend-uni/
├── App.vue                    # 根组件（全局样式、生命周期钩子）
├── main.js                    # 应用入口（创建app、注册Pinia、挂载全局组件）
├── pages.json                 # 页面路由配置（Tab栏、页面路径、窗口样式）
├── manifest.json              # 应用配置（AppID、权限声明、图标、启动图）
├── package.json               # npm依赖（vue3/pinia/axios等）
├── pages/                     # 页面组件目录
│   ├── index/index.vue        # 首页（推荐流、搜索栏、轮播）
│   ├── detail/detail.vue      # 图书详情（试读、评论、收藏、购书）
│   ├── shelf/shelf.vue        # 书架管理（书架列表、批量操作）
│   ├── chat/chat.vue          # AI问答助手（对话界面、历史回溯）
│   ├── profile/profile.vue    # 个人中心（阅读统计、趋势图、设置）
│   └── admin/admin.vue        # 管理后台（数据管理、权限控制）
├── components/                # 复用组件目录
│   ├── book-card.vue          # 图书卡片（封面/书名/作者/评分）
│   ├── comment-list.vue       # 评论列表（置顶/点赞/回复）
│   └── chat-widget.vue        # 聊天悬浮窗（可嵌入任意页面）
├── api/index.js               # API封装层（统一请求/响应拦截）
├── store/index.js             # Pinia Store（user/shelf/chat状态）
└── utils/                     # 工具函数
    ├── auth.js                # JWT Token管理（存储/读取/清除/过期检测）
    └── request.js             # axios实例（BaseURL/拦截器/错误处理）''')

H(doc, '2、师生端页面', 3)
# Each page at H4 level like template
pages_detail = [
    ('2.1 登录页面', '用户输入用户名和密码进行登录。页面包含用户名输入框（带图标）、密码输入框（支持显示/隐藏切换）、登录按钮（防重复提交）、注册跳转链接。登录成功后自动跳转首页，失败时显示错误提示。支持回车键提交。'),
    ('2.2 首页（推荐流）', '顶部为搜索栏（支持按书名/作者搜索），下方为个性化推荐流。已登录用户展示混合推荐结果（卡片式布局，显示封面、书名、作者、评分、推荐理由标签），未登录用户展示热门排行。支持下拉刷新、上拉加载更多。每条推荐卡片点击进入图书详情页。'),
    ('2.3 图书详情页', '展示图书完整信息：封面大图、书名、副标题、作者列表（可点击查看同作者图书）、出版社、ISBN、出版年份、评分（星级+数字）、标签（可点击搜索同标签图书）、简介（支持展开/收起）。功能操作区：试读按钮（未登录3页/登录10页）、收藏按钮（选择书架）、购书按钮（弹出多平台比价）。下方依次展示相似图书推荐卡片和书评评论区。'),
    ('2.4 书架页', '顶部为书架分类Tab（想读/在读/已读 + 自定义书架），支持左右滑动切换。每个书架显示图书数量徽标。书架内图书列表以网格或列表形式展示，每项显示封面缩略图、书名、作者、添加时间。支持长按弹出操作菜单（移动到其他书架/从书架移除）。右上角"+"按钮可创建自定义书架。'),
    ('2.5 AI问答助手页', '对话式界面设计：顶部标题栏显示"AI荐书助手"及清空对话按钮。中部为消息列表（用户消息右对齐蓝色气泡，AI消息左对齐白色气泡支持Markdown渲染）。底部为输入区域：文本输入框 + 发送按钮（输入为空时发送按钮隐藏）。AI回复末尾展示1-2个建议追问问题，可点击快速提问。支持键盘回车发送。'),
    ('2.6 个人中心页', '顶部用户信息卡片：头像、用户名、注册时间。阅读统计数字面板：已完成书籍数、在读书籍数、总评分次数、书架数量。最近7天阅读趋势柱状图（通过Canvas或图表库渲染）。热门标签词云展示。功能菜单列表：我的书架（跳转书架页）、修改密码、清除缓存、关于我们、退出登录。'),
]
for title, desc in pages_detail:
    H(doc, title, 4)
    P(doc, desc)

H(doc, '3、管理端页面', 3)
admin_pages = [
    ('3.1 管理后台首页', '管理员专属页面入口（需admin角色）。顶部显示管理功能导航卡片：图书管理、用户管理、评论管理、图谱管理、购书配置。每张卡片包含图标、标题和简要说明。页面仅对is_admin=true的用户可见，普通用户访问时显示无权限提示。'),
    ('3.2 图书管理页面', '图书列表以分页表格展示：ID、书名、作者、出版社、评分、热度、状态（新书/普通）、创建时间。支持按书名搜索、按标签筛选、按评分排序。操作列包含编辑（弹出编辑表单）、删除（二次确认）、配置购书链接。顶部"添加图书"按钮打开新增表单。支持JSON文件批量导入图书数据。'),
    ('3.3 用户管理页面', '用户列表展示：ID、用户名、邮箱、注册时间、状态（启用/禁用）。支持按用户名搜索。操作列包含查看详情（阅读统计、收藏列表）、禁用/启用账号（二次确认）、重置密码。不支持直接创建用户（用户只能通过注册入口创建）。'),
    ('3.4 评论管理页面', '评论列表展示：评论ID、内容摘要（截断50字）、评论者、所属图书、点赞数、置顶状态、发表时间。支持按图书筛选评论。操作列包含：置顶/取消置顶、删除评论（二次确认）。违规评论可一键删除并通知用户。'),
    ('3.5 知识图谱管理页面', '图谱概览统计卡片：节点总数、关系总数、图书节点数、作者节点数。操作按钮：初始化约束索引（幂等操作）、从MySQL同步实体到Neo4j（全量/增量）、清空图谱（危险操作需二次确认）。下方展示图谱子图可视化（通过ECharts/D3.js渲染力导向图）。'),
    ('3.6 系统设置页面', '推荐算法参数配置：KG权重（0.0~1.0滑块）、CF权重、热门权重、新书权重（四者之和自动归一化为1.0）。试读限制配置：未登录试读页数、登录后试读页数。LLM配置：API地址、模型名称、温度参数。所有配置实时生效（存入Redis缓存）。'),
]
for title, desc in admin_pages:
    H(doc, title, 4)
    P(doc, desc)

doc.add_page_break()

# ═══ PART 5: 单元模块设计 ═══
H(doc, '第五部分  单元模块设计', 1)

# ── 5.1 后端控制层设计 ──
H(doc, '一、后端控制层设计', 2)

H(doc, '1、类图设计', 3)
H(doc, '1.1 系统后端API路由层', 4)
P(doc, '系统控制层由FastAPI路由组成，按模块划分为五大路由组。通过app/api/v1/router.py统一注册，每个模块的端点独立在各自的endpoints/*.py文件中实现。所有需认证的端点通过Depends(get_current_user)依赖注入获取当前用户。')

diagram(doc, '''```mermaid
classDiagram
    class UserAPI {
        +POST /register 用户注册
        +POST /login 用户登录
        +GET /profile 获取用户画像
        +POST /history 记录阅读历史
        +GET /history 获取阅读历史
        +POST /bookmark 添加收藏
        +DELETE /bookmark/id 取消收藏
        +POST /rating 图书评分
        +POST /progress 更新阅读进度
        +GET /progress 获取阅读进度
        +GET /stats 阅读统计
    }
    class GraphAPI {
        +POST /entity 创建实体
        +POST /relation 创建关系
        +POST /init 初始化约束
        +POST /paths 图谱路径查询★
        +GET /subgraph/id 子图可视化
        +GET /stats 图谱统计
    }
    class RecommendAPI {
        +GET /home 首页推荐★
        +GET /similar/id 相似图书
        +GET /hot 热门推荐
        +PUT /weights 调整权重
    }
    class EcosystemAPI {
        +GET /trial/id 试读信息
        +GET /trial/id/content 试读内容
        +GET /comments/id 评论列表
        +POST /comments 发表评论
        +POST /comments/like 点赞评论
        +PUT /comments/id/pin 置顶评论
        +DELETE /comments/id 删除评论
        +GET /purchase/id 购书链接
        +PUT /purchase 配置购书
        +GET /shelves 书架列表
        +GET /shelves/name 书架图书
        +PUT /shelves/move 移动图书
    }
    class ChatAPI {
        +POST /message 发送消息
        +GET /history 获取历史
        +DELETE /history 清空历史
    }
    class Deps {
        +get_current_user()
        +get_current_user_optional()
        +get_admin_user()
        +get_db()
        +get_neo4j_session()
        +get_redis()
    }
    UserAPI ..> Deps
    GraphAPI ..> Deps
    RecommendAPI ..> Deps
    EcosystemAPI ..> Deps
    ChatAPI ..> Deps
```''', '图5-1 后端控制层API路由类图')

H(doc, '2、类的详细设计描述', 3)
P(doc, '2.1 系统后端API路由')

# Detailed class descriptions
class_details = [
    ('UserAPI（用户画像路由）', 'app/api/v1/endpoints/user.py', '成员A', [
        ('POST /register', 'UserRegister → TokenResponse', '注册新用户。校验用户名和邮箱唯一性，bcrypt哈希密码后存入MySQL，生成JWT Token直接返回（注册即登录）。'),
        ('POST /login', 'UserLogin → TokenResponse', '用户名+密码登录。查询users表，bcrypt验证密码哈希，通过后生成24小时有效的JWT Token。'),
        ('GET /profile', '→ UserProfile', '获取用户画像数据。【接口契约】模块三通过此接口获取标签偏好向量、偏好作者、高分图书列表用于推荐。'),
        ('POST /history', 'ReadingHistoryCreate → OK', '记录一条阅读历史。status可选read/reading/want_to_read，用于后续偏好计算。'),
        ('GET /history', '→ list[ReadingHistoryResponse]', '获取用户最近50条阅读历史，关联查询图书书名。支持分页参数limit。'),
        ('POST /bookmark', 'BookmarkCreate → OK', '将图书添加到指定书架收藏。默认书架为"默认书架"，支持自定义shelf_name。'),
        ('DELETE /bookmark/{book_id}', '→ OK', '按book_id取消收藏。'),
        ('POST /rating', 'RatingCreate → OK', '对图书评分（0.5~5.0）。自动更新图书的平均评分和评分人数。同一用户重复评分时更新原记录。'),
        ('POST /progress', 'ReadingProgressUpdate → OK', '更新阅读进度（百分比+当前页码）。使用(user_id, book_id)联合唯一索引实现upsert。'),
        ('GET /progress', '→ list[ReadingProgressResponse]', '获取用户所有在读书籍的阅读进度，按更新时间降序排列。'),
        ('GET /stats', '→ ReadingStats', '获取阅读统计数据：已完成书籍数、在读书籍数、评分次数、书架数量等。'),
    ]),
    ('GraphAPI（知识图谱路由）', 'app/api/v1/endpoints/graph.py', '成员B', [
        ('POST /entity', 'GraphEntityCreate → OK', '创建图谱实体节点（Book/Author/Tag/Publisher/Series）。通过MERGE实现幂等创建。'),
        ('POST /relation', 'GraphRelationCreate → OK', '创建两个实体间的关系（AUTHORED/TAGGED/PUBLISHED/SERIES_OF/SIMILAR）。'),
        ('POST /init', '→ OK', '初始化Neo4j唯一性约束和索引。幂等操作，重复执行不报错。'),
        ('POST /paths', 'GraphQueryRequest → GraphQueryResponse', '【核心接口·接口契约】从给定图书出发，沿五条路径进行图谱推理，返回候选图书及推理路径。模块三推荐引擎通过此接口获取KG推荐候选。'),
        ('GET /subgraph/{book_id}', '→ {nodes, edges}', '获取以某图书为中心的子图数据（节点+边），供前端ECharts/D3.js进行力导向图可视化。'),
        ('GET /stats', '→ {books, authors, tags, relations}', '图谱统计信息：各类节点数量和关系总数。'),
    ]),
    ('RecommendAPI（推荐路由）', 'app/api/v1/endpoints/recommend.py', '成员C', [
        ('GET /home', '→ RecommendListResponse', '【核心接口】首页个性化推荐。已登录用户：从模块一获取画像→从模块二获取图谱路径→四策略加权融合→生成推荐理由。未登录用户：返回热门推荐。'),
        ('GET /similar/{book_id}', '→ {similar_books}', '图书详情页"你可能也喜欢"——纯图谱推荐，返回与该书通过图谱关系关联的相似图书。'),
        ('GET /hot', '→ {hot_books}', '全站热门图书排行，按hot_score降序。用于冷启动和未登录用户的推荐。'),
        ('PUT /weights', 'RecommendWeights → OK', '运行时调整混合推荐的四种策略权重（KG/CF/Hot/New），管理员功能。'),
    ]),
    ('EcosystemAPI（阅读生态路由）', 'app/api/v1/endpoints/ecosystem.py', '成员D', [
        ('GET /trial/{book_id}', '→ TrialReadResponse', '获取试读权限信息。根据登录状态返回不同的试读页数限制（未登录3页/登录10页），同时返回当前阅读进度。'),
        ('GET /trial/{book_id}/content', '→ TrialContentResponse', '获取试读具体内容。当前使用图书简介作为试读文本，内容长度按登录状态限制。'),
        ('POST /comments', 'CommentCreate → OK', '发表图书评论。用户需登录，内容长度限制5000字符。'),
        ('GET /comments/{book_id}', '→ list[CommentResponse]', '获取图书评论列表，按置顶优先→点赞数降序→时间降序排列，支持分页。'),
        ('POST /comments/like', 'CommentLikeAction → OK', '点赞/取消点赞Toggle模式。已点赞则取消，未点赞则添加，同步更新评论的likes_count。'),
        ('PUT /comments/{id}/pin', '→ OK', '管理员置顶/取消置顶评论。'),
        ('DELETE /comments/{id}', '→ OK', '删除评论。作者本人或管理员可删除。'),
        ('GET /purchase/{book_id}', '→ PurchaseLinkResponse', '获取图书的多平台购书链接（京东/当当/淘宝）。'),
        ('PUT /purchase', 'PurchaseLinkUpdate → OK', '管理员配置图书的购书链接。'),
        ('GET /shelves', '→ list[ShelfResponse]', '获取用户所有书架列表（含默认三个书架和自定义书架），显示每个书架的图书数量。'),
        ('GET /shelves/{name}', '→ list[ShelfBookResponse]', '获取某个书架内的图书列表，含图书详情（封面、作者、评分等）。'),
        ('PUT /shelves/move', 'MoveBookRequest → OK', '将图书从一个书架移动到另一个书架。'),
    ]),
    ('ChatAPI（智能问答路由）', 'app/api/v1/endpoints/ai_chat.py', 'ALL', [
        ('POST /message', 'ChatRequest → ChatResponse', '【核心接口】处理用户自然语言消息。流程：意图识别→边界检查→上下文检索→LLM生成→保存历史。支持未登录用户（仅限功能问答和图书推荐）。'),
        ('GET /history', '→ ChatHistoryResponse', '获取当前用户的对话历史（最近50条消息）。'),
        ('DELETE /history', '→ ChatDeleteResponse', '清空当前用户的全部对话历史。'),
    ]),
]

for class_name, file_path, owner, methods in class_details:
    H(doc, f'{class_name}', 4)
    P(doc, f'文件位置：{file_path}　　负责人：{owner}')
    P(doc, '详细描述：')
    # Method table
    rows = []
    for method, signature, desc in methods:
        rows.append([method, signature, desc])
    styled_table(doc, ['方法/端点', '签名/数据模型', '功能说明'], rows)

# ── 5.2 业务逻辑层设计 ──
H(doc, '二、业务逻辑层设计', 2)

H(doc, '1、类图设计', 3)
H(doc, '1.1 系统后端Service层', 4)
P(doc, '业务逻辑层由五个核心Service类和一个辅助工具类组成。各Service独立封装业务逻辑，不依赖HTTP上下文，通过函数参数接收数据库会话。Service之间通过接口契约（Pydantic Schema）进行数据交换。')

diagram(doc, '''```mermaid
classDiagram
    class UserService {
        +register_user(db, username, email, password) User
        +authenticate_user(db, username, password) str
        +record_reading_history(db, user_id, book_id, status)
        +add_bookmark(db, user_id, book_id, shelf_name)
        +rate_book(db, user_id, book_id, rating)
        +build_user_profile(db, user_id) dict
        +update_reading_progress(db, user_id, book_id, percent, page)
        +get_reading_stats(db, user_id) dict
        -_compute_tag_preferences(db, user_id) dict
        -_compute_favorite_authors(db, user_id) list
        -_get_high_rated_books(db, user_id) list
    }
    class GraphService {
        +init_graph_constraints(session)
        +create_book_entity(session, book_id, title, props)
        +find_paths(session, book_id, max_hops, top_k, weights) dict
        +get_subgraph(session, book_id, depth) dict
        +get_stats(session) dict
    }
    class RecommendService {
        -user_profile: UserProfileForRecommend
        -graph_paths: GraphPathsForRecommend
        +recommend_cf(top_n) list
        +recommend_kg(top_n) list
        +recommend_hot(top_n) list
        +recommend_new(top_n) list
        +recommend_hybrid(top_n, weights) list
        +generate_reason(book_title, path_info) str$
    }
    class EcosystemService {
        +get_trial_info(db, book_id, user_id) dict
        +create_comment(db, user_id, book_id, content) BookComment
        +like_comment(db, user_id, comment_id) bool
        +pin_comment(db, comment_id, is_pinned)
        +get_user_bookshelves(db, user_id) list
        +move_book_to_shelf(db, user_id, book_id, shelf)
        +get_reading_stats(db, user_id) dict
    }
    class AIChatService {
        +process_message(db, user, message) dict
        +save_message(db, user_id, role, content, intent)
        +get_chat_history(db, user_id, limit) list
        -_classify_intent(message, user_role) dict
        -_build_context(intent, user, db, entities) str
        -_generate_answer(msg, intent, ctx, user, db, history) dict
    }
    class SecurityUtils {
        +hash_password(password) str$
        +verify_password(plain, hashed) bool$
        +create_access_token(data, expires) str$
        +decode_access_token(token) dict$
    }
    UserService ..> SecurityUtils
    RecommendService ..> UserService : 通过接口契约
    RecommendService ..> GraphService : 通过接口契约
```''', '图5-2 业务逻辑层Service类图')

# Per-module detailed class diagrams
H(doc, '3、各模块核心类图设计', 3)

modules_diagrams = [
    ('模块一：用户画像', '图5-3 模块一类图', '''```mermaid
classDiagram
    class User {
        +int id; +str username; +str email
        +str hashed_password; +bool is_active
    }
    class ReadingHistory {
        +int id; +int user_id; +int book_id
        +str status; +datetime read_at
    }
    class Bookmark {
        +int id; +int user_id; +int book_id
        +str shelf_name
    }
    class ReadingProgress {
        +int id; +int user_id; +int book_id
        +float progress_percent; +int current_page
    }
    class UserRating {
        +int id; +int user_id; +int book_id
        +float rating
    }
    class UserService {
        +register_user() User
        +authenticate_user() str
        +record_reading_history()
        +build_user_profile() dict
        +update_reading_progress()
        +get_reading_stats() dict
    }
    User "1" --> "*" ReadingHistory
    User "1" --> "*" Bookmark
    User "1" --> "*" ReadingProgress
    User "1" --> "*" UserRating
    UserService ..> User
    UserService ..> ReadingHistory
    UserService ..> Bookmark
    UserService ..> UserRating
```'''),
    ('模块二：知识图谱', '图5-4 模块二类图', '''```mermaid
classDiagram
    class GraphService {
        +init_graph_constraints(session)
        +create_book_entity(session, book_id, title, props)
        +create_relation(session, src, rel, tgt)
        +find_paths(session, book_id, max_hops, top_k, weights)
        +get_subgraph(session, book_id, depth)
        +get_stats(session)
    }
    class Book { +int id; +str title; +str isbn }
    class Author { +int id; +str name; +str bio }
    class Tag { +int id; +str name; +str category }
    class Publisher { +int id; +str name }
    class Series { +int id; +str name }
    class Neo4jDriver { +session(); +run(cypher) }
    GraphService --> Neo4jDriver
    GraphService ..> Book
    GraphService ..> Author
    GraphService ..> Tag
    Book "n" -- "m" Author
    Book "n" -- "m" Tag
    Book "n" -- "1" Publisher
    Book "n" -- "1" Series
```'''),
    ('模块三：个性化推荐', '图5-5 模块三类图', '''```mermaid
classDiagram
    class RecommendService {
        +recommend_cf(top_n) list
        +recommend_kg(top_n) list
        +recommend_hot(top_n) list
        +recommend_new(top_n) list
        +recommend_hybrid(top_n, weights) list
        +generate_reason(title, path) str$
    }
    class UserProfile { +tag_weights; +author_ids; +book_ids }
    class GraphPaths { +source_book_id; +candidates }
    class RecommendWeights { +kg=0.4; +cf=0.4; +hot=0.1; +new=0.1 }
    RecommendService --> UserProfile : 模块一提供
    RecommendService --> GraphPaths : 模块二提供
    RecommendService --> RecommendWeights
```'''),
    ('模块四：阅读生态', '图5-6 模块四类图', '''```mermaid
classDiagram
    class EcosystemService {
        +get_trial_info(db, book_id, user_id) dict
        +get_trial_content(db, book_id, user_id) dict
        +create_comment(db, user_id, book_id, content) BookComment
        +get_book_comments(db, book_id, page, size) list
        +like_comment(db, user_id, comment_id) bool
        +pin_comment(db, comment_id, is_pinned)
        +delete_comment(db, comment_id, user_id, is_admin)
        +get_user_bookshelves(db, user_id) list
        +move_book_to_shelf(db, user_id, book_id, shelf)
        +get_purchase_links(db, book_id) dict
        +get_reading_stats(db, user_id) dict
    }
    class BookComment { +int id; +int user_id; +int book_id; +text content; +int likes_count; +bool is_pinned }
    class CommentLike { +int id; +int user_id; +int comment_id }
    EcosystemService --> BookComment
    EcosystemService --> CommentLike
```'''),
    ('模块五：智能问答助手', '图5-7 模块五类图', '''```mermaid
classDiagram
    class AIChatService {
        +process_message(db, user, message) dict
        +save_message(db, user_id, role, content, intent)
        +get_chat_history(db, user_id, limit) list
        +delete_chat_history(db, user_id) int
        -_classify_intent(msg, role) dict
        -_build_context(intent, user, db, entities) str
        -_check_bounds(msg, intent) str
        -_generate_answer(msg, intent, ctx, user, db) dict
        -_fallback_intent(msg) dict
        -_fallback_answer(intent, ctx) dict
    }
    class ChatHistory { +int id; +int user_id; +enum role; +text content; +str intent_type }
    AIChatService --> ChatHistory
    AIChatService ..> LLMClient : OpenAI兼容API
```'''),
]

for name, title, code in modules_diagrams:
    P(doc, f'{name}模块核心类的职责和关系如下所示。')
    diagram(doc, code, title)

H(doc, '2、类的详细设计描述', 3)
P(doc, '1.1 系统后端Service层')

service_details = [
    ('UserService（用户画像服务）', 'app/services/user_service.py', [
        ('register_user(db, username, email, password) → User', '注册新用户。调用hash_password()对密码进行bcrypt哈希，创建User ORM对象并commit到MySQL的users表，返回新创建的User实例。'),
        ('authenticate_user(db, username, password) → str', '用户认证。查询users表按username查找，调用verify_password()验证bcrypt哈希，通过后调用create_access_token()生成JWT Token（有效期24小时）。'),
        ('build_user_profile(db, user_id) → dict', '【核心方法·接口契约】构建用户画像。从reading_history和bookmarks统计标签频次并归一化为tag_weights字典（如{"科幻":0.85,"历史":0.15}），按作者频次降序取top20为favorite_author_ids，从user_ratings筛选评分>=4.0的图书为high_rated_book_ids。输出格式匹配UserProfileForRecommend Schema。'),
        ('record_reading_history(db, user_id, book_id, status)', '记录阅读行为。status可选read/reading/want_to_read。用于后续偏好计算的数据积累。'),
        ('rate_book(db, user_id, book_id, rating)', '评分。使用(user_id, book_id)唯一索引实现upsert——已存在则更新评分，不存在则新建。评分后自动调用_update_book_avg_rating()重新计算图书平均分。'),
        ('update_reading_progress(db, user_id, book_id, percent, page)', '多端阅读进度同步。使用(user_id, book_id)联合唯一索引实现upsert，已存在记录则更新progress_percent和current_page，不存在则插入新记录。'),
        ('_compute_tag_preferences(db, user_id) → dict', '内部方法。从阅读历史关联图书标签，统计每个标签的出现频次，归一化为0~1的权重字典。'),
        ('get_reading_stats(db, user_id) → dict', '获取阅读统计数据。统计已完成/在读书籍数、评分次数和均价、书架数量、评论数量，以及最近7天每日阅读趋势数据。'),
    ]),
    ('GraphService（知识图谱服务）', 'app/services/graph_service.py', [
        ('init_graph_constraints(session)', '初始化Neo4j唯一性约束。为Book.book_id、Author.author_id、Tag.tag_id、Publisher.publisher_id、Series.series_id创建UNIQUE约束，确保实体节点不重复。'),
        ('create_book_entity(session, book_id, title, **props)', '创建或更新图书节点。使用MERGE语句实现幂等操作——节点已存在则更新属性，不存在则创建。'),
        ('find_paths(session, book_id, max_hops, top_k, path_weights) → dict', '【核心方法·接口契约】多跳路径推理。定义五条参数化Cypher查询（P1同作者weight=1.0、P2同标签weight=0.8、P3同系列weight=0.6、P4同出版社weight=0.5、P5多跳推理weight=0.7仅max_hops>=2时执行）。五条路径的查询结果汇聚到candidates_map按book_id合并，同一图书被多条路径命中时final_score累加，最后按得分降序返回top_k候选。所有Cypher查询使用$参数化方式防止注入。'),
        ('get_subgraph(session, book_id, depth) → dict', '获取以某图书为中心的子图数据。使用可变长度路径查询MATCH path=(b)-[*1..depth]-()，解析路径中的节点和关系，返回{nodes, edges}格式供前端ECharts/D3.js渲染力导向图。'),
        ('get_stats(session) → dict', '统计图谱中的各类节点数量和关系总数：{books, authors, tags, relations}。'),
    ]),
    ('RecommendService（推荐引擎服务）', 'app/services/recommend_service.py', [
        ('__init__(user_profile, graph_paths)', '构造函数。接收UserProfileForRecommend（模块一提供的用户画像）和可选的GraphPathsForRecommend（模块二提供的图谱查询结果）。'),
        ('recommend_cf(top_n, exclude_ids) → list', 'ItemCF协同过滤推荐。构建用户-图书评分稀疏矩阵（scipy.sparse.csr_matrix），使用sklearn.metrics.pairwise.cosine_similarity计算图书间余弦相似度。从用户高分图书出发，推荐Top-N最相似的未知图书。相似度矩阵建议离线预计算并缓存到Redis。'),
        ('recommend_kg(top_n, exclude_ids) → list', '知识图谱推理推荐。遍历用户的高分图书列表，为每本书调用模块二的find_paths()获取图谱推荐候选，合并多源路径结果，按final_score降序排列去重。'),
        ('recommend_hot(top_n, exclude_ids) → list', '热门推荐。按图书的hot_score降序排列，用于新用户冷启动和未登录用户的推荐。'),
        ('recommend_new(top_n, exclude_ids) → list', '新书推荐。筛选is_new=True的图书，按入库时间降序排列，提供时效性推荐。'),
        ('recommend_hybrid(top_n, weights, exclude_ids) → list', '【核心方法】混合推荐。四种策略独立计算候选列表，按RecommendWeights配置（默认kg=0.4, cf=0.4, hot=0.1, new=0.1）进行加权融合：final_score = Σ(策略得分 × 策略权重)。同一本书被多个策略命中时分数累加，最后按得分降序返回top_n。'),
        ('generate_reason(book_title, path_info) → str', '推荐理由生成。根据path_type映射自然语言模板：author→"因为你也喜欢{作者}的作品"、tag→"这本书和{标签}相关"、series→"与{系列}属于同一丛书"、cf→"与你读过的高分图书相似"、hot→"近期热门图书"、new→"新书上架"。'),
    ]),
    ('EcosystemService（阅读生态服务）', 'app/services/ecosystem_service.py', [
        ('get_trial_info(db, book_id, user_id) → dict', '获取试读权限。根据user_id是否为None判断登录状态，返回对应的试读页数限制（TRIAL_PAGES_ANONYMOUS=3/TRIAL_PAGES_LOGGED_IN=10），同时查询reading_progress表获取当前阅读进度。'),
        ('get_trial_content(db, book_id, user_id) → dict', '获取试读具体内容。当前使用图书description字段作为试读文本，内容长度按登录状态限制（登录2000字符/未登录600字符）。生产环境可对接PDF.js/EPUB解析器。'),
        ('create_comment(db, user_id, book_id, content) → BookComment', '发表书评。创建BookComment ORM对象并commit，内容长度限制5000字符。'),
        ('get_book_comments(db, book_id, page, page_size) → list', '获取评论列表。排序规则：is_pinned DESC → likes_count DESC → created_at DESC。支持分页。'),
        ('like_comment(db, user_id, comment_id) → bool', 'Toggle点赞。查询comment_likes表检查是否已点赞——已点赞则DELETE（返回False），未点赞则INSERT（返回True）。操作后调用_update_comment_likes()同步更新book_comments表的likes_count聚合字段。'),
        ('pin_comment(db, comment_id, is_pinned)', '管理员置顶。更新book_comments表的is_pinned字段。'),
        ('delete_comment(db, comment_id, user_id, is_admin)', '删除评论。权限控制：管理员可删除任意评论，普通用户仅可删除自己的评论。'),
        ('get_user_bookshelves(db, user_id) → list', '获取书架列表。从bookmarks表按shelf_name分组统计数量，确保三个默认书架（想读/在读/已读）始终在列表中，自定义书架追加其后。'),
        ('move_book_to_shelf(db, user_id, book_id, new_shelf)', '移动图书到另一个书架。修改bookmark记录的shelf_name字段。'),
        ('get_reading_stats(db, user_id) → dict', '综合阅读统计。统计阅读历史数量、完成/在读/想读书籍数、评分次数与均分、书架数、评论数、在读书籍数，以及最近7天每日阅读时长趋势（每次阅读记录估算15分钟）。同时统计用户TOP8热门标签。'),
    ]),
    ('AIChatService（智能问答服务）', 'app/services/ai_chat_service.py', [
        ('process_message(db, user, message) → dict', '【核心方法】对话处理主入口。处理流程：确定用户角色(anonymous/user/admin)→意图识别(_classify_intent)→边界检查(_check_bounds)→检索业务上下文(_build_context)→获取对话历史(最近30条)→LLM生成回答(_generate_answer)→保存对话记录(save_message×2)。'),
        ('_classify_intent(message, user_role) → dict', '意图识别。LLM可用时通过INTENT_PROMPT让LLM输出JSON分类结果{"intent","entities","confidence"}（temperature=0.1确保稳定）。LLM不可用时降级为Python关键词匹配（如含"推荐/介绍"→book_rec，含"我/我的"→personal_qa）。六种意图：function_qa/book_rec/book_qa/personal_qa/admin_help/kg_assist，以及out_of_scope超范围拒绝。'),
        ('_build_context(intent, user, db, entities) → str', '根据意图从MySQL检索相关业务数据：function_qa→系统功能说明文档、book_rec→热门图书TOP10+可用标签列表+用户偏好、book_qa→按关键词搜索图书详情、personal_qa→用户阅读历史/书架/评分数据、admin_help→管理员操作指引文档、kg_assist→图书简介文本。'),
        ('_generate_answer(msg, intent, ctx, user, db, history) → dict', 'LLM生成回答。组装消息列表：[System Prompt] + 历史对话(最近20条message) + [业务上下文+用户问题]，调用OpenAI兼容API生成回答。LLM不可用时使用_fallback_answer()降级方案（预置模板回复）。')]),
    ('SecurityUtils（安全工具）', 'app/core/security.py', [
        ('hash_password(password) → str', '使用bcrypt算法对密码进行哈希。通过passlib.context.CryptContext实现，schemes=["bcrypt"]。'),
        ('verify_password(plain, hashed) → bool', '验证明文密码与bcrypt哈希是否匹配。'),
        ('create_access_token(data, expires_delta) → str', '生成JWT Token。默认过期时间24小时，使用HS256算法签名，payload包含sub(user_id)和username。'),
        ('decode_access_token(token) → dict', '解码JWT Token。验证签名和过期时间，失败返回None。'),
    ]),
]

for svc_name, svc_path, methods in service_details:
    H(doc, f'{svc_name}', 4)
    P(doc, f'文件位置：{svc_path}')
    P(doc, '详细描述：')
    rows = []
    for method, desc in methods:
        rows.append([method, desc])
    styled_table(doc, ['方法签名', '功能说明'], rows)

# ── 5.3 推荐算法设计 ──
H(doc, '三、推荐算法设计', 2)

H(doc, '1、ItemCF协同过滤算法', 3)
P(doc, 'ItemCF（Item-based Collaborative Filtering）基于物品的协同过滤算法。核心思想：如果两本图书被相似的用户群体评分，则认为它们相似。')
P(doc, '算法步骤：')
P(doc, '步骤1：从MySQL user_ratings表构建用户-图书评分矩阵（rows=users, cols=books, values=ratings），使用scipy.sparse.csr_matrix存储为稀疏矩阵。')
P(doc, '步骤2：计算图书间的余弦相似度。cos(A,B) = (A·B)/(|A|·|B|)，使用sklearn.metrics.pairwise.cosine_similarity。过滤评分人数<10的冷门图书以提升计算效率和推荐质量。')
P(doc, '步骤3：从用户画像的high_rated_book_ids（评分≥4.0的图书）出发，对每本高分图书取Top-N最相似图书。多本高分书命中同一本书时，累加（相似度×评分）作为最终推荐分数。排除用户已读/已评分的图书。')
P(doc, '步骤4：优化——相似度矩阵离线预计算，结果存入Redis缓存（TTL=6小时，定时刷新），避免每次请求重新计算。增量更新策略：新增评分>100条时触发矩阵重计算。')

H(doc, '2、知识图谱路径推理算法', 3)
P(doc, '知识图谱推荐基于图遍历的候选发现策略。从用户感兴趣的高分图书（种子节点）出发，沿图谱中的多种关系路径进行多跳探索。')
styled_table(doc,
    ['路径编号', '路径名称', 'Cypher查询模式', '跳数', '权重', '推荐语义'],
    [
        ['P1', '同作者', '(Book)-[:AUTHORED]-(Author)-[:AUTHORED]-(Book)', '1', '1.0', '同作者的其他作品，最强推荐信号'],
        ['P2', '同标签', '(Book)-[:TAGGED]-(Tag)-[:TAGGED]-(Book)', '1', '0.8', '共享标签越多越相关，按共享标签数排序'],
        ['P3', '同系列', '(Book)-[:SERIES_OF]-(Series)-[:SERIES_OF]-(Book)', '1', '0.6', '同一丛书/系列的续作或前作'],
        ['P4', '同出版社', '(Book)-[:PUBLISHED]-(Publisher)-[:PUBLISHED]-(Book)', '1', '0.5', '同出版社的同类图书，信号较弱'],
        ['P5', '多跳推理', '(Book)-Author-(Book)-Tag-(Book)', '≥2', '0.7', '综合推理：同作者其他书的标签关联图书'],
    ]
)
P(doc, '权重计算：每个候选图书的final_score = Σ(命中路径的权重)。例如《流浪地球》同时被P1（同作者刘慈欣，1.0）和P2（同标签科幻，0.8）命中，则final_score=1.0+0.8=1.8。路径权重可通过API参数的path_weights字典运行时覆盖。')

H(doc, '3、混合推荐融合算法', 3)
P(doc, '混合推荐将四种策略的推荐结果进行加权融合，兼顾精度（CF）、可解释性（KG）、覆盖率（Hot）和时效性（New）。')
code_block(doc, '''融合公式:
  final_score = kg_score × W_kg + cf_score × W_cf + hot_score × W_hot + new_score × W_new

默认权重（总权重=1.0）:
  W_kg  = 0.4  （知识图谱推理 — 可解释性强，推荐理由可溯源）
  W_cf  = 0.4  （协同过滤 — 利用群体智慧，精度高）
  W_hot = 0.1  （热门推荐 — 冷启动保底，确保新用户有推荐）
  W_new = 0.1  （新书推荐 — 时效性，促进新书发现）''')
P(doc, '融合流程：四种策略独立计算候选列表→创建merged字典(key=book_id)→遍历四种策略结果，对每本候选图书将score×权重累加到accumulated_score→按accumulated_score降序排列截取top_n→为每条推荐调用generate_reason()生成自然语言推荐理由。')

diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收推荐请求
    接收推荐请求 --> 检查登录状态
    检查登录状态 --> 热门推荐: 未登录
    热门推荐 --> 返回推荐列表
    检查登录状态 --> 获取用户画像: 已登录
    获取用户画像 --> 并行四种策略
    state 并行四种策略 {
        KG推荐 --> KG结果
        CF推荐 --> CF结果
        Hot推荐 --> Hot结果
        New推荐 --> New结果
    }
    KG结果 --> 加权融合
    CF结果 --> 加权融合
    Hot结果 --> 加权融合
    New结果 --> 加权融合
    加权融合 --> 排序取TopN
    排序取TopN --> 生成推荐理由
    生成推荐理由 --> 返回推荐列表
    返回推荐列表 --> [*]
```''', '图5-8 混合推荐算法活动图')

H(doc, '4、意图识别算法（模块五）', 3)
P(doc, '意图识别采用LLM优先+关键词降级的双层策略，确保在各种环境下都能提供可用的意图分类结果。')
P(doc, '第一层（LLM识别）：当LLM API可用时，将用户消息和精心设计的INTENT_PROMPT发送给LLM，要求以JSON格式返回{"intent":"xxx","entities":[...],"confidence":0.0-1.0}。使用temperature=0.1确保分类结果稳定一致。')
P(doc, '第二层（关键词降级）：当LLM不可用时，使用Python关键词匹配规则进行分类：含"推荐/介绍/好看/适合/入门"→book_rec，含"作者/出版社/标签/ISBN"→book_qa，含"我/我的/收藏/书架/历史/记录"→personal_qa，含"怎么/如何/在哪/功能/使用"→function_qa，含"添加/删除/管理/后台"→admin_help。降级方案置信度为0.5。')

diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收用户消息
    接收用户消息 --> 意图识别
    意图识别 --> 超范围拒绝: out_of_scope
    超范围拒绝 --> [*]
    意图识别 --> 合法意图: 6种业务意图
    合法意图 --> 关键词边界检查
    关键词边界检查 --> 超范围拒绝: 含禁用词
    关键词边界检查 --> 检索业务上下文: 通过
    检索业务上下文 --> 获取对话历史
    获取对话历史 --> LLM生成回答
    LLM生成回答 --> 保存对话记录
    保存对话记录 --> [*]
```''', '图5-9 智能问答对话处理活动图')

doc.add_page_break()

# ═══ PART 6: 数据库设计 ═══
H(doc, '第六部分  数据库设计', 1)

H(doc, '一、数据库整体结构图', 2)
P(doc, '系统采用MySQL+Neo4j双数据库架构。MySQL存储结构化实体属性和业务数据（共15张表），Neo4j存储实体间的语义关系图谱（5种节点类型+5种关系类型）。两个数据库通过统一的实体ID进行关联。')

diagram(doc, '''```mermaid
erDiagram
    User ||--o{ ReadingHistory : 阅读
    User ||--o{ Bookmark : 收藏
    User ||--o{ ReadingProgress : 进度
    User ||--o{ UserRating : 评分
    User ||--o{ BookComment : 评论
    User ||--o{ CommentLike : 点赞
    User ||--o{ ChatHistory : 对话
    Book ||--o{ ReadingHistory : 被读
    Book ||--o{ Bookmark : 被收藏
    Book ||--o{ ReadingProgress : 进度
    Book ||--o{ UserRating : 被评分
    Book ||--o{ BookComment : 评论
    Book }o--|| Publisher : 属于
    Book }o--|| Series : 系列
    Book }o--o{ Author : 多对多
    Book }o--o{ Tag : 多对多
    BookComment ||--o{ CommentLike : 被点赞
```''', '图6-1 MySQL数据库ER图')

P(doc, 'Neo4j图数据库存储实体间的语义关系图谱，包含5种实体节点（Book/Author/Tag/Publisher/Series）和5种关系类型（AUTHORED/TAGGED/PUBLISHED/SERIES_OF/SIMILAR）。图谱与MySQL通过统一实体ID关联。各关系类型在推荐推理中具有不同权重。')
diagram(doc, '''```mermaid
graph LR
    subgraph 实体节点
        B["Book<br/>book_id, title, isbn"]
        A["Author<br/>author_id, name"]
        T["Tag<br/>tag_id, name, category"]
        P["Publisher<br/>publisher_id, name"]
        S["Series<br/>series_id, name"]
    end
    B ---|"AUTHORED (1.0)"| A
    B ---|"TAGGED (0.8)"| T
    B ---|"PUBLISHED (0.5)"| P
    B ---|"SERIES_OF (0.6)"| S
    B ---|"SIMILAR (0.7)"| B
```''', '图6-2 Neo4j知识图谱Schema')

H(doc, '二、详细表结构', 2)

# Complete table definitions at H4 level like template
tables_def = [
    ('1、users表结构', '用户信息表，存储系统用户的基本信息和认证数据。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '用户唯一标识'),
        ('username', 'VARCHAR(64)', 'UNIQUE, NOT NULL, INDEX', '用户名，用于登录和显示'),
        ('email', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '邮箱地址'),
        ('hashed_password', 'VARCHAR(256)', 'NOT NULL', 'bcrypt哈希后的密码'),
        ('is_active', 'BOOLEAN', 'DEFAULT TRUE', '账户启用状态，禁用后不可登录'),
        ('is_admin', 'BOOLEAN', 'DEFAULT FALSE', '管理员权限标识'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '注册时间'),
        ('updated_at', 'DATETIME', 'ON UPDATE NOW()', '信息更新时间'),
    ]),
    ('2、books表结构', '图书主表，存储图书的核心元数据。系统中数据量最大的表，同时存储模块四的购书链接。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '图书唯一标识'),
        ('title', 'VARCHAR(256)', 'NOT NULL, INDEX', '书名'),
        ('subtitle', 'VARCHAR(256)', '', '副标题'),
        ('isbn', 'VARCHAR(20)', 'UNIQUE', '国际标准书号ISBN'),
        ('publisher_id', 'INTEGER', 'FK→publishers.id', '出版社ID'),
        ('series_id', 'INTEGER', 'FK→series.id', '丛书系列ID'),
        ('publication_year', 'INTEGER', '', '出版年份'),
        ('description', 'TEXT', '', '图书简介/内容摘要'),
        ('cover_url', 'VARCHAR(512)', '', '封面图片URL'),
        ('page_count', 'INTEGER', '', '总页数'),
        ('language', 'VARCHAR(32)', "DEFAULT 'zh-CN'", '语言代码'),
        ('avg_rating', 'FLOAT', 'DEFAULT 0.0', '用户平均评分（由评分触发器更新）'),
        ('rating_count', 'INTEGER', 'DEFAULT 0', '评分总人数'),
        ('is_new', 'BOOLEAN', 'DEFAULT FALSE', '新书标记（新书上架推荐用）'),
        ('hot_score', 'FLOAT', 'DEFAULT 0.0', '热度分值（综合阅读/收藏/评分计算）'),
        ('purchase_url_jd', 'VARCHAR(512)', '', '京东购买链接（模块四）'),
        ('purchase_url_dd', 'VARCHAR(512)', '', '当当购买链接（模块四）'),
        ('purchase_url_tb', 'VARCHAR(512)', '', '淘宝购买链接（模块四）'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '入库时间'),
    ]),
    ('3、authors表结构', '作者信息表。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '作者唯一标识'),
        ('name', 'VARCHAR(128)', 'UNIQUE, NOT NULL, INDEX', '作者姓名'),
        ('bio', 'TEXT', '', '作者简介/生平'),
        ('avatar_url', 'VARCHAR(512)', '', '作者头像URL'),
    ]),
    ('4、tags表结构', '主题标签表，用于图书分类和偏好计算。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '标签唯一标识'),
        ('name', 'VARCHAR(64)', 'UNIQUE, NOT NULL, INDEX', '标签名称（如"科幻""编程"）'),
        ('category', 'VARCHAR(32)', '', '标签分类（大类，如"文学""科技"）'),
    ]),
    ('5、publishers表结构', '出版社信息表。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '出版社唯一标识'),
        ('name', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '出版社名称'),
    ]),
    ('6、series表结构', '丛书系列信息表。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '系列唯一标识'),
        ('name', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '系列名称（如"三体"系列）'),
        ('description', 'TEXT', '', '系列描述'),
    ]),
    ('7、reading_history表结构（模块一）', '阅读历史记录表，记录用户的每次阅读行为，是用户画像计算的核心数据来源。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '用户ID'),
        ('book_id', 'INTEGER', 'FK→books.id', '图书ID'),
        ('status', 'VARCHAR(20)', "DEFAULT 'read'", '阅读状态：read/reading/want_to_read'),
        ('read_at', 'DATETIME', 'DEFAULT NOW()', '阅读时间'),
    ]),
    ('8、search_logs表结构（模块一）', '搜索日志表，记录用户的搜索关键词，用于分析用户兴趣和标签偏好。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '用户ID（未登录搜索为NULL）'),
        ('keyword', 'VARCHAR(256)', 'NOT NULL', '搜索关键词'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '搜索时间'),
    ]),
    ('9、bookmarks表结构（模块一/四）', '书架收藏表，记录用户将图书收藏到哪个书架。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '收藏记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '用户ID'),
        ('book_id', 'INTEGER', 'FK→books.id', '图书ID'),
        ('shelf_name', 'VARCHAR(64)', "DEFAULT '默认书架'", '所在书架名称（想读/在读/已读/自定义）'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '收藏时间'),
    ]),
    ('10、reading_progress表结构（模块一/四）', '阅读进度同步表，使用联合唯一索引支持多端进度同步的upsert操作。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '用户ID'),
        ('book_id', 'INTEGER', 'FK→books.id', '图书ID'),
        ('progress_percent', 'FLOAT', 'DEFAULT 0.0', '进度百分比（0.0~100.0）'),
        ('current_page', 'INTEGER', 'DEFAULT 0', '当前阅读页码'),
        ('updated_at', 'DATETIME', 'ON UPDATE NOW()', '最后同步时间'),
    ]),
    ('11、user_ratings表结构（模块一）', '用户评分表，使用联合唯一索引确保同一用户对同一本书只有一条评分记录。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '评分记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '用户ID'),
        ('book_id', 'INTEGER', 'FK→books.id', '图书ID'),
        ('rating', 'FLOAT', 'NOT NULL', '评分值（0.5~5.0，步长0.5）'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '评分时间'),
    ]),
    ('12、book_comments表结构（模块四）', '书评评论表，存储用户发表的图书评论。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '评论ID'),
        ('user_id', 'INTEGER', 'FK→users.id, INDEX', '评论者用户ID'),
        ('book_id', 'INTEGER', 'FK→books.id, INDEX', '被评论图书ID'),
        ('content', 'TEXT', 'NOT NULL', '评论内容（最大5000字符）'),
        ('likes_count', 'INTEGER', 'DEFAULT 0', '被点赞数（聚合字段，由触发器更新）'),
        ('is_pinned', 'BOOLEAN', 'DEFAULT FALSE', '管理员置顶标记'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '评论时间'),
    ]),
    ('13、comment_likes表结构（模块四）', '评论点赞关联表，记录用户对评论的点赞关系。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '点赞记录ID'),
        ('user_id', 'INTEGER', 'FK→users.id', '点赞用户ID'),
        ('comment_id', 'INTEGER', 'FK→book_comments.id', '被点赞评论ID'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '点赞时间'),
    ]),
    ('14、chat_history表结构（模块五）', 'AI对话历史表，存储用户与智能问答助手的多轮对话记录。', [
        ('id', 'INTEGER', 'PK, AUTO_INCREMENT', '消息ID'),
        ('user_id', 'INTEGER', 'FK→users.id, CASCADE', '用户ID（级联删除）'),
        ('role', 'ENUM', "('user','assistant')", '消息角色：user=用户提问，assistant=AI回答'),
        ('content', 'TEXT', 'NOT NULL', '消息内容'),
        ('intent_type', 'VARCHAR(32)', '', '意图类型（function_qa/book_rec等）'),
        ('created_at', 'DATETIME', 'DEFAULT NOW()', '消息时间'),
    ]),
    ('15、中间表：book_author', '图书-作者多对多关联中间表。', [
        ('book_id', 'INTEGER', 'PK, FK→books.id', '图书ID（联合主键）'),
        ('author_id', 'INTEGER', 'PK, FK→authors.id', '作者ID（联合主键）'),
    ]),
    ('16、中间表：book_tag', '图书-标签多对多关联中间表。', [
        ('book_id', 'INTEGER', 'PK, FK→books.id', '图书ID（联合主键）'),
        ('tag_id', 'INTEGER', 'PK, FK→tags.id', '标签ID（联合主键）'),
    ]),
]

for title, desc, fields in tables_def:
    H(doc, title, 4)
    if desc: P(doc, desc)
    styled_table(doc,
        ['序号', '列名', '数据类型', '约束/默认值', '注释'],
        [[str(j), *f] for j, f in enumerate(fields, 1)]
    )

H(doc, '三、外键信息', 2)
P(doc, '下表列出数据库中所有外键约束关系，用于维护数据引用完整性。')
styled_table(doc,
    ['外键名称', '父表(主键表)', '父键列', '子表(外键表)', '外键列', '关系基数', '说明'],
    [
        ['fk_history_user', 'users', 'id', 'reading_history', 'user_id', '1..*', '一个用户可有多条阅读历史'],
        ['fk_history_book', 'books', 'id', 'reading_history', 'book_id', '1..*', '一本书可被多个用户阅读'],
        ['fk_bookmark_user', 'users', 'id', 'bookmarks', 'user_id', '1..*', '一个用户可收藏多本书'],
        ['fk_bookmark_book', 'books', 'id', 'bookmarks', 'book_id', '1..*', '一本书可被多个用户收藏'],
        ['fk_progress_user', 'users', 'id', 'reading_progress', 'user_id', '1..*', '一个用户可有多个阅读进度'],
        ['fk_progress_book', 'books', 'id', 'reading_progress', 'book_id', '1..*', '一本书可被多个用户阅读'],
        ['fk_rating_user', 'users', 'id', 'user_ratings', 'user_id', '1..*', '一个用户可有多个评分'],
        ['fk_rating_book', 'books', 'id', 'user_ratings', 'book_id', '1..*', '一本书可获多个评分'],
        ['fk_comment_user', 'users', 'id', 'book_comments', 'user_id', '1..*', '一个用户可发表多条评论'],
        ['fk_comment_book', 'books', 'id', 'book_comments', 'book_id', '1..*', '一本书可有多条评论'],
        ['fk_like_user', 'users', 'id', 'comment_likes', 'user_id', '1..*', '一个用户可点赞多条评论'],
        ['fk_like_comment', 'book_comments', 'id', 'comment_likes', 'comment_id', '1..*', '一条评论可被多个用户点赞'],
        ['fk_chat_user', 'users', 'id', 'chat_history', 'user_id', '1..*', '一个用户可有多条对话记录（级联删除）'],
        ['fk_book_publisher', 'publishers', 'id', 'books', 'publisher_id', '1..*', '一个出版社出版多本书'],
        ['fk_book_series', 'series', 'id', 'books', 'series_id', '0..1..*', '一本书可选属于一个丛书系列'],
        ['fk_ba_book', 'books', 'id', 'book_author', 'book_id', '*..*', '图书-作者多对多关联'],
        ['fk_ba_author', 'authors', 'id', 'book_author', 'author_id', '*..*', '作者-图书多对多关联'],
        ['fk_bt_book', 'books', 'id', 'book_tag', 'book_id', '*..*', '图书-标签多对多关联'],
        ['fk_bt_tag', 'tags', 'id', 'book_tag', 'tag_id', '*..*', '标签-图书多对多关联'],
    ]
)

doc.add_page_break()

# ═══ PART 7: 补充设计和说明 ═══
H(doc, '第七部分  补充设计和说明', 1)

H(doc, '一、编译运行环境设计', 2)

H(doc, '1、前端运行环境设计', 3)
styled_table(doc,
    ['环境项', '配置要求', '说明'],
    [
        ['操作系统', 'Windows 10+ / macOS 11+ / Linux', '开发环境支持三大平台'],
        ['Node.js', 'v18.0.0 及以上', 'JavaScript运行时环境'],
        ['npm', 'v9.0.0 及以上', '包管理器'],
        ['HBuilderX', '最新稳定版', 'uni-app官方IDE（可选，也可用VSCode）'],
        ['微信开发者工具', '最新稳定版', '微信小程序调试预览（可选）'],
        ['Android Studio', '最新稳定版', 'Android打包（可选）'],
        ['Xcode', '最新稳定版', 'iOS打包（仅macOS，可选）'],
    ]
)

H(doc, '2、后端运行环境设计', 3)
styled_table(doc,
    ['环境项', '配置要求', '说明'],
    [
        ['操作系统', 'Windows 10+ / macOS 11+ / Linux', '开发和生产环境'],
        ['Python', '3.11 及以上', '解释器语言'],
        ['pip', '最新版', 'Python包管理器'],
        ['Docker', '24.x 及以上', '容器运行时（用于MySQL/Neo4j/Redis/ES）'],
        ['Docker Compose', 'v2.x 及以上', '多容器编排'],
        ['MySQL', '8.0 (Docker)', '关系型数据库，端口3306'],
        ['Neo4j', '5.x (Docker)', '图数据库，Bolt端口7687，HTTP端口7474'],
        ['Redis', '7.x (Docker)', '缓存数据库，端口6379'],
        ['ElasticSearch', '8.x (Docker)', '搜索引擎，端口9200'],
        ['Git', '2.x 及以上', '版本控制'],
    ]
)

H(doc, '二、包路径与目录结构设计', 2)

H(doc, '1、前端目录结构设计', 3)
code_block(doc, '''frontend-uni/                        # uni-app前端项目根目录
├── App.vue                           # 根组件（全局样式、生命周期）
├── main.js                           # 应用入口（创建app、注册Pinia）
├── pages.json                        # 页面路由配置（Tab栏+页面栈）
├── manifest.json                     # 应用配置（AppID/权限/图标）
├── package.json                      # npm依赖声明
├── pages/                            # 页面组件
│   ├── index/index.vue               #   首页（推荐流/搜索/轮播）
│   ├── detail/detail.vue             #   图书详情（试读/评论/收藏）
│   ├── shelf/shelf.vue               #   书架管理（分类/批量操作）
│   ├── chat/chat.vue                 #   AI问答助手（对话界面）
│   ├── profile/profile.vue           #   个人中心（统计/趋势/设置）
│   └── admin/admin.vue               #   管理后台（数据管理面板）
├── components/                       # 复用组件
│   ├── book-card.vue                 #   图书卡片组件
│   ├── comment-list.vue              #   评论列表组件
│   └── chat-widget.vue               #   聊天悬浮窗组件
├── api/index.js                      # API封装层（统一请求/响应拦截）
├── store/index.js                    # Pinia Store（user/shelf/chat）
└── utils/                            # 工具函数
    ├── auth.js                       #   JWT Token管理
    └── request.js                    #   axios实例（BaseURL/拦截器）''')

H(doc, '2、后端目录结构设计', 3)
code_block(doc, '''app/                                # FastAPI后端主目录
├── main.py                           # 应用主入口（FastAPI实例/CORS/路由注册）
├── core/                             # 核心基础设施层
│   ├── config.py                     #   Settings配置中心（环境变量/.env加载）
│   ├── database.py                   #   数据库连接管理（MySQL/Neo4j/Redis驱动）
│   └── security.py                   #   JWT生成验证 + bcrypt密码哈希
├── models/                           # SQLAlchemy ORM数据模型层
│   ├── user.py                       #   模块一：User/ReadingHistory/Bookmark/
│   │                                 #          ReadingProgress/UserRating/SearchLog
│   ├── book.py                       #   模块二：Book/Author/Publisher/Tag/Series
│   │                                 #          +book_author/book_tag中间表
│   ├── ecosystem.py                  #   模块四：BookComment/CommentLike
│   └── chat.py                       #   模块五：ChatHistory
├── schemas/                          # Pydantic请求/响应验证模型层
│   ├── user.py                       #   模块一：注册/登录/画像/行为/进度Schema
│   ├── book.py                       #   模块二：图书详情/图谱查询/实体创建Schema
│   ├── recommend.py                  #   模块三：推荐请求/响应/权重/接口契约Schema
│   ├── ecosystem.py                  #   模块四：评论/试读/购书/书架/统计Schema
│   └── chat.py                       #   模块五：对话请求/响应/历史Schema
├── services/                         # 业务逻辑层（核心）
│   ├── user_service.py               #   模块一：认证/行为采集/画像建模/进度同步
│   ├── graph_service.py              #   模块二：图谱构建/路径推理(5条Cypher)/可视化
│   ├── recommend_service.py          #   模块三：CF推荐/KG推荐/混合融合/理由生成
│   ├── ecosystem_service.py          #   模块四：试读/书评/购书/书架管理/统计
│   └── ai_chat_service.py            #   模块五：意图识别/上下文检索/LLM对话
└── api/                              # API路由层（HTTP端点）
    ├── deps.py                       #   依赖注入（get_current_user/admin/optional）
    └── v1/
        ├── router.py                 #   API v1总路由（注册五大模块路由）
        └── endpoints/
            ├── user.py               #     模块一：13个端点
            ├── graph.py              #     模块二：6个端点
            ├── recommend.py          #     模块三：4个端点
            ├── ecosystem.py          #     模块四：15个端点
            └── ai_chat.py            #     模块五：3个端点''')

# ── Save ──
output_path = os.path.join(OUTPUT_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v2.0_struct.docx')
doc.save(output_path)
print(f"Saved: {output_path}")

# Cleanup
shutil.rmtree(TEMP_DIR, ignore_errors=True)
print("Done!")
