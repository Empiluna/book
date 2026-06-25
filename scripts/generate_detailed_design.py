"""
生成《基于知识图谱的个性化荐书系统》详细设计说明书
依据: D04_详细设计-1.pdf + D05_详细设计-2.pdf 要求
参考: 03-《升学空间站》详细设计文档_软工方向样例.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ═══════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        for p in cell.paragraphs:
            p.style = cell._tc.getparent().getparent().get_or_add_pPr() if hasattr(cell._tc, 'get_or_add_pPr') else None
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if bold:
                    run.bold = True
        if header:
            set_cell_shading(cell, '2F5496')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    return row

def create_styled_table(doc, headers, rows_data, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], '2F5496')
    # 数据行
    for row_data in rows_data:
        add_table_row(table, row_data)
    doc.add_paragraph()
    return table

def add_code_block(doc, code_text, title=""):
    """添加代码块（Mermaid图或代码）"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 添加底纹
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    return p

def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_body_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_mermaid_diagram(doc, diagram_code, title, description=""):
    """添加Mermaid图（用代码块表示，后续可渲染）"""
    if description:
        add_body_text(doc, description)
    add_code_block(doc, diagram_code, title)
    p = doc.add_paragraph()
    run = p.add_run("▲ 上图可通过 Mermaid 渲染工具（如 mermaid.live）渲染为矢量图")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

# ═══════════════════════════════════════════════
# 主文档生成
# ═══════════════════════════════════════════════

def generate_document():
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ═══════════════════════════════════
    # 封面
    # ═══════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('基于知识图谱的个性化荐书系统')
    run.font.size = Pt(28)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('详 细 设 计 说 明 书')
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        f'版本：v1.0',
        f'日期：{datetime.now().strftime("%Y年%m月%d日")}',
        '项目团队：成员A · 成员B · 成员C · 成员D',
        '文档状态：评审中',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ═══════════════════════════════════
    # 目录页
    # ═══════════════════════════════════
    add_heading_styled(doc, '目  录', 1)
    toc_items = [
        ('第一部分', '引言', 1),
        ('  一、', '编写目的', 2),
        ('  二、', '项目背景', 2),
        ('  三、', '定义', 2),
        ('  四、', '参考资料', 2),
        ('第二部分', '项目概述', 1),
        ('  一、', '系统功能概述', 2),
        ('  二、', '模块划分', 2),
        ('第三部分', '总体设计', 1),
        ('  一、', '技术架构设计', 2),
        ('  二、', '分层架构设计', 2),
        ('  三、', '核心控制流程', 2),
        ('  四、', '开发环境配置', 2),
        ('第四部分', '数据库设计', 1),
        ('  一、', '数据库选型与双库架构', 2),
        ('  二、', 'ER图设计', 2),
        ('  三、', 'MySQL数据表结构设计', 2),
        ('  四、', 'Neo4j图谱Schema设计', 2),
        ('  五、', '接口表设计', 2),
        ('第五部分', '界面设计', 1),
        ('  一、', '前端技术栈概述', 2),
        ('  二、', '页面框架结构设计', 2),
        ('  三、', '主要页面设计', 2),
        ('第六部分', '单元模块设计', 1),
        ('  一、', '模块一：用户画像模块', 2),
        ('  二、', '模块二：知识图谱模块', 2),
        ('  三、', '模块三：个性化推荐模块', 2),
        ('  四、', '模块四：阅读生态模块', 2),
        ('  五、', '模块五：智能问答助手模块', 2),
        ('第七部分', '数据访问层设计', 1),
        ('第八部分', '功能模块接口设计', 1),
        ('第九部分', '算法设计', 1),
        ('第十部分', '项目文件目录结构', 1),
    ]
    for num, title_text, level in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}{title_text}')
        run.font.size = Pt(12) if level == 1 else Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if level == 1:
            run.bold = True
        p.paragraph_format.line_spacing = 1.8

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第一部分：引言
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第一部分  引言', 1)

    add_heading_styled(doc, '一、编写目的', 2)
    add_body_text(doc, '编写本文的主要目的是把需求分析得到的用例模型和概要设计确定的技术方案，转换为具体的软件结构和数据结构。设计软件结构的具体任务是：将一个复杂系统按功能进行模块划分、建立模块的层次结构及调用关系、确定模块间的接口及人机界面等。数据结构设计包括数据特征的描述、确定数据的结构特性、以及数据库的设计。')
    add_body_text(doc, '本设计是指导详细设计和项目实施的重要指导性文件，也是后续项目实施、系统集成测试的核心指导文件以及系统集成测试的重要依据。本文档面向项目开发团队（成员A/B/C/D），为各模块编码实现提供精确的技术规范。')

    add_heading_styled(doc, '二、项目背景', 2)
    add_body_text(doc, '在信息爆炸的时代，读者面临"选书难"的痛点——面对海量图书，如何快速找到适合自己的读物？传统的图书推荐多基于简单的分类浏览或畅销榜单，缺乏个性化和可解释性。')
    add_body_text(doc, '本系统"基于知识图谱的个性化荐书系统"利用知识图谱技术，将图书、作者、标签、出版社、丛书等实体及其关系构建成语义网络，并结合用户画像和协同过滤算法，实现可解释的个性化图书推荐。系统支持Web端和移动端（uni-app），提供从发现图书、在线阅读到书评互动的完整阅读生态。')
    add_body_text(doc, '系统采用FastAPI作为后端框架，MySQL+Neo4j双数据库架构（关系型存储实体属性，图数据库存储语义关系），Redis缓存加速，并集成大语言模型提供智能问答助手。')

    add_heading_styled(doc, '三、定义', 2)
    add_heading_styled(doc, '1、技术类', 3)

    tech_terms = [
        ('FastAPI', '基于Python的现代Web框架，支持异步处理、自动生成Swagger文档、Pydantic数据验证。本系统选择FastAPI而非Flask/Django，主要因为其异步支持和高性能特性。'),
        ('Neo4j', '图数据库管理系统，使用Cypher查询语言。本系统使用Neo4j存储图书、作者、标签、出版社等实体及它们之间的关系图谱，支持多跳路径推理。'),
        ('MySQL', '关系型数据库管理系统。本系统使用MySQL存储用户、图书属性、评分、评论、阅读记录等结构化数据。'),
        ('Redis', '内存键值存储数据库。本系统使用Redis缓存热门推荐结果、用户Session、相似度矩阵等高频访问数据。'),
        ('SQLAlchemy', 'Python ORM框架，提供对象关系映射。本系统使用SQLAlchemy管理MySQL数据库操作。'),
        ('JWT', 'JSON Web Token，用于无状态用户认证。用户登录后获得Token，后续API请求在Header中携带Token进行身份验证。'),
        ('ItemCF', '基于物品的协同过滤算法。通过计算图书间的相似度（余弦相似度），从用户已评分的高分图书出发，推荐相似图书。'),
        ('uni-app', '跨平台前端框架，一套代码编译到Android/iOS/小程序/H5。本系统使用uni-app开发移动端。'),
        ('LLM', '大语言模型（Large Language Model）。本系统模块五集成LLM实现智能问答助手，支持自然语言荐书和功能咨询。'),
    ]
    for name, desc in tech_terms:
        p = doc.add_paragraph()
        run = p.add_run(f'（1）{name}')
        run.bold = True
        run.font.size = Pt(12)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(desc)
        run2.font.size = Pt(12)
        p2.paragraph_format.first_line_indent = Cm(0.74)

    add_heading_styled(doc, '2、业务类', 3)
    biz_terms = [
        ('用户画像', '通过分析用户的阅读历史、搜索记录、收藏偏好和评分数据，构建用户的标签偏好向量、作者偏好和类别偏好，为个性化推荐引擎提供数据支撑。'),
        ('知识图谱推理', '从用户感兴趣的高分图书出发，沿图谱中的作者、标签、系列、出版社等关系路径进行多跳探索，发现候选推荐图书并生成可解释的推理路径。'),
        ('混合推荐', '融合知识图谱推理推荐（40%）、协同过滤推荐（40%）、热门推荐（10%）和新书推荐（10%）四种策略，按可配置权重进行加权融合，生成最终推荐列表。'),
        ('推荐理由生成', '根据推荐来源路径类型（同作者/同标签/同系列/协同过滤/热门等），自动生成自然语言的推荐理由，增强推荐的可解释性和用户信任度。'),
        ('阅读生态', '涵盖电子书试读、书评社区、实体书购书链接、书架管理、阅读统计等面向用户的完整功能集合。'),
        ('智能问答助手', '基于大语言模型的对话式助手，支持功能问答、自然语言荐书、图书知识问答、个人阅读查询和管理员操作指引等场景。'),
    ]
    for name, desc in biz_terms:
        p = doc.add_paragraph()
        run = p.add_run(f'（{biz_terms.index((name, desc))+1}）{name}')
        run.bold = True
        run.font.size = Pt(12)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(desc)
        run2.font.size = Pt(12)
        p2.paragraph_format.first_line_indent = Cm(0.74)

    add_heading_styled(doc, '四、参考资料', 2)
    refs = [
        '《基于知识图谱的个性化荐书系统-需求说明书》v1.0',
        '《基于知识图谱的个性化荐书系统》概要设计说明书 v3.2',
        '《升学空间站》详细设计文档_软工方向样例',
        '实训任务04-详细设计-1（D04）— 详细设计要求与规范',
        '实训任务05-详细设计-2（D05）— 详细设计内容与交付物',
        'FastAPI官方文档：https://fastapi.tiangolo.com/',
        'Neo4j Python Driver文档：https://neo4j.com/docs/python-manual/current/',
        'SQLAlchemy ORM文档：https://docs.sqlalchemy.org/',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] {ref}')
        run.font.size = Pt(10.5)

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第二部分：项目概述
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第二部分  项目概述', 1)

    add_heading_styled(doc, '一、系统功能概述', 2)
    add_body_text(doc, '《基于知识图谱的个性化荐书系统》旨在利用知识图谱技术和混合推荐算法，为读者提供可解释的个性化图书推荐服务。系统围绕五大核心模块构建：')
    add_body_text(doc, '（1）用户画像模块：负责用户注册登录、阅读行为采集（历史记录、搜索、收藏、评分）、用户兴趣建模（标签偏好向量、作者偏好、类别偏好）以及阅读进度同步。是系统的数据底座，为推荐引擎提供用户偏好数据。')
    add_body_text(doc, '（2）知识图谱模块：负责构建和维护书籍知识图谱（Book-Author-Tag-Publisher-Series实体及关系），实现多跳路径查询与推理（5条Cypher查询路径），为推荐引擎提供图谱推理能力。同时提供图谱可视化数据接口。')
    add_body_text(doc, '（3）个性化推荐模块：系统的核心算法模块。实现ItemCF协同过滤推荐、知识图谱推理推荐、热门推荐和新书推荐四种策略，并通过加权融合生成混合推荐结果。支持首页推荐、相似图书推荐、策略切换和推荐理由生成。')
    add_body_text(doc, '（4）阅读生态模块：面向用户的交互功能集合，包括电子书在线试读（权限分级：未登录3页/登录10页）、书评社区（发表/点赞/置顶/删除）、实体书购书链接（多平台比价跳转）、书架收藏管理（自定义书架/移动/删除）和阅读统计分析。')
    add_body_text(doc, '（5）智能问答助手模块：基于大语言模型的对话式助手，实现意图识别、知识检索增强、LLM回答生成和对话历史管理。支持功能问答、自然语言荐书、图书知识问答、个人阅读查询和管理员操作指引六大场景。')

    add_heading_styled(doc, '二、模块划分', 2)
    create_styled_table(doc,
        ['模块编号', '模块名称', '负责人', '核心职责', '依赖关系'],
        [
            ['模块一', '用户画像', '成员A', '用户认证、行为采集、兴趣建模、进度同步', '被模块三、模块四依赖'],
            ['模块二', '知识图谱', '成员B', '图谱构建、路径推理、候选发现、可视化', '被模块三依赖'],
            ['模块三', '个性化推荐', '成员C', '协同过滤、图谱推荐、混合融合、理由生成', '依赖模块一、模块二'],
            ['模块四', '阅读生态', '成员D', '试读、书评、购书、书架、统计', '依赖模块一(JWT认证)'],
            ['模块五', '智能问答', 'ALL', '意图识别、LLM回答、对话管理', '依赖模块一+系统数据'],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第三部分：总体设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第三部分  总体设计', 1)

    add_heading_styled(doc, '一、技术架构设计', 2)
    add_body_text(doc, '系统采用B/S架构与前后端分离设计模式。后端基于FastAPI框架，前端采用uni-app跨平台框架。整体架构分为表现层、业务逻辑层、数据服务层三层，同时引入缓存层和搜索层。各层之间通过明确定义的接口契约进行通信。')

    add_mermaid_diagram(doc, '''```mermaid
graph TB
    subgraph 表现层["表现层 (Presentation Layer)"]
        Web["Web前端<br/>(Vue/uni-app)"]
        Mobile["移动端<br/>(Android/iOS/小程序)"]
        H5["H5端"]
    end

    subgraph 网关层["网关层"]
        Nginx["Nginx<br/>反向代理/负载均衡"]
    end

    subgraph 业务层["业务逻辑层 (Business Layer)"]
        M1["模块一<br/>用户画像"]
        M2["模块二<br/>知识图谱"]
        M3["模块三<br/>个性化推荐"]
        M4["模块四<br/>阅读生态"]
        M5["模块五<br/>智能问答助手"]
    end

    subgraph 数据层["数据服务层 (Data Layer)"]
        MySQL["MySQL 8.0<br/>关系型数据库<br/>(用户/图书/评论)"]
        Neo4j["Neo4j<br/>图数据库<br/>(知识图谱)"]
        Redis["Redis<br/>缓存<br/>(Session/推荐缓存)"]
        ES["ElasticSearch<br/>搜索引擎<br/>(全文检索)"]
    end

    subgraph 外部服务["外部服务"]
        LLM_API["LLM API<br/>(OpenAI兼容)"]
        BookStore["购书平台<br/>(京东/当当/淘宝)"]
    end

    Web --> Nginx
    Mobile --> Nginx
    H5 --> Nginx
    Nginx --> M1
    Nginx --> M2
    Nginx --> M3
    Nginx --> M4
    Nginx --> M5

    M1 --> MySQL
    M1 --> Redis
    M2 --> Neo4j
    M2 --> MySQL
    M3 --> MySQL
    M3 --> Redis
    M4 --> MySQL
    M5 --> MySQL
    M5 --> LLM_API
    M4 --> BookStore

    M3 -.->|接口契约| M1
    M3 -.->|接口契约| M2
    M4 -.->|JWT认证| M1
```''', '图3-1 系统技术架构图', '系统技术架构如图3-1所示，展示了五层架构及各模块间的关系。')

    add_heading_styled(doc, '二、分层架构设计', 2)

    add_heading_styled(doc, '1、表现层 (Presentation Layer)', 3)
    add_body_text(doc, '表现层基于uni-app框架构建，采用Vue 3组件化开发模式。前端通过RESTful API与后端交互，使用axios进行HTTP请求，通过Pinia进行状态管理，使用vue-router进行页面路由。支持一套代码编译到Android、iOS、微信小程序和H5多端运行。')

    add_heading_styled(doc, '2、业务逻辑层 (Business Logic Layer)', 3)
    add_body_text(doc, '业务逻辑层基于FastAPI框架实现，采用分层架构：')
    add_body_text(doc, '（1）API层（endpoints/）：处理HTTP请求/响应，参数验证，路由分发。使用Pydantic模型进行请求体和响应体的序列化与验证。')
    add_body_text(doc, '（2）服务层（services/）：封装核心业务逻辑。各模块的业务逻辑独立在各自的service文件中实现，通过函数调用或类方法组织。服务层不直接处理HTTP请求，便于单元测试。')
    add_body_text(doc, '（3）数据访问层（models/ + core/database.py）：通过SQLAlchemy ORM操作MySQL，通过Neo4j Python Driver操作图数据库，通过redis-py操作缓存。数据库会话通过FastAPI依赖注入（Depends）管理生命周期。')

    add_heading_styled(doc, '3、数据服务层 (Data Service Layer)', 3)
    add_body_text(doc, '数据服务层包含四种存储引擎：')
    add_body_text(doc, '（1）MySQL 8.0：存储结构化关系数据——用户信息、图书属性（书名、ISBN、简介）、作者信息、出版社、标签、阅读历史、评分、评论、收藏、阅读进度等。使用SQLAlchemy ORM进行对象关系映射。')
    add_body_text(doc, '（2）Neo4j：存储知识图谱——Book、Author、Tag、Publisher、Series五种实体节点，以及AUTHORED、TAGGED、PUBLISHED、SERIES_OF、SIMILAR五种关系类型。使用Cypher查询语言进行多跳路径推理。')
    add_body_text(doc, '（3）Redis：缓存层——热门推荐结果缓存、用户会话缓存、相似度矩阵缓存。减少数据库查询压力，提升响应速度。')
    add_body_text(doc, '（4）ElasticSearch：搜索引擎——图书全文检索，支持中文分词。为知识图谱模块的搜索功能提供支持。')

    add_heading_styled(doc, '三、核心控制流程', 2)

    add_heading_styled(doc, '1、用户登录认证流程', 3)
    add_mermaid_diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant F as 前端(uni-app)
    participant API as FastAPI后端
    participant DB as MySQL

    U->>F: 输入用户名密码
    F->>API: POST /api/v1/user/login
    API->>DB: 查询用户表
    DB-->>API: 返回用户记录
    API->>API: verify_password(bcrypt)
    alt 密码正确
        API->>API: create_access_token(JWT)
        API-->>F: 返回Token + 用户信息
        F->>F: 存储Token到本地
        F-->>U: 跳转首页
    else 密码错误
        API-->>F: 401 用户名或密码错误
        F-->>U: 显示错误提示
    end
```''', '图3-2 用户登录认证时序图')

    add_heading_styled(doc, '2、个性化推荐流程', 3)
    add_mermaid_diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant F as 前端
    participant R as 推荐模块(模块三)
    participant M1 as 用户画像(模块一)
    participant M2 as 知识图谱(模块二)
    participant DB as MySQL/Neo4j/Redis

    U->>F: 访问首页
    F->>R: GET /api/v1/recommend/home
    R->>M1: GET /api/v1/user/profile
    M1->>DB: 查询阅读历史/收藏/评分
    DB-->>M1: 行为数据
    M1->>M1: 计算标签偏好向量
    M1-->>R: 用户画像(tag_weights, authors, books)

    loop 每本高分图书
        R->>M2: POST /api/v1/graph/paths
        M2->>DB: Cypher多跳路径查询
        DB-->>M2: 候选图书+推理路径
        M2-->>R: 图谱推荐结果
    end

    R->>R: ItemCF协同过滤
    R->>R: 热门推荐 + 新书推荐
    R->>R: 加权融合(0.4+0.4+0.1+0.1)
    R->>R: 生成推荐理由
    R-->>F: 推荐列表(含理由)
    F-->>U: 展示个性化推荐
```''', '图3-3 个性化推荐时序图')

    add_heading_styled(doc, '3、智能问答对话流程', 3)
    add_mermaid_diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant F as 前端
    participant C as 问答模块(模块五)
    participant LLM as LLM API
    participant DB as MySQL

    U->>F: 输入自然语言问题
    F->>C: POST /api/v1/chat/message
    C->>C: 意图识别(LLM/关键词)
    alt 超出业务范围
        C-->>F: 礼貌拒绝+引导
    else 合法意图
        C->>DB: 检索业务上下文
        DB-->>C: 图书/用户/功能数据
        C->>DB: 获取对话历史(最近30条)
        DB-->>C: 历史消息
        C->>LLM: System Prompt + 上下文 + 问题
        LLM-->>C: AI回答
        C->>DB: 保存对话记录
        C-->>F: 回答+建议追问
    end
    F-->>U: 显示AI回复
```''', '图3-4 智能问答对话时序图')

    add_heading_styled(doc, '四、开发环境配置', 2)
    create_styled_table(doc,
        ['类别', '工具/技术', '版本/说明'],
        [
            ['后端框架', 'FastAPI', 'Python 3.11+'],
            ['ORM', 'SQLAlchemy', '2.0+'],
            ['关系数据库', 'MySQL', '8.0 (Docker)'],
            ['图数据库', 'Neo4j', '5.x (Docker)'],
            ['缓存', 'Redis', '7.x (Docker)'],
            ['搜索引擎', 'ElasticSearch', '8.x (Docker)'],
            ['认证', 'JWT (python-jose)', 'Bearer Token, 24h过期'],
            ['密码哈希', 'bcrypt (passlib)', '—'],
            ['前端框架', 'uni-app (Vue 3)', 'HBuilderX / CLI'],
            ['状态管理', 'Pinia', '—'],
            ['HTTP客户端', 'axios', '—'],
            ['LLM', 'OpenAI兼容API', 'Claude/GPT等'],
            ['容器化', 'Docker + Docker Compose', '—'],
            ['开发服务器', 'Uvicorn', '--reload热重载'],
            ['API文档', 'Swagger + ReDoc', '自动生成'],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第四部分：数据库设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第四部分  数据库设计', 1)

    add_heading_styled(doc, '一、数据库选型与双库架构', 2)
    add_body_text(doc, '系统采用MySQL + Neo4j双数据库架构。MySQL负责存储结构化实体属性数据（用户信息、图书元数据、阅读记录、评分评论等关系型数据），Neo4j负责存储实体间的语义关系图谱（Book-Author-Tag-Publisher-Series关系网络）。两个数据库通过统一的实体ID（book_id, author_id, tag_id等）进行关联。')
    add_body_text(doc, '选择双数据库架构的原因：（1）图书推荐场景天然适合图结构建模（作者写了哪些书、书有哪些标签等）；（2）Neo4j的Cypher查询语言对多跳路径查询有优化，性能优于关系数据库的多次JOIN；（3）MySQL在事务处理、用户管理、评论等典型关系型场景中表现优秀。')

    add_heading_styled(doc, '二、ER图设计', 2)
    add_body_text(doc, '以下ER图展示了MySQL数据库中核心实体及其关系。Neo4j图谱Schema在下一节单独说明。')

    add_mermaid_diagram(doc, '''```mermaid
erDiagram
    User ||--o{ ReadingHistory : "有"
    User ||--o{ Bookmark : "收藏"
    User ||--o{ ReadingProgress : "记录进度"
    User ||--o{ UserRating : "评分"
    User ||--o{ BookComment : "发表评论"
    User ||--o{ ChatHistory : "对话"

    Book ||--o{ ReadingHistory : "被阅读"
    Book ||--o{ Bookmark : "被收藏"
    Book ||--o{ ReadingProgress : "有进度"
    Book ||--o{ UserRating : "被评分"
    Book ||--o{ BookComment : "有评论"
    Book }o--|| Publisher : "属于"
    Book }o--|| Series : "属于"
    Book }o--o{ Author : "多对多"
    Book }o--o{ Tag : "多对多"

    BookComment ||--o{ CommentLike : "被点赞"
    User ||--o{ CommentLike : "点赞"

    User {
        int id PK
        string username UK
        string email UK
        string hashed_password
        bool is_active
        bool is_admin
        datetime created_at
    }

    Book {
        int id PK
        string title
        string isbn UK
        int publisher_id FK
        int series_id FK
        float avg_rating
        float hot_score
        string cover_url
        text description
    }

    Author {
        int id PK
        string name UK
        text bio
    }

    Tag {
        int id PK
        string name UK
        string category
    }

    Publisher {
        int id PK
        string name UK
    }

    Series {
        int id PK
        string name UK
        text description
    }

    ReadingHistory {
        int id PK
        int user_id FK
        int book_id FK
        string status
        datetime read_at
    }

    UserRating {
        int id PK
        int user_id FK
        int book_id FK
        float rating
        datetime created_at
    }

    BookComment {
        int id PK
        int user_id FK
        int book_id FK
        text content
        int likes_count
        bool is_pinned
        datetime created_at
    }

    Bookmark {
        int id PK
        int user_id FK
        int book_id FK
        string shelf_name
        datetime created_at
    }

    ReadingProgress {
        int id PK
        int user_id FK
        int book_id FK
        float progress_percent
        int current_page
        datetime updated_at
    }

    ChatHistory {
        int id PK
        int user_id FK
        enum role
        text content
        string intent_type
        datetime created_at
    }
```''', '图4-1 MySQL数据库ER图', '图4-1展示了系统中11张核心数据表及其关系。为简化展示，SearchLog表和CommentLike中间表的关系已在主体关系中体现。')

    add_heading_styled(doc, '三、MySQL数据表结构设计', 2)

    # Users表
    add_heading_styled(doc, '1、用户表（users）', 3)
    add_body_text(doc, '存储系统用户的基本信息，是用户认证和权限管理的基础表。')
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '用户唯一标识'],
            ['username', 'VARCHAR(64)', 'UNIQUE, NOT NULL, INDEX', '用户名，用于登录'],
            ['email', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '邮箱地址'],
            ['hashed_password', 'VARCHAR(256)', 'NOT NULL', 'bcrypt哈希后的密码'],
            ['is_active', 'BOOLEAN', 'DEFAULT TRUE', '账户启用状态'],
            ['is_admin', 'BOOLEAN', 'DEFAULT FALSE', '管理员标识'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '注册时间'],
            ['updated_at', 'DATETIME', 'ON UPDATE NOW()', '更新时间'],
        ]
    )

    # Books表
    add_heading_styled(doc, '2、图书表（books）', 3)
    add_body_text(doc, '存储图书的核心元数据，是系统中数据量最大的表。同时存储模块四所需的购书链接字段。')
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '图书唯一标识'],
            ['title', 'VARCHAR(256)', 'NOT NULL, INDEX', '书名'],
            ['subtitle', 'VARCHAR(256)', '', '副标题'],
            ['isbn', 'VARCHAR(20)', 'UNIQUE', 'ISBN编号'],
            ['publisher_id', 'INTEGER', 'FK → publishers.id', '出版社ID'],
            ['series_id', 'INTEGER', 'FK → series.id', '丛书系列ID'],
            ['publication_year', 'INTEGER', '', '出版年份'],
            ['description', 'TEXT', '', '图书简介'],
            ['cover_url', 'VARCHAR(512)', '', '封面图片URL'],
            ['page_count', 'INTEGER', '', '总页数'],
            ['language', 'VARCHAR(32)', "DEFAULT 'zh-CN'", '语言'],
            ['avg_rating', 'FLOAT', 'DEFAULT 0.0', '平均评分'],
            ['rating_count', 'INTEGER', 'DEFAULT 0', '评分人数'],
            ['is_new', 'BOOLEAN', 'DEFAULT FALSE', '是否新书'],
            ['hot_score', 'FLOAT', 'DEFAULT 0.0', '热度分值'],
            ['purchase_url_jd', 'VARCHAR(512)', '', '京东购买链接（模块四）'],
            ['purchase_url_dd', 'VARCHAR(512)', '', '当当购买链接（模块四）'],
            ['purchase_url_tb', 'VARCHAR(512)', '', '淘宝购买链接（模块四）'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '入库时间'],
        ]
    )

    # 作者/标签/出版社/丛书
    add_heading_styled(doc, '3、作者表（authors）', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '作者唯一标识'],
            ['name', 'VARCHAR(128)', 'UNIQUE, NOT NULL, INDEX', '作者姓名'],
            ['bio', 'TEXT', '', '作者简介'],
            ['avatar_url', 'VARCHAR(512)', '', '作者头像URL'],
        ]
    )

    add_heading_styled(doc, '4、标签表（tags）', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '标签唯一标识'],
            ['name', 'VARCHAR(64)', 'UNIQUE, NOT NULL, INDEX', '标签名称'],
            ['category', 'VARCHAR(32)', '', '标签分类（如科幻、编程）'],
        ]
    )

    add_heading_styled(doc, '5、出版社表（publishers）', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '出版社唯一标识'],
            ['name', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '出版社名称'],
        ]
    )

    add_heading_styled(doc, '6、丛书系列表（series）', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '系列唯一标识'],
            ['name', 'VARCHAR(128)', 'UNIQUE, NOT NULL', '系列名称'],
            ['description', 'TEXT', '', '系列描述'],
        ]
    )

    # 中间表
    add_heading_styled(doc, '7、多对多中间表', 3)
    add_body_text(doc, 'book_author表：关联图书与作者（多对多），字段为book_id(FK) + author_id(FK)，联合主键。')
    add_body_text(doc, 'book_tag表：关联图书与标签（多对多），字段为book_id(FK) + tag_id(FK)，联合主键。')

    # 模块一表
    add_heading_styled(doc, '8、阅读历史表（reading_history）【模块一】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID'],
            ['book_id', 'INTEGER', 'FK → books.id', '图书ID'],
            ['status', 'VARCHAR(20)', "DEFAULT 'read'", '状态: read/reading/want_to_read'],
            ['read_at', 'DATETIME', 'DEFAULT NOW()', '阅读时间'],
        ]
    )

    add_heading_styled(doc, '9、搜索日志表（search_logs）【模块一】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID（可空）'],
            ['keyword', 'VARCHAR(256)', 'NOT NULL', '搜索关键词'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '搜索时间'],
        ]
    )

    add_heading_styled(doc, '10、收藏表（bookmarks）【模块一】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '收藏ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID'],
            ['book_id', 'INTEGER', 'FK → books.id', '图书ID'],
            ['shelf_name', 'VARCHAR(64)', "DEFAULT '默认书架'", '所在书架名称'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '收藏时间'],
        ]
    )

    add_heading_styled(doc, '11、阅读进度表（reading_progress）【模块一/四】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '记录ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID'],
            ['book_id', 'INTEGER', 'FK → books.id', '图书ID'],
            ['progress_percent', 'FLOAT', 'DEFAULT 0.0', '进度百分比 0.0~100.0'],
            ['current_page', 'INTEGER', 'DEFAULT 0', '当前页码'],
            ['updated_at', 'DATETIME', 'ON UPDATE NOW()', '最后更新时间'],
        ]
    )

    add_heading_styled(doc, '12、用户评分表（user_ratings）【模块一】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '评分ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID'],
            ['book_id', 'INTEGER', 'FK → books.id', '图书ID'],
            ['rating', 'FLOAT', 'NOT NULL', '评分 0.5~5.0'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '评分时间'],
        ]
    )

    add_heading_styled(doc, '13、书评表（book_comments）【模块四】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '评论ID'],
            ['user_id', 'INTEGER', 'FK → users.id, INDEX', '用户ID'],
            ['book_id', 'INTEGER', 'FK → books.id, INDEX', '图书ID'],
            ['content', 'TEXT', 'NOT NULL', '评论内容'],
            ['likes_count', 'INTEGER', 'DEFAULT 0', '点赞数'],
            ['is_pinned', 'BOOLEAN', 'DEFAULT FALSE', '管理员置顶'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '评论时间'],
        ]
    )

    add_heading_styled(doc, '14、评论点赞表（comment_likes）【模块四】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '点赞ID'],
            ['user_id', 'INTEGER', 'FK → users.id', '用户ID'],
            ['comment_id', 'INTEGER', 'FK → book_comments.id', '评论ID'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '点赞时间'],
        ]
    )

    add_heading_styled(doc, '15、对话历史表（chat_history）【模块五】', 3)
    create_styled_table(doc,
        ['字段名', '类型', '约束', '说明'],
        [
            ['id', 'INTEGER', 'PK, AUTO_INCREMENT', '消息ID'],
            ['user_id', 'INTEGER', 'FK → users.id, CASCADE', '用户ID'],
            ['role', 'ENUM', "user/assistant", '角色：user=用户 assistant=AI'],
            ['content', 'TEXT', 'NOT NULL', '消息内容'],
            ['intent_type', 'VARCHAR(32)', '', '意图类型'],
            ['created_at', 'DATETIME', 'DEFAULT NOW()', '消息时间'],
        ]
    )

    add_heading_styled(doc, '四、Neo4j图谱Schema设计', 2)
    add_body_text(doc, '知识图谱在Neo4j中存储，包含5种实体节点和5种关系类型。图谱与MySQL通过统一的实体ID进行数据关联。')

    add_mermaid_diagram(doc, '''```mermaid
graph LR
    subgraph 实体节点
        B["Book<br/>book_id, title, isbn"]
        A["Author<br/>author_id, name"]
        T["Tag<br/>tag_id, name, category"]
        P["Publisher<br/>publisher_id, name"]
        S["Series<br/>series_id, name"]
    end

    B ---|"AUTHORED<br/>weight:1.0"| A
    B ---|"TAGGED<br/>weight:0.8"| T
    B ---|"PUBLISHED<br/>weight:0.5"| P
    B ---|"SERIES_OF<br/>weight:0.6"| S
    B ---|"SIMILAR<br/>weight:0.7"| B
```''', '图4-2 Neo4j知识图谱Schema', '图4-2展示了图谱中的实体节点类型（Node Labels）和关系类型（Relationship Types），权重值用于推荐排序。所有关系采用无向设计（双向查询均可），Cypher查询时可从任意方向遍历。')

    add_heading_styled(doc, '五、接口表设计', 2)
    add_body_text(doc, '以下是模块间接口的数据交换约定表，定义了各模块间的输入输出格式。这些约定在代码中以Pydantic Schema的形式实现。')

    create_styled_table(doc,
        ['接口编号', '提供方', '消费方', 'HTTP方法', '路径', '数据类型'],
        [
            ['IF-01', '模块一', '模块三', 'GET', '/api/v1/user/profile', 'UserProfileForRecommend'],
            ['IF-02', '模块一', '模块四', 'JWT', 'Authorization Header', 'JWT Token Payload'],
            ['IF-03', '模块二', '模块三', 'POST', '/api/v1/graph/paths', 'GraphQueryRequest/Response'],
            ['IF-04', '模块三', '模块四/前端', 'GET', '/api/v1/recommend/home', 'RecommendListResponse'],
            ['IF-05', '模块三', '模块四/前端', 'GET', '/api/v1/recommend/similar/{id}', 'SimilarBooks'],
            ['IF-06', '模块四', '前端', 'GET', '/api/v1/ecosystem/trial/{id}', 'TrialReadResponse'],
            ['IF-07', '模块四', '前端', 'GET/POST', '/api/v1/ecosystem/comments', 'CommentResponse'],
            ['IF-08', '模块五', '前端', 'POST', '/api/v1/chat/message', 'ChatRequest/Response'],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第五部分：界面设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第五部分  界面设计', 1)

    add_heading_styled(doc, '一、前端技术栈概述', 2)
    add_body_text(doc, '前端采用uni-app跨平台框架，基于Vue 3 Composition API进行组件化开发。技术栈包括：uni-app作为主体框架（支持编译到Android/iOS/小程序/H5），Vue 3提供响应式数据绑定和组合式API，Pinia进行全局状态管理，vue-router实现页面路由导航，axios封装HTTP请求拦截器（自动附加JWT Token），Vite作为构建工具提供快速热更新。')

    add_heading_styled(doc, '二、页面框架结构设计', 2)
    add_body_text(doc, 'uni-app采用底部Tab导航结构，包含6个主要页面和2个功能页面：')

    add_code_block(doc, '''frontend-uni/
├── App.vue                      # 根组件（全局样式、生命周期）
├── main.js                      # 应用入口（挂载Pinia、全局组件）
├── pages.json                   # 页面路由配置（Tab栏定义）
├── manifest.json                # 应用配置（App权限、图标）
├── package.json                 # 依赖管理
├── pages/
│   ├── index/index.vue          # 首页（推荐流、轮播、热门）
│   ├── detail/detail.vue        # 图书详情页（试读、评论、收藏）
│   ├── shelf/shelf.vue          # 书架页（我的书架管理）
│   ├── chat/chat.vue            # AI问答助手页
│   ├── profile/profile.vue      # 个人中心（阅读统计、设置）
│   └── admin/admin.vue          # 管理后台（数据管理）
├── components/
│   ├── book-card.vue            # 图书卡片组件（复用）
│   ├── comment-list.vue         # 评论列表组件（复用）
│   └── chat-widget.vue          # 聊天悬浮窗组件（复用）
├── api/
│   └── index.js                 # API封装（统一请求、Token注入）
├── store/
│   └── index.js                 # Pinia状态管理（用户、书架）
└── utils/
    ├── auth.js                  # JWT Token管理（存储/读取/清除）
    └── request.js               # axios封装（拦截器、错误处理）''', '图5-1 前端项目目录结构')

    add_heading_styled(doc, '三、主要页面设计', 2)

    add_heading_styled(doc, '1、首页（index）', 3)
    add_body_text(doc, '首页为Tab栏默认页面，顶部展示搜索栏，下方为推荐流内容区。已登录用户展示混合推荐结果（含推荐理由），未登录用户展示热门排行。每条推荐卡片显示封面、书名、作者、评分，点击可进入图书详情页。支持下拉刷新和上拉加载更多。')

    add_heading_styled(doc, '2、图书详情页（detail）', 3)
    add_body_text(doc, '展示图书完整信息：封面、书名、作者、出版社、ISBN、评分、标签、简介。功能区提供试读按钮（权限分级）、收藏到书架（选择书架）、购买实体书（多平台链接）。评论区展示用户书评（置顶优先），支持发表评论和点赞。底部展示"相似图书"推荐卡片。')

    add_heading_styled(doc, '3、书架页（shelf）', 3)
    add_body_text(doc, '展示用户全部书架，默认包含"想读/在读/已读"三个书架，支持创建自定义书架。每个书架显示图书数量，点击进入书架内图书列表。支持长按图书进行移动/删除操作。列表项展示封面缩略图、书名、作者和添加时间。')

    add_heading_styled(doc, '4、AI问答助手页（chat）', 3)
    add_body_text(doc, '对话式界面，顶部显示欢迎语和能力介绍。消息列表展示用户与AI的多轮对话，支持Markdown渲染。输入框支持文字输入和发送，输入时显示发送按钮。AI回复末尾提供1-2个建议追问问题，可点击快速提问。支持清空对话历史。')

    add_heading_styled(doc, '5、个人中心页（profile）', 3)
    add_body_text(doc, '展示用户头像、用户名、阅读统计数据（已完成/在读书籍数、评分数量等）。阅读趋势图（最近7天柱状图）。热门标签词云展示。提供书架快捷入口、设置选项（修改密码、清除缓存）和退出登录按钮。')

    add_heading_styled(doc, '6、管理后台页（admin）', 3)
    add_body_text(doc, '管理员专属页面（需admin角色），提供图书管理（添加/编辑/删除）、用户管理（查看/禁用）、评论管理（删除/置顶）、购书链接配置、图谱初始化导入等功能面板。使用卡片式布局分类展示管理功能。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第六部分：单元模块设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第六部分  单元模块设计', 1)

    # ── 模块一 ──
    add_heading_styled(doc, '一、模块一：用户画像模块', 2)
    add_heading_styled(doc, '1、模块概述', 3)
    add_body_text(doc, '用户画像模块是系统的数据底座，负责用户全生命周期的数据管理。包含三大子模块：阅读行为采集（历史记录、搜索日志、收藏偏好、评分采集）、用户兴趣建模（标签偏好向量、作者偏好、类别偏好、动态更新）、阅读进度同步（多端同步、自动保存、upsert机制）。')

    add_heading_styled(doc, '2、模块类图', 3)
    add_mermaid_diagram(doc, '''```mermaid
classDiagram
    class User {
        +int id
        +str username
        +str email
        +str hashed_password
        +bool is_active
        +bool is_admin
    }

    class ReadingHistory {
        +int id
        +int user_id
        +int book_id
        +str status
        +datetime read_at
    }

    class Bookmark {
        +int id
        +int user_id
        +int book_id
        +str shelf_name
    }

    class ReadingProgress {
        +int id
        +int user_id
        +int book_id
        +float progress_percent
        +int current_page
    }

    class UserRating {
        +int id
        +int user_id
        +int book_id
        +float rating
    }

    class SearchLog {
        +int id
        +int user_id
        +str keyword
    }

    class UserService {
        +register_user(db, username, email, password) User
        +authenticate_user(db, username, password) str
        +record_reading_history(db, user_id, book_id, status)
        +record_search(db, user_id, keyword)
        +add_bookmark(db, user_id, book_id, shelf_name)
        +remove_bookmark(db, user_id, book_id)
        +rate_book(db, user_id, book_id, rating)
        +build_user_profile(db, user_id) dict
        +update_reading_progress(db, user_id, book_id, percent, page)
        +get_reading_progress(db, user_id) list
        +get_reading_stats(db, user_id) dict
        -_compute_tag_preferences(db, user_id) dict
        -_compute_favorite_authors(db, user_id) list
        -_compute_favorite_tags(db, user_id) list
        -_get_high_rated_books(db, user_id) list
    }

    class UserAPI {
        +POST /register
        +POST /login
        +GET /profile
        +PUT /profile
        +POST /history
        +GET /history
        +POST /bookmark
        +DELETE /bookmark/{id}
        +POST /rating
        +POST /progress
        +GET /progress
        +GET /stats
    }

    User "1" --> "*" ReadingHistory
    User "1" --> "*" Bookmark
    User "1" --> "*" ReadingProgress
    User "1" --> "*" UserRating
    User "1" --> "*" SearchLog
    UserAPI --> UserService
    UserService --> User
    UserService --> ReadingHistory
    UserService --> Bookmark
    UserService --> ReadingProgress
    UserService --> UserRating
```''', '图6-1 模块一：用户画像模块类图')

    add_heading_styled(doc, '3、用户注册登录活动图', 3)
    add_mermaid_diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 访问登录页
    访问登录页 --> 已有账号: 是
    访问登录页 --> 注册新用户: 否

    注册新用户 --> 填写注册信息
    填写注册信息 --> 验证表单
    验证表单 --> 用户名已存在: 重复
    用户名已存在 --> 填写注册信息
    验证表单 --> 提交注册: 通过
    提交注册 --> 密码bcrypt哈希
    密码bcrypt哈希 --> 存入MySQL
    存入MySQL --> 生成JWTToken
    生成JWTToken --> 登录成功

    已有账号 --> 输入用户名密码
    输入用户名密码 --> 提交登录
    提交登录 --> 查询用户表
    查询用户表 --> 用户不存在: 查无此人
    用户不存在 --> 输入用户名密码
    查询用户表 --> 验证密码: 找到用户
    验证密码 --> 密码错误: 不匹配
    密码错误 --> 输入用户名密码
    验证密码 --> 生成JWTToken: 匹配

    登录成功 --> 前端存储Token
    前端存储Token --> 跳转首页
    跳转首页 --> [*]
```''', '图6-2 用户注册登录活动图')

    add_heading_styled(doc, '4、类详细设计描述', 3)

    add_body_text(doc, '（1）UserService类：模块一的核心服务类，包含所有用户相关业务逻辑的静态方法。register_user()方法接收用户名、邮箱和密码，通过passlib的bcrypt算法对密码进行哈希后存储到MySQL的users表，返回新创建的User对象。authenticate_user()方法根据用户名查询用户记录，使用verify_password()验证密码哈希，验证通过后调用create_access_token()生成24小时有效的JWT Token。')
    add_body_text(doc, '（2）build_user_profile()方法：这是模块间接口契约的核心方法，为模块三推荐引擎提供用户偏好数据。实现逻辑为：从reading_history和bookmarks表统计用户阅读/收藏图书的标签频次→归一化为0~1的tag_weights字典（如{"科幻":0.85,"历史":0.15}）；从历史阅读记录按作者频次降序排列取top20作为favorite_author_ids；从user_ratings表筛选评分≥4.0的图书ID列表作为high_rated_book_ids。')
    add_body_text(doc, '（3）update_reading_progress()方法：支持多端阅读进度同步。使用(user_id, book_id)联合唯一索引实现upsert语义——如果该用户对该书的进度记录已存在则更新progress_percent和current_page，否则插入新记录。')

    # ── 模块二 ──
    add_heading_styled(doc, '二、模块二：知识图谱模块', 2)
    add_heading_styled(doc, '1、模块概述', 3)
    add_body_text(doc, '知识图谱模块负责构建和维护Neo4j中的书籍知识图谱，并为推荐引擎提供多跳路径推理能力。包含三大子模块：知识图谱构建（实体创建、关系建立、约束索引管理）、图谱查询与推理（五条Cypher路径查询、路径权重计算、候选图书生成）、图谱可视化（子图数据接口、统计信息）。')

    add_heading_styled(doc, '2、模块类图', 3)
    add_mermaid_diagram(doc, '''```mermaid
classDiagram
    class GraphService {
        -dict PATH_WEIGHTS
        +init_graph_constraints(session)
        +create_book_entity(session, book_id, title, props)
        +create_relation(session, src_type, src_id, rel, tgt_type, tgt_id)
        +find_paths(session, book_id, max_hops, top_k, weights) dict
        +get_subgraph(session, book_id, depth) dict
        +get_stats(session) dict
    }

    class GraphAPI {
        +POST /entity (创建实体)
        +POST /relation (创建关系)
        +POST /init (初始化约束)
        +POST /paths (图谱路径查询) ←核心
        +GET /subgraph/{book_id} (子图数据)
        +GET /stats (统计信息)
    }

    class Neo4jDriver {
        +session()
        +run(cypher, params)
    }

    GraphAPI --> GraphService
    GraphService --> Neo4jDriver
    Neo4jDriver --> Neo4j : "bolt://localhost:7687"
```''', '图6-3 模块二：知识图谱模块类图')

    add_heading_styled(doc, '3、图谱路径推理活动图', 3)
    add_mermaid_diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收图谱查询请求
    接收图谱查询请求 --> 解析参数: book_id, max_hops, top_k
    解析参数 --> 获取源图书信息

    获取源图书信息 --> 路径1_同作者: 执行Cypher查询
    获取源图书信息 --> 路径2_同标签: 执行Cypher查询
    获取源图书信息 --> 路径3_同系列: 执行Cypher查询
    获取源图书信息 --> 路径4_同出版社: 执行Cypher查询

    路径1_同作者 --> 检查max_hops: ≥2?
    路径2_同标签 --> 检查max_hops: ≥2?
    路径3_同系列 --> 检查max_hops: ≥2?
    路径4_同出版社 --> 检查max_hops: ≥2?

    检查max_hops --> 路径5_多跳推理: 是(执行Author→Book→Tag→Book)
    检查max_hops --> 合并所有路径结果: 否

    路径5_多跳推理 --> 合并所有路径结果

    合并所有路径结果 --> 去重并按final_score排序
    去重并按final_score排序 --> 取TopK候选
    取TopK候选 --> 返回GraphQueryResponse
    返回GraphQueryResponse --> [*]
```''', '图6-4 知识图谱路径推理活动图')

    add_heading_styled(doc, '4、类详细设计描述', 3)
    add_body_text(doc, '（1）GraphService类：模块二的核心服务类，所有方法均为静态方法。类属性PATH_WEIGHTS定义了五种关系类型的默认权重：AUTHORED=1.0（作者关系权重最高，因为同作者是强推荐信号），TAGGED=0.8（标签关系），SIMILAR=0.7（相似关系），SERIES_OF=0.6（系列关系），PUBLISHED=0.5（出版社关系较弱）。')
    add_body_text(doc, '（2）find_paths()方法：核心图谱推理方法。接收book_id（源图书ID）、max_hops（最大跳数，默认3）、top_k（返回候选项数，默认20）、path_weights（可覆盖的路径权重）四个参数。内部定义了五条参数化Cypher查询模板，分别沿AUTHORED→TAGGED→SERIES_OF→PUBLISHED→AUTHORED_TAG多跳路径探索候选图书。所有Cypher查询均使用参数化方式（$book_id, $limit），防止注入攻击。查询结果汇聚到candidates_map字典中按book_id合并，同一图书被多条路径命中时final_score累加，最后按得分降序排列截取top_k。')
    add_body_text(doc, '（3）五条Cypher查询路径说明：')
    add_body_text(doc, '路径1（同作者）：MATCH (source:Book)-[:AUTHORED]-(a:Author)-[:AUTHORED]-(candidate:Book)，权重1.0。从源图书出发，找到同作者的其他作品。')
    add_body_text(doc, '路径2（同标签）：MATCH (source:Book)-[:TAGGED]-(t:Tag)-[:TAGGED]-(candidate:Book)，权重0.8。共享标签越多得分越高。')
    add_body_text(doc, '路径3（同系列）：MATCH (source:Book)-[:SERIES_OF]-(s:Series)-[:SERIES_OF]-(candidate:Book)，权重0.6。')
    add_body_text(doc, '路径4（同出版社）：MATCH (source:Book)-[:PUBLISHED]-(p:Publisher)-[:PUBLISHED]-(candidate:Book)，权重0.5。')
    add_body_text(doc, '路径5（多跳）：MATCH (source:Book)-[:AUTHORED]-(a:Author)-[:AUTHORED]-(mid:Book)-[:TAGGED]-(t:Tag)-[:TAGGED]-(candidate:Book)，权重0.7。仅当max_hops≥2时执行。')

    # ── 模块三 ──
    add_heading_styled(doc, '三、模块三：个性化推荐模块', 2)
    add_heading_styled(doc, '1、模块概述', 3)
    add_body_text(doc, '个性化推荐模块是系统的核心算法模块，负责融合多源信号生成个性化推荐。包含五大子模块：ItemCF协同过滤推荐（构建用户-图书评分矩阵、余弦相似度计算）、知识图谱推理推荐（调用模块二接口合并多路径结果）、混合推荐策略（四策略加权融合，权重可配置）、推荐场景分发（首页/详情页/热门/新书/策略切换）、推荐理由生成（基于路径类型的自然语言解释）。')

    add_heading_styled(doc, '2、模块类图', 3)
    add_mermaid_diagram(doc, '''```mermaid
classDiagram
    class RecommendService {
        -UserProfileForRecommend user_profile
        -GraphPathsForRecommend graph_paths
        +recommend_cf(top_n) list
        +recommend_kg(top_n) list
        +recommend_hot(top_n) list
        +recommend_new(top_n) list
        +recommend_hybrid(top_n, weights) list
        +generate_reason(book_title, path_info) str$
    }

    class UserProfileForRecommend {
        +int user_id
        +dict tag_weights
        +list favorite_author_ids
        +list favorite_tag_ids
        +list high_rated_book_ids
    }

    class GraphPathsForRecommend {
        +int source_book_id
        +list candidates
    }

    class RecommendWeights {
        +float kg_weight = 0.4
        +float cf_weight = 0.4
        +float hot_weight = 0.1
        +float new_weight = 0.1
    }

    class RecommendAPI {
        +GET /home (首页推荐)
        +GET /similar/{book_id} (相似图书)
        +GET /hot (热门推荐)
        +PUT /weights (调整权重)
    }

    RecommendAPI --> RecommendService
    RecommendService --> UserProfileForRecommend
    RecommendService --> GraphPathsForRecommend
    RecommendService --> RecommendWeights
```''', '图6-5 模块三：个性化推荐模块类图')

    add_heading_styled(doc, '3、混合推荐算法活动图', 3)
    add_mermaid_diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收推荐请求
    接收推荐请求 --> 检查登录状态

    检查登录状态 --> 热门推荐: 未登录(冷启动)
    热门推荐 --> 返回推荐列表

    检查登录状态 --> 获取用户画像: 已登录
    获取用户画像 --> 并行执行四种策略

    state 并行执行四种策略 {
        策略1_ItemCF --> CF结果列表
        策略2_图谱推理 --> KG结果列表
        策略3_热门推荐 --> Hot结果列表
        策略4_新书推荐 --> New结果列表
    }

    CF结果列表 --> 加权融合
    KG结果列表 --> 加权融合
    Hot结果列表 --> 加权融合
    New结果列表 --> 加权融合

    加权融合 --> 合并去重: final_score = kg×0.4 + cf×0.4 + hot×0.1 + new×0.1
    合并去重 --> 按score降序排列
    按score降序排列 --> 取TopN
    取TopN --> 生成推荐理由
    生成推荐理由 --> 返回推荐列表
    返回推荐列表 --> [*]
```''', '图6-6 混合推荐算法活动图')

    add_heading_styled(doc, '4、类详细设计描述', 3)
    add_body_text(doc, '（1）RecommendService类：模块三的核心服务类。构造函数接收UserProfileForRecommend（模块一提供的用户画像）和可选的GraphPathsForRecommend（模块二提供的图谱推理结果）。内部实现了四种推荐策略方法：recommend_cf()基于物品的协同过滤、recommend_kg()基于知识图谱的推理推荐、recommend_hot()基于热度分的推荐（冷启动保底）、recommend_new()基于新书上架时间的推荐（时效性）。')
    add_body_text(doc, '（2）recommend_hybrid()方法：混合推荐的核心入口。使用RecommendWeights配置（默认kg_weight=0.4, cf_weight=0.4, hot_weight=0.1, new_weight=0.1）进行加权融合。融合算法：为每本候选图书维护一个merged字典（key=book_id），同一本书被多个策略命中时分数累加（乘以各自策略权重），最终按累加分数降序排列返回top_n。')
    add_body_text(doc, '（3）generate_reason()静态方法：根据path_type生成自然语言推荐理由。模板映射表包含8种理由类型——author→"因为你也喜欢{作者}的作品"，tag→"这本书和{标签}相关"，series→"这本书与{系列}属于同一丛书"，publisher→"由{出版社}出版"，cf→"与你读过的高分图书相似"，hot→"近期热门图书"，new→"新书上架"。')

    # ── 模块四 ──
    add_heading_styled(doc, '四、模块四：阅读生态模块', 2)
    add_heading_styled(doc, '1、模块概述', 3)
    add_body_text(doc, '阅读生态模块是系统的用户交互核心，负责面向用户的完整阅读功能。包含五大子模块：电子书在线试读（权限分级控制、进度自动保存）、书评社区（发表评论、点赞/取消、管理员置顶、删除）、实体书购书链接（京东/当当/淘宝多平台比价）、书架与收藏管理（自定义书架、图书移动、删除）、阅读统计（阅读时长、完成数量、趋势图表数据）。')

    add_heading_styled(doc, '2、模块类图', 3)
    add_mermaid_diagram(doc, '''```mermaid
classDiagram
    class EcosystemService {
        +get_trial_info(db, book_id, user_id) dict
        +get_trial_content(db, book_id, user_id) dict
        +create_comment(db, user_id, book_id, content) BookComment
        +get_book_comments(db, book_id, page, size) list
        +like_comment(db, user_id, comment_id) bool
        +pin_comment(db, comment_id, is_pinned)
        +delete_comment(db, comment_id, user_id, is_admin)
        +update_purchase_links(db, book_id, urls)
        +get_purchase_links(db, book_id) dict
        +get_user_bookshelves(db, user_id) list
        +get_shelf_books(db, user_id, shelf_name) list
        +move_book_to_shelf(db, user_id, book_id, shelf)
        +remove_book_from_shelf(db, user_id, book_id)
        +delete_shelf(db, user_id, shelf_name)
        +get_reading_stats(db, user_id) dict
        -_update_comment_likes(db, comment_id)
    }

    class BookComment {
        +int id
        +int user_id
        +int book_id
        +text content
        +int likes_count
        +bool is_pinned
    }

    class CommentLike {
        +int id
        +int user_id
        +int comment_id
    }

    class EcosystemAPI {
        +GET /trial/{book_id}
        +GET /trial/{book_id}/content
        +GET /comments/{book_id}
        +POST /comments
        +POST /comments/like
        +PUT /comments/{id}/pin
        +DELETE /comments/{id}
        +GET /purchase/{book_id}
        +PUT /purchase
        +GET /shelves
        +GET /shelves/{name}
        +PUT /shelves/move
        +DELETE /shelves/book
        +DELETE /shelves/{name}
    }

    EcosystemAPI --> EcosystemService
    EcosystemService --> BookComment
    EcosystemService --> CommentLike
```''', '图6-7 模块四：阅读生态模块类图')

    add_heading_styled(doc, '3、书评点赞流程时序图', 3)
    add_mermaid_diagram(doc, '''```mermaid
sequenceDiagram
    participant U as 用户
    participant F as 前端
    participant API as 生态API
    participant SVC as EcosystemService
    participant DB as MySQL

    U->>F: 点击评论点赞按钮
    F->>API: POST /api/v1/ecosystem/comments/like
    API->>API: 验证JWT Token
    API->>SVC: like_comment(db, user_id, comment_id)
    SVC->>DB: 查询comment_likes表(user_id+comment_id)
    DB-->>SVC: 返回查询结果

    alt 未点赞
        SVC->>DB: INSERT comment_likes
        SVC->>SVC: _update_comment_likes()
        SVC->>DB: UPDATE book_comments SET likes_count=count
        SVC-->>API: 返回 True(已点赞)
        API-->>F: 点赞成功
        F-->>U: 图标变红/数字+1
    else 已点赞
        SVC->>DB: DELETE comment_likes
        SVC->>SVC: _update_comment_likes()
        SVC->>DB: UPDATE book_comments SET likes_count=count
        SVC-->>API: 返回 False(已取消)
        API-->>F: 取消成功
        F-->>U: 图标恢复/数字-1
    end
```''', '图6-8 书评点赞Toggle时序图')

    add_heading_styled(doc, '4、类详细设计描述', 3)
    add_body_text(doc, '（1）试读服务：get_trial_info()根据user_id是否为None判断登录状态，返回对应的试读页数限制（未登录3页/登录10页，由Settings中的TRIAL_PAGES_ANONYMOUS和TRIAL_PAGES_LOGGED_IN配置）。同时查询reading_progress表获取用户当前阅读进度。get_trial_content()返回试读具体内容，当前使用图书简介作为试读文本（生产环境可对接PDF.js/EPUB解析器），内容长度根据用户登录状态限制（登录2000字符/未登录600字符）。')
    add_body_text(doc, '（2）书评社区：create_comment()创建新评论记录。get_book_comments()按置顶优先→点赞数降序→时间降序排序，支持分页。like_comment()实现Toggle模式——如果用户已点赞则取消点赞（DELETE），未点赞则添加点赞（INSERT），操作后通过_update_comment_likes()更新评论的likes_count聚合字段。delete_comment()支持权限控制：作者本人或管理员可删除。')
    add_body_text(doc, '（3）书架管理：使用三个默认书架（想读/在读/已读）和自定义书架。get_user_bookshelves()从bookmarks表按shelf_name分组统计数量，确保三个默认书架始终出现在列表中。move_book_to_shelf()修改bookmark记录的shelf_name字段。delete_shelf()删除自定义书架及其下所有收藏（默认书架不可删除，返回-1）。')

    # ── 模块五 ──
    add_heading_styled(doc, '五、模块五：智能问答助手模块', 2)
    add_heading_styled(doc, '1、模块概述', 3)
    add_body_text(doc, '智能问答助手模块是基于大语言模型的对话式AI助手，为用户提供自然语言交互入口。包含五大核心能力：意图识别（LLM对用户问题进行六类分类）、知识检索增强（根据意图从MySQL检索业务上下文数据）、边界控制（拒绝超出荐书业务范围的问题）、LLM对话生成（结合System Prompt+上下文+历史生成回答）、对话历史管理（存储/检索/删除多轮对话记录）。')

    add_heading_styled(doc, '2、模块类图', 3)
    add_mermaid_diagram(doc, '''```mermaid
classDiagram
    class AIChatService {
        +process_message(db, user, message) dict
        +save_message(db, user_id, role, content, intent) ChatHistory
        +get_chat_history(db, user_id, limit) list
        +delete_chat_history(db, user_id) int
        -_classify_intent(message, user_role) dict
        -_fallback_intent(message) dict
        -_build_context(intent, user, db, entities) str
        -_check_bounds(message, intent) str
        -_generate_answer(msg, intent, ctx, user, db, history) dict
        -_fallback_answer(intent, ctx) dict
        -_generate_suggestions(intent) list
    }

    class ChatHistory {
        +int id
        +int user_id
        +enum role
        +text content
        +str intent_type
        +datetime created_at
    }

    class ChatAPI {
        +POST /message (发送消息)
        +GET /history (获取历史)
        +DELETE /history (清空历史)
    }

    ChatAPI --> AIChatService
    AIChatService --> ChatHistory
    AIChatService --> LLMClient : "OpenAI兼容API"
```''', '图6-9 模块五：智能问答助手模块类图')

    add_heading_styled(doc, '3、对话处理活动图', 3)
    add_mermaid_diagram(doc, '''```mermaid
stateDiagram-v2
    [*] --> 接收用户消息
    接收用户消息 --> 确定用户角色

    确定用户角色 --> 意图识别: 调用LLM或关键词降级

    意图识别 --> 超范围拒绝: out_of_scope
    超范围拒绝 --> 返回礼貌拒绝回复
    返回礼貌拒绝回复 --> 保存对话历史
    保存对话历史 --> [*]

    意图识别 --> 合法意图: 六种业务意图
    合法意图 --> 关键词边界检查

    关键词边界检查 --> 超范围拒绝: 含禁用词
    关键词边界检查 --> 检索业务上下文: 通过

    检索业务上下文 --> 获取对话历史: 从MySQL查询最近30条
    获取对话历史 --> 构建LLM消息: System+历史+上下文+问题
    构建LLM消息 --> 调用LLM生成回答

    调用LLM生成回答 --> 降级回复: LLM不可用
    调用LLM生成回答 --> AI回复: 成功
    降级回复 --> 保存对话历史
    AI回复 --> 生成建议追问
    生成建议追问 --> 保存对话历史
    保存对话历史 --> [*]
```''', '图6-10 智能问答对话处理活动图')

    add_heading_styled(doc, '4、类详细设计描述', 3)
    add_body_text(doc, '（1）process_message()方法：模块五的核心入口。处理流程为：确定用户角色（anonymous/user/admin）→ 意图识别（_classify_intent调用LLM或降级为关键词匹配）→ 边界检查（_check_bounds检查out_of_scope或含禁用关键词）→ 构建上下文（_build_context根据意图从MySQL检索图书/用户/功能数据）→ 获取对话历史（get_chat_history最近30条消息）→ LLM生成回答（_generate_answer或降级_fallback_answer）→ 保存对话记录（user消息和assistant消息均保存）。')
    add_body_text(doc, '（2）意图识别体系：定义六种业务意图——function_qa（功能问答）、book_rec（自然语言荐书）、book_qa（图书信息查询）、personal_qa（个人阅读查询）、admin_help（管理员帮助）、kg_assist（知识图谱辅助），以及out_of_scope（超范围拒绝）。LLM可用时通过INTENT_PROMPT让LLM进行JSON格式输出，LLM不可用时降级为Python关键词匹配（如含"推荐/介绍"判定为book_rec，含"我/我的"判定为personal_qa）。')
    add_body_text(doc, '（3）System Prompt设计：定义AI助手的角色边界——可以回答功能问答、自然语言荐书、图书知识问答、个人阅读问答、管理员帮助、知识图谱辅助六类问题；不可以回答超出业务范围的问题（天气、新闻、编程等）。回复风格要求简洁、友好、有引导性，回复末尾建议1-2个相关问题。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第七部分：数据访问层设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第七部分  数据访问层设计', 1)

    add_heading_styled(doc, '一、MySQL数据访问层', 2)
    add_body_text(doc, 'MySQL数据访问采用SQLAlchemy ORM框架。数据库连接通过app/core/database.py中的SessionLocal工厂函数管理，采用连接池（pool_size=20, max_overflow=40）和连接预检（pool_pre_ping=True）保证高可用。所有数据库会话通过FastAPI的Depends(get_db)依赖注入获取，自动管理事务生命周期（yield后自动close）。')

    add_heading_styled(doc, '1、Session管理', 3)
    add_code_block(doc, '''# app/core/database.py - 核心数据库连接管理
engine = create_engine(
    settings.DATABASE_URL,
    pool_size=20,           # 常驻连接池大小
    max_overflow=40,        # 最大溢出连接数
    pool_pre_ping=True,     # 连接前检测可用性
)

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

def get_db():
    """FastAPI依赖注入：获取MySQL会话，请求结束自动关闭"""
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()''', '代码7-1 MySQL会话管理')

    add_heading_styled(doc, '2、ORM模型设计', 3)
    add_body_text(doc, '所有ORM模型位于app/models/目录，使用SQLAlchemy的declarative_base()作为基类。模型按模块组织：models/user.py（模块一：6张表——User/ReadingHistory/SearchLog/Bookmark/ReadingProgress/UserRating）、models/book.py（模块二：5张表+2张中间表——Book/Author/Publisher/Tag/Series + book_author/book_tag）、models/ecosystem.py（模块四：2张表——BookComment/CommentLike）、models/chat.py（模块五：1张表——ChatHistory）。')
    add_body_text(doc, '模型间关系通过ForeignKey和relationship声明。例如User模型通过relationship关联了reading_history、bookmarks、reading_progress、ratings、comments、chat_history六个子表。Book模型通过secondary参数关联book_author和book_tag中间表，实现与Author和Tag的多对多关系。')

    add_heading_styled(doc, '二、Neo4j数据访问层', 2)
    add_body_text(doc, 'Neo4j数据访问通过官方Python Driver实现。驱动实例通过单例模式（get_neo4j_driver()）管理，使用Bolt协议连接（bolt://localhost:7687）。每次API请求通过Depends(get_neo4j_session)获取新的session对象，请求结束后自动关闭。所有Cypher查询使用参数化方式（$book_id, $limit），防止Cypher注入攻击。')

    add_code_block(doc, '''# app/core/database.py - Neo4j连接管理
_neo4j_driver = None  # 全局单例

def get_neo4j_driver():
    """获取Neo4j driver（全局单例，应用生命周期内复用）"""
    global _neo4j_driver
    if _neo4j_driver is None:
        _neo4j_driver = GraphDatabase.driver(
            settings.NEO4J_URI,
            auth=(settings.NEO4J_USER, settings.NEO4J_PASSWORD),
        )
    return _neo4j_driver

def get_neo4j_session():
    """FastAPI依赖注入：获取Neo4j会话，请求结束自动关闭"""
    driver = get_neo4j_driver()
    with driver.session() as session:
        yield session''', '代码7-2 Neo4j连接管理')

    add_heading_styled(doc, '三、Redis数据访问层', 2)
    add_body_text(doc, 'Redis客户端通过单例模式管理，使用redis-py库。连接参数通过Settings配置（REDIS_HOST, REDIS_PORT），启用decode_responses=True自动将字节解码为字符串。Redis用于缓存热门推荐结果（减少MySQL查询压力）、用户Session数据（加速认证）和ItemCF相似度矩阵（避免重复计算）。')

    add_code_block(doc, '''# app/core/database.py - Redis连接管理
_redis_client = None

def get_redis():
    """获取Redis客户端（全局单例）"""
    global _redis_client
    if _redis_client is None:
        _redis_client = _redis.Redis(
            host=settings.REDIS_HOST,
            port=settings.REDIS_PORT,
            decode_responses=True,
        )
    return _redis_client''', '代码7-3 Redis连接管理')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第八部分：功能模块接口设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第八部分  功能模块接口设计', 1)

    add_heading_styled(doc, '一、模块间接口契约', 2)
    add_body_text(doc, '模块间接口契约是各模块独立开发、并行推进的基础。以下四个核心契约已在代码中以Pydantic Schema形式实现，修改前必须与上下游成员沟通。')

    add_heading_styled(doc, '契约1：模块一 → 模块三（用户画像）', 3)
    add_code_block(doc, '''response = {
    "user_id": 1,
    "tag_weights": {"科幻": 0.85, "人工智能": 0.42, "历史": 0.15},
    "favorite_author_ids": [1, 5, 23],
    "favorite_tag_ids": [1, 2, 7],
    "high_rated_book_ids": [101, 203, 405]
}
# 方法: GET /api/v1/user/profile
# Schema: UserProfileForRecommend (app/schemas/recommend.py)
# 负责人: A  |  消费者: C''', '接口契约 IF-01')

    add_heading_styled(doc, '契约2：模块二 → 模块三（图谱推理）', 3)
    add_code_block(doc, '''request = {
    "book_id": 101, "max_hops": 3, "top_k": 20,
    "author_weight": 1.0, "tag_weight": 0.8,
    "publisher_weight": 0.5, "series_weight": 0.6
}

response = {
    "source_book_id": 101, "source_book_title": "三体",
    "candidates": [{
        "book_id": 104, "book_title": "流浪地球",
        "paths": [{"path_type": "author", "via": "刘慈欣",
                   "hop_count": 1, "weight": 1.0}],
        "final_score": 0.9
    }]
}
# 方法: POST /api/v1/graph/paths
# Schema: GraphQueryRequest/Response (app/schemas/book.py)
# 负责人: B  |  消费者: C''', '接口契约 IF-03')

    add_heading_styled(doc, '契约3：模块三 → 前端（推荐结果）', 3)
    add_code_block(doc, '''response = {
    "user_id": 1, "strategy": "hybrid",
    "items": [{
        "book": {"id": 104, "title": "流浪地球",
                 "authors": ["刘慈欣"], "cover_url": "...",
                 "avg_rating": 4.5, "tags": ["科幻"]},
        "score": 0.92,
        "reason": "因为你也喜欢刘慈欣的作品",
        "reason_type": "author"
    }],
    "generated_at": "2026-06-25T10:30:00"
}
# 方法: GET /api/v1/recommend/home
# Schema: RecommendListResponse (app/schemas/recommend.py)
# 负责人: C  |  消费者: D (前端展示)''', '接口契约 IF-04')

    add_heading_styled(doc, '二、RESTful API完整清单', 2)
    add_body_text(doc, '以下是系统全部API端点清单，按模块分类。所有需要认证的接口在Header中携带 Authorization: Bearer <jwt_token>。')

    # 模块一API
    add_heading_styled(doc, '1、模块一 API（/api/v1/user）', 3)
    create_styled_table(doc,
        ['方法', '路径', '认证', '功能', '请求/响应Schema'],
        [
            ['POST', '/register', '否', '用户注册', 'UserRegister → TokenResponse'],
            ['POST', '/login', '否', '用户登录', 'UserLogin → TokenResponse'],
            ['GET', '/profile', '是', '获取用户画像（模块三调用）', '→ UserProfile'],
            ['PUT', '/profile', '是', '更新用户偏好', 'UserProfileUpdate → OK'],
            ['POST', '/history', '是', '记录阅读历史', 'ReadingHistoryCreate → OK'],
            ['GET', '/history', '是', '获取阅读历史', '→ list[ReadingHistoryResponse]'],
            ['POST', '/bookmark', '是', '添加收藏', 'BookmarkCreate → OK'],
            ['DELETE', '/bookmark/{id}', '是', '取消收藏', '→ OK'],
            ['GET', '/bookmarks', '是', '获取收藏列表', '→ list[BookmarkResponse]'],
            ['POST', '/rating', '是', '图书评分(0.5~5.0)', 'RatingCreate → OK'],
            ['POST', '/progress', '是', '更新阅读进度', 'ReadingProgressUpdate → OK'],
            ['GET', '/progress', '是', '获取全部进度', '→ list[ReadingProgressResponse]'],
            ['GET', '/stats', '是', '阅读统计', '→ ReadingStats'],
        ]
    )

    # 模块二API
    add_heading_styled(doc, '2、模块二 API（/api/v1/graph）', 3)
    create_styled_table(doc,
        ['方法', '路径', '认证', '功能', '请求/响应Schema'],
        [
            ['POST', '/entity', '是(管理)', '创建图谱实体', 'GraphEntityCreate → OK'],
            ['POST', '/relation', '是(管理)', '创建图谱关系', 'GraphRelationCreate → OK'],
            ['POST', '/init', '是(管理)', '初始化约束索引', '→ OK'],
            ['POST', '/paths', '否', '图谱路径查询（模块三调用）', 'GraphQueryRequest → Response'],
            ['GET', '/subgraph/{id}', '否', '获取子图可视化数据', '→ {nodes, edges}'],
            ['GET', '/stats', '否', '图谱统计', '→ {books, authors, tags, relations}'],
        ]
    )

    # 模块三API
    add_heading_styled(doc, '3、模块三 API（/api/v1/recommend）', 3)
    create_styled_table(doc,
        ['方法', '路径', '认证', '功能', '请求/响应Schema'],
        [
            ['GET', '/home', '可选', '首页推荐（核心接口）', '→ RecommendListResponse'],
            ['GET', '/similar/{id}', '否', '相似图书推荐', '→ {similar_books}'],
            ['GET', '/hot', '否', '热门推荐', '→ {hot_books}'],
            ['PUT', '/weights', '是', '调整混合推荐权重', 'RecommendWeights → OK'],
        ]
    )

    # 模块四API
    add_heading_styled(doc, '4、模块四 API（/api/v1/ecosystem）', 3)
    create_styled_table(doc,
        ['方法', '路径', '认证', '功能', '请求/响应Schema'],
        [
            ['GET', '/trial/{id}', '可选', '试读权限信息', '→ TrialReadResponse'],
            ['GET', '/trial/{id}/content', '可选', '试读内容', '→ TrialContentResponse'],
            ['GET', '/comments/{id}', '否', '图书评论列表', '→ list[CommentResponse]'],
            ['POST', '/comments', '是', '发表书评', 'CommentCreate → OK'],
            ['POST', '/comments/like', '是', '点赞/取消评论', 'CommentLikeAction → OK'],
            ['PUT', '/comments/{id}/pin', '是(管理)', '置顶评论', '→ OK'],
            ['DELETE', '/comments/{id}', '是', '删除评论', '→ OK'],
            ['GET', '/purchase/{id}', '否', '购书链接', '→ PurchaseLinkResponse'],
            ['PUT', '/purchase', '是(管理)', '配置购书链接', 'PurchaseLinkUpdate → OK'],
            ['GET', '/shelves', '是', '书架列表', '→ list[ShelfResponse]'],
            ['GET', '/shelves/{name}', '是', '书架图书', '→ list[ShelfBookResponse]'],
            ['PUT', '/shelves/move', '是', '移动图书', 'MoveBookRequest → OK'],
            ['DELETE', '/shelves/book', '是', '移除图书', 'RemoveBookRequest → OK'],
            ['DELETE', '/shelves/{name}', '是', '删除书架', '→ OK'],
        ]
    )

    # 模块五API
    add_heading_styled(doc, '5、模块五 API（/api/v1/chat）', 3)
    create_styled_table(doc,
        ['方法', '路径', '认证', '功能', '请求/响应Schema'],
        [
            ['POST', '/message', '可选', '发送对话消息', 'ChatRequest → ChatResponse'],
            ['GET', '/history', '是', '获取对话历史', '→ ChatHistoryResponse'],
            ['DELETE', '/history', '是', '清空对话历史', '→ ChatDeleteResponse'],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第九部分：算法设计
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第九部分  算法设计', 1)

    add_heading_styled(doc, '一、ItemCF协同过滤算法', 2)
    add_heading_styled(doc, '1、算法原理', 3)
    add_body_text(doc, 'ItemCF（Item-based Collaborative Filtering）基于物品的协同过滤算法。核心思想：如果两个图书被相似的用户群体评分，则认为它们相似。计算图书间的余弦相似度矩阵，从用户已评分的高分图书出发，推荐Top-N最相似的未知图书。')

    add_heading_styled(doc, '2、算法步骤', 3)
    add_body_text(doc, '步骤1：从MySQL user_ratings表构建用户-图书评分矩阵（rows=users, cols=books, values=ratings）。使用scipy.sparse.csr_matrix存储为稀疏矩阵（大多数用户只评分了少量图书）。')
    add_body_text(doc, '步骤2：计算图书间的余弦相似度。使用sklearn.metrics.pairwise.cosine_similarity计算矩阵转置后的列间余弦相似度：cos(A, B) = (A·B) / (|A|·|B|)。过滤评分人数<10的冷门图书以提升计算效率和推荐质量。')
    add_body_text(doc, '步骤3：从用户画像的high_rated_book_ids（评分≥4.0的图书）出发，对每本高分图书找到Top-N最相似的图书。多本高分书命中同一本书时，累加（相似度×评分）作为最终推荐分数。排除用户已读/已评分的图书。')
    add_body_text(doc, '步骤4：优化策略——相似度矩阵离线预计算，结果存入Redis缓存（设置TTL定期刷新），避免每次请求重新计算。')

    add_heading_styled(doc, '二、知识图谱路径推理算法', 2)
    add_heading_styled(doc, '1、算法原理', 3)
    add_body_text(doc, '知识图谱推荐基于图遍历的候选发现策略。从用户感兴趣的高分图书（种子节点）出发，沿图谱中的多种关系路径进行多跳探索，发现通过语义关系关联的候选图书。每条关系路径赋予预定义权重，同一候选图书被多条路径命中时权重累加。')

    add_heading_styled(doc, '2、五条推理路径', 3)
    create_styled_table(doc,
        ['路径编号', '路径名称', 'Cypher模式', '跳数', '权重', '语义含义'],
        [
            ['P1', '同作者路径', '(Book)-[:AUTHORED]-(Author)-[:AUTHORED]-(Book)', '1', '1.0', '同作者的其他作品，强推荐信号'],
            ['P2', '同标签路径', '(Book)-[:TAGGED]-(Tag)-[:TAGGED]-(Book)', '1', '0.8', '共享标签越多越相关'],
            ['P3', '同系列路径', '(Book)-[:SERIES_OF]-(Series)-[:SERIES_OF]-(Book)', '1', '0.6', '同一丛书系列'],
            ['P4', '同出版社路径', '(Book)-[:PUBLISHED]-(Publisher)-[:PUBLISHED]-(Book)', '1', '0.5', '同出版社，信号较弱'],
            ['P5', '多跳推理', '(Book)-Author-(Book)-Tag-(Book)', '≥2', '0.7', '综合推理，发现间接关联'],
        ]
    )

    add_heading_styled(doc, '3、路径权重计算', 3)
    add_body_text(doc, '每个候选图书的final_score = Σ(命中路径的权重)。例如：《流浪地球》同时被P1（同作者刘慈欣，权重1.0）和P2（同标签科幻，权重0.8）命中，则final_score = 1.0 + 0.8 = 1.8。权重可通过API参数path_weights运行时覆盖（如管理员调整推荐策略时降低出版社权重）。')

    add_heading_styled(doc, '三、混合推荐融合算法', 2)
    add_heading_styled(doc, '1、加权融合公式', 3)
    add_code_block(doc, '''final_score = kg_score × W_kg + cf_score × W_cf + hot_score × W_hot + new_score × W_new

默认权重配置:
  W_kg  = 0.4  （知识图谱推理 — 可解释性强）
  W_cf  = 0.4  （协同过滤 — 精度高）
  W_hot = 0.1  （热门推荐 — 冷启动保底）
  W_new = 0.1  （新书推荐 — 时效性）

总权重 = 1.0（可根据业务需求运行时调整）''', '公式9-1 混合推荐加权融合公式')

    add_heading_styled(doc, '2、融合流程', 3)
    add_body_text(doc, '（1）四种推荐策略独立计算各自的候选列表（各自返回score在0~1之间的候选图书）')
    add_body_text(doc, '（2）创建merged字典（key=book_id，value={book_title, accumulated_score, reasons[]}）')
    add_body_text(doc, '（3）遍历四种策略结果，对每本候选图书：如果不在merged中则初始化条目，将策略得分×策略权重累加到accumulated_score，将路径信息追加到reasons[]')
    add_body_text(doc, '（4）对merged中的所有值按accumulated_score降序排列，截取top_n返回')
    add_body_text(doc, '（5）为每条推荐调用generate_reason()生成自然语言推荐理由')

    add_heading_styled(doc, '四、意图识别算法（模块五）', 2)
    add_body_text(doc, '意图识别采用LLM优先+关键词降级的双层策略：')
    add_body_text(doc, '（1）LLM识别（优先）：当LLM API可用时，将用户消息和INTENT_PROMPT发送给LLM，要求输出JSON格式的意图分类结果 {"intent": "xxx", "entities": [...], "confidence": 0.0-1.0}。使用temperature=0.1确保分类结果稳定。')
    add_body_text(doc, '（2）关键词降级（备选）：当LLM不可用时，使用Python关键词匹配进行意图分类。规则包括：含"推荐/介绍/好看/适合"→book_rec；含"作者/出版社/标签/ISBN"→book_qa；含"我/我的/收藏/书架/历史"→personal_qa；含"怎么/如何/在哪/功能/使用"→function_qa；含"添加/删除/管理/后台"→admin_help。降级方案置信度为0.5。')

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 第十部分：项目文件目录结构
    # ═══════════════════════════════════════
    add_heading_styled(doc, '第十部分  项目文件目录结构', 1)
    add_body_text(doc, '项目采用前后端分离的目录结构，后端基于FastAPI按模块组织代码，前端采用uni-app跨平台框架。以下展示完整的项目目录及文件职责说明。')

    add_code_block(doc, '''PythonProject6/
├── app/                                    # FastAPI后端主目录
│   ├── __init__.py
│   ├── main.py                             # 应用主入口（FastAPI实例、CORS、路由注册）
│   ├── core/                               # 核心基础设施
│   │   ├── __init__.py
│   │   ├── config.py                       # Settings配置中心（环境变量加载）
│   │   ├── database.py                     # 数据库连接管理（MySQL/Neo4j/Redis）
│   │   └── security.py                     # JWT生成/验证、bcrypt密码哈希
│   │
│   ├── models/                             # SQLAlchemy ORM数据模型
│   │   ├── __init__.py
│   │   ├── user.py                         # 模块一：User/ReadingHistory/Bookmark/Progress/Rating/SearchLog
│   │   ├── book.py                         # 模块二：Book/Author/Publisher/Tag/Series + 中间表
│   │   ├── ecosystem.py                    # 模块四：BookComment/CommentLike
│   │   └── chat.py                         # 模块五：ChatHistory
│   │
│   ├── schemas/                            # Pydantic请求/响应验证模型
│   │   ├── __init__.py
│   │   ├── user.py                         # 模块一：注册/登录/画像/行为/进度Schema
│   │   ├── book.py                         # 模块二：图书详情/图谱查询/实体创建Schema
│   │   ├── recommend.py                    # 模块三：推荐请求/响应/权重/接口契约Schema
│   │   ├── ecosystem.py                    # 模块四：评论/试读/购书/书架/统计Schema
│   │   └── chat.py                         # 模块五：对话请求/响应/历史Schema
│   │
│   ├── services/                           # 业务逻辑层（核心）
│   │   ├── __init__.py
│   │   ├── user_service.py                 # 模块一：认证/行为采集/画像建模/进度同步
│   │   ├── graph_service.py                # 模块二：图谱构建/路径推理/可视化
│   │   ├── recommend_service.py            # 模块三：CF推荐/KG推荐/混合融合/理由生成
│   │   ├── ecosystem_service.py            # 模块四：试读/书评/购书/书架/统计
│   │   └── ai_chat_service.py              # 模块五：意图识别/上下文检索/LLM对话
│   │
│   └── api/                                # API层（HTTP端点）
│       ├── __init__.py
│       ├── deps.py                         # 依赖注入（get_current_user/admin/optional）
│       └── v1/
│           ├── __init__.py
│           ├── router.py                   # API v1总路由（注册五大模块路由）
│           └── endpoints/
│               ├── __init__.py
│               ├── user.py                 # 模块一：13个端点
│               ├── graph.py                # 模块二：6个端点
│               ├── recommend.py            # 模块三：4个端点
│               ├── ecosystem.py            # 模块四：15个端点
│               └── ai_chat.py              # 模块五：3个端点
│
├── frontend-uni/                           # uni-app前端（跨平台）
│   ├── App.vue                             # 根组件
│   ├── main.js                             # 应用入口（挂载Pinia、全局组件注册）
│   ├── pages.json                          # 页面路由配置（Tab栏定义）
│   ├── manifest.json                       # 应用配置（App权限、图标、版本）
│   ├── package.json                        # 依赖管理（Vue3/Pinia/axios等）
│   ├── pages/                              # 页面组件
│   │   ├── index/index.vue                 # 首页（推荐流）
│   │   ├── detail/detail.vue               # 图书详情页
│   │   ├── shelf/shelf.vue                 # 书架管理页
│   │   ├── chat/chat.vue                   # AI问答助手页
│   │   ├── profile/profile.vue             # 个人中心页
│   │   └── admin/admin.vue                 # 管理后台页
│   ├── components/                         # 复用组件
│   │   ├── book-card.vue                   # 图书卡片组件
│   │   ├── comment-list.vue                # 评论列表组件
│   │   └── chat-widget.vue                 # 聊天悬浮窗组件
│   ├── api/index.js                        # API封装（统一请求拦截、Token注入）
│   ├── store/index.js                      # Pinia状态管理（用户信息、书架数据）
│   └── utils/                              # 工具函数
│       ├── auth.js                         # JWT Token管理（存储/获取/清除）
│       └── request.js                      # axios封装（BaseURL/拦截器/错误处理）
│
├── frontend/                               # Web开发调试环境（备选HTML前端）
│
├── docs/                                   # 项目文档
│   ├── ARCHITECTURE.md                     # 系统架构文档
│   ├── API_CONTRACTS.md                    # 模块间接口契约
│   ├── MODULE_1_用户画像模块.md            # 模块一开发指南
│   ├── MODULE_2_知识图谱模块.md            # 模块二开发指南
│   ├── MODULE_3_个性化推荐模块.md          # 模块三开发指南
│   ├── MODULE_4_阅读生态模块.md            # 模块四开发指南
│   └── DEVELOPMENT_GUIDE.md               # 开发环境搭建与工作流
│
├── document/                               # 设计文档（需求/概要/详细设计）
│   ├── 03-基于知识图谱的个性化荐书系统-需求说明书.docx
│   ├── 《基于知识图谱的个性化荐书系统》概要设计说明书_v3.2.docx
│   └── 详细设计/
│       ├── D04_详细设计-1.pdf              # 实训任务：详细设计要求
│       ├── D05_详细设计-2.pdf              # 实训任务：详细设计内容
│       └── 03-《升学空间站》详细设计文档_软工方向样例.docx
│
├── scripts/                                # 脚本工具
│   ├── seed_data.py                        # MySQL种子数据填充脚本
│   ├── init_neo4j.cypher                   # Neo4j约束+示例数据Cypher脚本
│   └── generate_detailed_design.py         # 本详细设计文档生成脚本
│
├── tests/                                  # 测试目录
│   ├── conftest.py                         # pytest fixtures (test_db, test_client)
│   ├── mock_data.py                        # Mock数据（模块三独立开发用）
│   ├── test_user_service.py
│   ├── test_graph_service.py
│   ├── test_recommend_service.py
│   └── test_ecosystem_service.py
│
├── docker-compose.yml                      # Docker基础设施编排（MySQL/Neo4j/Redis/ES）
├── requirements.txt                        # Python依赖清单
├── .env                                    # 环境变量配置（数据库连接/密钥/推荐参数）
└── README.md                               # 项目说明''', '图10-1 项目完整文件目录结构')

    add_body_text(doc, '目录结构设计原则：')
    add_body_text(doc, '（1）按模块分包：models/、schemas/、services/、endpoints/均按模块划分文件，每个文件顶部标注模块归属和负责人。')
    add_body_text(doc, '（2）分层解耦：API层（endpoints/）只处理HTTP请求/响应，不包含业务逻辑；服务层（services/）封装纯业务逻辑，不依赖HTTP；数据层（models/ + core/database.py）封装数据访问。')
    add_body_text(doc, '（3）接口契约驱动：schemas/中的Pydantic模型即为接口契约的代码化实现，模块间数据交换严格按Schema定义。')
    add_body_text(doc, '（4）文档即代码：docs/目录下为各模块的开发指南，document/目录下为正式的设计文档（需求→概要→详细设计）。')

    # ── 保存文档 ──
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'document', '详细设计')
    output_path = os.path.join(output_dir, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.0.docx')

    doc.save(output_path)
    print(f'Detailed design document generated: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_document()
