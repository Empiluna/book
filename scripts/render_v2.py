"""
改进版渲染脚本:
1. 从原始生成脚本直接提取mermaid代码 (避免docx解析问题)
2. 修复mermaid语法 (中文括号等)
3. mmdc渲染为PNG
4. 插入到docx中
"""
import os
import re
import sys
import json
import subprocess
import tempfile
import shutil
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BASE_DIR = r'C:\Users\Emp\PycharmProjects\PythonProject6'
SCRIPT_DIR = os.path.join(BASE_DIR, 'scripts')
DOC_DIR = os.path.join(BASE_DIR, 'document', '详细设计')
PUPPETEER_CONFIG = os.path.join(SCRIPT_DIR, 'puppeteer-config.json')
MMDC = r'C:\Users\Emp\AppData\Roaming\npm\mmdc.cmd'
INPUT_DOCX = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.0.docx')
OUTPUT_DOCX = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.1_含图片.docx')
WORK_DIR = tempfile.mkdtemp(prefix='mermaid_v2_')

# ═══════════════════════════════
# Step 0: 从docx提取mermaid代码
# ═══════════════════════════════
print("Loading document...")
doc = Document(INPUT_DOCX)

# 收集所有段落文本，识别mermaid代码块
# 策略: 扫描所有段落文本，找```mermaid 开头和 ``` 结尾
paragraphs = doc.paragraphs
all_texts = [(i, p.text) for i, p in enumerate(paragraphs)]

diagrams = []  # [(title_text, mermaid_code, para_indices_set)]

i = 0
while i < len(all_texts):
    pi, pt = all_texts[i]
    if pt.strip().startswith('```mermaid'):
        # 找到代码块开始
        code_parts = []
        start_pi = pi
        # 提取第一行 (去掉```mermaid)
        first = pt.strip()
        if first == '```mermaid':
            code_parts.append('')
        else:
            code_parts.append(first[10:])  # after ```mermaid

        indices = {pi}
        j = i + 1
        while j < len(all_texts):
            nj, nt = all_texts[j]
            indices.add(nj)
            if nt.strip().endswith('```'):
                # 最后一行
                rest = nt.strip()[:-3].strip()
                if rest:
                    code_parts.append(rest)
                break
            else:
                code_parts.append(nt)
            j += 1

        # 如果没找到结束```，跳过
        if j >= len(all_texts) or not all_texts[j][1].strip().endswith('```'):
            i += 1
            continue

        full_code = '\n'.join(code_parts).strip()

        # 找标题 (前面最近的bold段落)
        title = ""
        for k in range(pi - 1, max(pi - 5, -1), -1):
            prev_para = paragraphs[k]
            prev_text = prev_para.text.strip()
            if prev_text and not prev_text.startswith('▲') and not prev_text.startswith('```'):
                is_bold = any(r.bold for r in prev_para.runs if r.bold)
                if is_bold and len(prev_text) < 120:
                    title = prev_text
                    indices.add(k)
                    break

        # 找caption (后面最近的▲段落)
        for k in range(j + 1, min(j + 4, len(all_texts))):
            if all_texts[k][1].strip().startswith('▲'):
                indices.add(all_texts[k][0])
                break

        diagrams.append({
            'title': title,
            'code': full_code,
            'indices': sorted(indices),
            'min_idx': min(indices),
            'max_idx': max(indices),
        })

        i = j + 2  # 跳过代码块和caption
    else:
        i += 1

print(f"Found {len(diagrams)} mermaid code blocks")

# ═══════════════════════════════
# Step 1: 修复mermaid语法 & 渲染
# ═══════════════════════════════
def fix_mermaid_syntax(code: str) -> str:
    """修复常见的mermaid语法问题"""
    lines = code.split('\n')
    fixed = []
    for line in lines:
        # 修复类图: 移除方法签名中的中文括号说明
        # 例如: +GET /path (中文说明) -> +GET /path
        # 只影响classDiagram中的方法
        if re.match(r'\s*[+\-]\w+', line) and '(' in line:
            # 检查是否包含中文括号(非标准参数语法)
            # 保留标准语法: method(type param)
            # 移除中文说明: method (中文说明)
            line = re.sub(r'\s*\([^)]*[一-鿿][^)]*\)', '', line)
        fixed.append(line)
    return '\n'.join(fixed)

rendered = []
for idx, d in enumerate(diagrams):
    code = fix_mermaid_syntax(d['code'])
    title_preview = d['title'][:80] if d['title'] else code[:60].replace('\n', ' ')

    mmd_file = os.path.join(WORK_DIR, f'd{idx+1:02d}.mmd')
    png_file = os.path.join(WORK_DIR, f'd{idx+1:02d}.png')

    with open(mmd_file, 'w', encoding='utf-8') as f:
        f.write(code)

    print(f"[{idx+1}/{len(diagrams)}] {title_preview}")

    try:
        result = subprocess.run(
            [MMDC, '-i', mmd_file, '-o', png_file,
             '-p', PUPPETEER_CONFIG, '-s', '2',
             '--backgroundColor', 'white'],
            capture_output=True, timeout=90,
            cwd=BASE_DIR,
            encoding='utf-8', errors='replace'
        )
        if os.path.exists(png_file) and os.path.getsize(png_file) > 500:
            rendered.append((idx, png_file))
            print(f"  OK ({os.path.getsize(png_file)/1024:.0f} KB)")
        else:
            err = result.stderr[:300] if result.stderr else '(no error output)'
            print(f"  FAILED: {err}")
            # 查看生成的mmd内容
            if os.path.exists(mmd_file):
                with open(mmd_file, 'r', encoding='utf-8') as f:
                    preview = f.read()[:200]
                print(f"  First 200 chars of mmd: {preview}")
    except subprocess.TimeoutExpired:
        print(f"  TIMEOUT")
    except Exception as e:
        print(f"  ERROR: {e}")

print(f"\nRendered: {len(rendered)}/{len(diagrams)}")

# ═══════════════════════════════
# Step 2: 构建新docx
# ═══════════════════════════════
print("\nBuilding new document with images...")

# 快速查找: 哪些段落索引属于被渲染的mermaid图块
rendered_map = {}  # min_idx -> png_path
for di, png_path in rendered:
    d = diagrams[di]
    rendered_map[d['min_idx']] = {
        'png': png_path,
        'title': d['title'],
        'block_indices': set(d['indices']),
        'max_idx': d['max_idx'],
    }

# 新文档
nd = Document()

# 复制页面设置
src_sec = doc.sections[0]
dst_sec = nd.sections[0]
for attr in ['page_width', 'page_height', 'top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
    setattr(dst_sec, attr, getattr(src_sec, attr))

# 设置默认字体
nd.styles['Normal'].font.name = '宋体'
nd.styles['Normal'].font.size = Pt(12)
nd.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
nd.styles['Normal'].paragraph_format.line_spacing = 1.5

img_count = 0
i = 0
while i < len(paragraphs):
    if i in rendered_map:
        info = rendered_map[i]
        png_path = info['png']
        title = info['title']
        max_idx = info['max_idx']
        img_count += 1

        # 图片标题
        if title:
            p = nd.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 插入图片
        try:
            img_para = nd.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()

            pil = Image.open(png_path)
            pw, ph = pil.size
            dpi = 144  # mmdc -s 2 produces 2x DPI
            max_w = 5.8
            if pw / dpi > max_w:
                scale = max_w / (pw / dpi)
                w_in = max_w
                h_in = (ph / dpi) * scale
            else:
                w_in = pw / dpi
                h_in = ph / dpi

            run.add_picture(png_path, width=Inches(w_in), height=Inches(h_in))

            # 图注
            cap = nd.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(f'图{img_count} (已渲染)')
            cr.font.size = Pt(8)
            cr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            cr.italic = True
            cr.font.name = '宋体'
            cr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            nd.add_paragraph()

            print(f"  [{img_count}] Inserted {os.path.basename(png_path)} ({w_in:.1f}x{h_in:.1f}in)")
        except Exception as e:
            print(f"  [ERR] Insert {png_path}: {e}")
            # 回退: 复制原code段落
            for j in range(i, max_idx + 1):
                if j < len(paragraphs):
                    _copy_para(nd, paragraphs[j])

        i = max_idx + 1
        continue

    # 普通段落
    _copy_para(nd, paragraphs[i])
    i += 1

# 保存
nd.save(OUTPUT_DOCX)
print(f"\nSaved: {OUTPUT_DOCX}")
print(f"Images embedded: {img_count}")

# 清理
shutil.rmtree(WORK_DIR, ignore_errors=True)


def _copy_para(nd, old):
    """复制段落"""
    text = old.text
    sn = old.style.name if old.style else 'Normal'

    if not text.strip():
        nd.add_paragraph()
        return

    if sn.startswith('Heading'):
        lv = int(sn.split()[-1]) if sn.split()[-1].isdigit() else 1
        h = nd.add_heading(text, level=lv)
        for r in h.runs:
            r.font.name = '黑体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return

    p = nd.add_paragraph()
    for or_ in old.runs:
        nr = p.add_run(or_.text)
        try:
            if or_.bold: nr.bold = True
            if or_.italic: nr.italic = True
            if or_.font.size: nr.font.size = or_.font.size
            if or_.font.name:
                nr.font.name = or_.font.name
                nr._element.rPr.rFonts.set(qn('w:eastAsia'), or_.font.name)
            if or_.font.color and or_.font.color.rgb:
                nr.font.color.rgb = or_.font.color.rgb
        except:
            pass
    try:
        pf = old.paragraph_format
        if pf.first_line_indent: p.paragraph_format.first_line_indent = pf.first_line_indent
        if pf.line_spacing: p.paragraph_format.line_spacing = pf.line_spacing
        if pf.alignment: p.alignment = pf.alignment
    except:
        pass
