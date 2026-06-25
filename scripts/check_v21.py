import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'C:\Users\Emp\PycharmProjects\PythonProject6\document\详细设计\《基于知识图谱的个性化荐书系统》详细设计说明书_v2.1_最终版.docx')

headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
h1 = len([h for h in headings if h.style.name == 'Heading 1'])
h2 = len([h for h in headings if h.style.name == 'Heading 2'])
h3 = len([h for h in headings if h.style.name == 'Heading 3'])
h4 = len([h for h in headings if h.style.name == 'Heading 4'])
print(f'Headings: {len(headings)} (H1:{h1} H2:{h2} H3:{h3} H4:{h4})')
print(f'Tables: {len(doc.tables)}')

blips = doc.element.findall('.//' + qn('a:blip'))
print(f'Images embedded: {len(blips)}')

mmd = 0
for p in doc.paragraphs:
    if p.text.strip().startswith('```mermaid'):
        mmd += 1
print(f'Mermaid code fallbacks: {mmd}')
print()

for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        print(f'  {p.text[:80]}')
