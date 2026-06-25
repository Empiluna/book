"""
从详细设计.docx中提取Mermaid图表 → mmdc渲染为PNG → 生成带图片的新docx
"""
import os
import re
import sys
import json
import subprocess
import tempfile
import shutil
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BASE_DIR = r'C:\Users\Emp\PycharmProjects\PythonProject6'
SCRIPT_DIR = os.path.join(BASE_DIR, 'scripts')
DOC_DIR = os.path.join(BASE_DIR, 'document', '详细设计')
PUPPETEER_CONFIG = os.path.join(SCRIPT_DIR, 'puppeteer-config.json')
INPUT_DOCX = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.0.docx')
OUTPUT_DOCX = os.path.join(DOC_DIR, '《基于知识图谱的个性化荐书系统》详细设计说明书_v1.1_含图片.docx')
WORK_DIR = tempfile.mkdtemp(prefix='mermaid_render_')

print(f"Work directory: {WORK_DIR}")

# ═══════════════════════════════════════
# Step 1: 从文档中提取Mermaid代码块
# ═══════════════════════════════════════
print("\n[1/4] Extracting Mermaid diagrams from document...")
doc = Document(INPUT_DOCX)

# 分析文档结构：识别Mermaid图块
# 每个图块 = (title段落idx, code段落idx列表, caption段落idx, diagram_code, title_text)
diagram_blocks = []
paragraphs = doc.paragraphs

i = 0
while i < len(paragraphs):
    text = paragraphs[i].text.strip()
    if text.startswith('```mermaid'):
        # 找到代码块起始
        code_start = i
        code_paragraphs = [i]
        code_text = text

        # 向前查找后续的代码段落
        j = i + 1
        while j < len(paragraphs):
            next_text = paragraphs[j].text
            code_paragraphs.append(j)
            code_text += '\n' + next_text
            if '```' in next_text and j > code_start:
                break
            j += 1
        code_end = j + 1

        # 向前查找标题
        title_idx = None
        title_text = ""
        if i > 0:
            prev = paragraphs[i - 1]
            prev_text = prev.text.strip()
            if prev_text and not prev_text.startswith('▲') and not prev_text.startswith('```'):
                # 检查是否是bold（很可能是图标题）
                is_bold = any(r.bold for r in prev.runs if r.bold)
                if is_bold and len(prev_text) < 100:
                    title_idx = i - 1
                    title_text = prev_text

        # 向后查找caption
        caption_idx = None
        if code_end < len(paragraphs):
            cap_text = paragraphs[code_end].text.strip()
            if cap_text.startswith('▲'):
                caption_idx = code_end

        block = {
            'title_idx': title_idx,
            'code_indices': code_paragraphs,
            'caption_idx': caption_idx,
            'code': code_text,
            'title': title_text,
        }
        diagram_blocks.append(block)
        i = code_end + 1 if caption_idx else code_end
    else:
        i += 1

print(f"  Found {len(diagram_blocks)} Mermaid diagrams")

# ═══════════════════════════════════════
# Step 2: 渲染每个图为PNG
# ═══════════════════════════════════════
print("\n[2/4] Rendering diagrams with mmdc...")
rendered = []  # [(block_index, img_path), ...]

for idx, block in enumerate(diagram_blocks):
    code = block['code']
    # 清理代码：去掉 ```mermaid 和 ``` 标记
    clean_code = code.strip()
    if clean_code.startswith('```mermaid'):
        clean_code = clean_code[10:]
    if clean_code.endswith('```'):
        clean_code = clean_code[:-3]
    clean_code = clean_code.strip()

    mmd_path = os.path.join(WORK_DIR, f'diagram_{idx+1:02d}.mmd')
    png_path = os.path.join(WORK_DIR, f'diagram_{idx+1:02d}.png')

    with open(mmd_path, 'w', encoding='utf-8') as f:
        f.write(clean_code)

    title_preview = block['title'][:60] if block['title'] else clean_code[:60].replace('\n', ' ')
    print(f"  [{idx+1}/{len(diagram_blocks)}] {title_preview}...")

    try:
        result = subprocess.run(
            [r'C:\Users\Emp\AppData\Roaming\npm\mmdc.cmd', '-i', mmd_path, '-o', png_path,
             '-p', PUPPETEER_CONFIG,
             '-s', '2',  # scale factor
             '--backgroundColor', 'white'],
            capture_output=True, text=True, timeout=60,
            cwd=BASE_DIR
        )
        if os.path.exists(png_path) and os.path.getsize(png_path) > 0:
            rendered.append((idx, png_path))
            size_kb = os.path.getsize(png_path) / 1024
            print(f"    OK ({size_kb:.0f} KB)")
        else:
            print(f"    FAILED: {result.stderr[:200] if result.stderr else 'No output'}")
    except subprocess.TimeoutExpired:
        print(f"    TIMEOUT")
    except Exception as e:
        print(f"    ERROR: {e}")

print(f"\n  Successfully rendered: {len(rendered)}/{len(diagram_blocks)}")

if not rendered:
    print("No diagrams rendered! Aborting.")
    sys.exit(1)

# ═══════════════════════════════════════
# Step 3: 生成含图片的新文档
# ═══════════════════════════════════════
print("\n[3/4] Building new document with images...")

# 构建段落索引到图片的映射
block_info = {}  # title_idx -> (png_path, caption_text)
for idx, png_path in rendered:
    block = diagram_blocks[idx]
    ti = block['title_idx']
    caption = ""
    if block['caption_idx'] is not None:
        caption = paragraphs[block['caption_idx']].text
    # 需要删除的段落索引集合
    removed = set()
    if block['title_idx'] is not None:
        removed.add(block['title_idx'])
    removed.update(block['code_indices'])
    if block['caption_idx'] is not None:
        removed.add(block['caption_idx'])

    # 用最小的index作为插入点
    insert_at = min(removed)
    block_info[insert_at] = {
        'png_path': png_path,
        'title': block['title'],
        'caption': caption,
        'removed': removed,
    }

# 创建新文档
new_doc = Document()

# 复制页面设置
src_section = doc.sections[0]
dst_section = new_doc.sections[0]
dst_section.page_width = src_section.page_width
dst_section.page_height = src_section.page_height
dst_section.top_margin = src_section.top_margin
dst_section.bottom_margin = src_section.bottom_margin
dst_section.left_margin = src_section.left_margin
dst_section.right_margin = src_section.right_margin

# 设置默认样式
style = new_doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 逐段落处理
diagram_counter = 0
i = 0
while i < len(paragraphs):
    if i in block_info:
        info = block_info[i]
        png_path = info['png_path']
        title = info['title']
        caption = info['caption']
        removed = info['removed']
        diagram_counter += 1

        # 插入图片标题
        if title:
            p = new_doc.add_paragraph()
            run = p.add_run(f'图{diagram_counter}  {title.replace("图3-1", "").replace("图4-1", "").replace("图4-2", "").replace("图5-1", "").replace("图6-1", "").replace("图6-2", "").replace("图6-3", "").replace("图6-4", "").replace("图6-5", "").replace("图6-6", "").replace("图6-7", "").replace("图6-8", "").replace("图6-9", "").replace("图6-10", "").replace("图10-1", "").strip()}')
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 插入图片
        try:
            img_para = new_doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()

            # 计算图片尺寸：保持宽高比，最大宽6英寸
            from PIL import Image
            pil_img = Image.open(png_path)
            img_w, img_h = pil_img.size
            max_w_inches = 5.8
            if img_w / 96 > max_w_inches:  # 96 DPI assumption
                scale = max_w_inches / (img_w / 96)
                w_inches = max_w_inches
                h_inches = (img_h / 96) * scale
            else:
                w_inches = img_w / 96
                h_inches = img_h / 96

            run.add_picture(png_path, width=Inches(w_inches), height=Inches(h_inches))

            # 图注
            if caption:
                cap_para = new_doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption.replace('▲', '').strip())
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                cap_run.italic = True
                cap_run.font.name = '宋体'
                cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 图片后间距
            new_doc.add_paragraph()
            print(f"  Inserted image: {os.path.basename(png_path)} ({w_inches:.1f}x{h_inches:.1f} inches)")
        except Exception as e:
            print(f"  [ERROR] Failed to insert {png_path}: {e}")

        # 跳过已处理的段落
        i = max(removed) + 1
        continue

    # 复制普通段落
    para = paragraphs[i]
    _copy_paragraph_to_doc(new_doc, para)
    i += 1

# ═══════════════════════════════════════
# Step 4: 复制表格
# ═══════════════════════════════════════
print("\n[4/4] Copying tables...")
# 表格通过element复制（跨文档）
for table in doc.tables:
    try:
        # 在文档末尾添加表格
        new_tbl = OxmlElement('w:tbl')
        # 深拷贝表格XML
        tbl_element = table._tbl
        # 简单的复制: 添加一个空段落然后在其后插入表格
        new_doc.add_paragraph()
        last_para = new_doc.paragraphs[-1]
        last_para._element.addnext(tbl_element)
    except Exception as e:
        print(f"  [WARN] Could not copy table: {e}")

print(f"\nSaving: {OUTPUT_DOCX}")
new_doc.save(OUTPUT_DOCX)

# 清理
shutil.rmtree(WORK_DIR, ignore_errors=True)

print(f"\n{'='*60}")
print(f"Done! Generated document with {diagram_counter} rendered diagrams.")
print(f"Output: {OUTPUT_DOCX}")
print(f"{'='*60}")


def _copy_paragraph_to_doc(new_doc, old_para):
    """复制段落到新文档，尽可能保持格式"""
    text = old_para.text
    style_name = old_para.style.name if old_para.style else 'Normal'

    if not text.strip():
        new_doc.add_paragraph()
        return

    if style_name.startswith('Heading'):
        # 标题
        try:
            level = int(style_name.split()[-1])
        except:
            level = 1
        h = new_doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:
        p = new_doc.add_paragraph()
        # 复制runs
        for old_run in old_para.runs:
            new_run = p.add_run(old_run.text)
            try:
                if old_run.bold: new_run.bold = True
                if old_run.italic: new_run.italic = True
                if old_run.font.size: new_run.font.size = old_run.font.size
                if old_run.font.name:
                    new_run.font.name = old_run.font.name
                if old_run.font.color and old_run.font.color.rgb:
                    new_run.font.color.rgb = old_run.font.color.rgb
            except:
                pass

        # 复制段落格式
        try:
            pf = old_para.paragraph_format
            if pf.first_line_indent:
                p.paragraph_format.first_line_indent = pf.first_line_indent
            if pf.line_spacing:
                p.paragraph_format.line_spacing = pf.line_spacing
            if pf.alignment:
                p.alignment = pf.alignment
        except:
            pass
